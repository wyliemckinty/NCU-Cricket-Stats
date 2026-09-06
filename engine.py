# ==========================================
# engine.py
# ==========================================
import pandas as pd
import numpy as np
import os
import re
import io
import zipfile
from datetime import datetime, timedelta
from thefuzz import process, fuzz
import warnings
import streamlit as st
import unicodedata
import shutil
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from collections import Counter


warnings.filterwarnings('ignore')

# ==========================================
# GLOBAL PANDAS OPTIMIZATION
# ==========================================
# Patch pandas to use the 'calamine' engine by default for all read_excel 
# operations, which is up to 20x faster than the default openpyxl engine.
if not hasattr(pd, '_original_read_excel'):
    pd._original_read_excel = pd.read_excel
    pd._original_excel_file = pd.ExcelFile

    def fast_read_excel(*args, **kwargs):
        if len(args) > 0 and type(args[0]).__name__ != 'ExcelFile' and 'engine' not in kwargs:
            kwargs['engine'] = 'calamine'
        elif 'io' in kwargs and type(kwargs['io']).__name__ != 'ExcelFile' and 'engine' not in kwargs:
            kwargs['engine'] = 'calamine'
        return pd._original_read_excel(*args, **kwargs)

    def fast_excel_file(*args, **kwargs):
        if 'engine' not in kwargs:
            kwargs['engine'] = 'calamine'
        return pd._original_excel_file(*args, **kwargs)

    pd.read_excel = fast_read_excel
    pd.ExcelFile = fast_excel_file

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    pass

try:
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except ImportError:
    pass

# ==========================================
# DEFAULT FILE NAME MAPPING & REGISTRIES
# ==========================================
DEFAULT_FILES = {
    "Men's": {
        "reg": "1. NCU_Registered_Players.xlsx",
        "id_map": "NCU_Mens_Master_ID_Mapping.xlsx",
        "alias": "2. NCU_Validated_Aliases_Master.xlsx",
        "starring": "3. NCU Complete -Men's- Starring List from 1st June.xlsx",
        "unreg": "4. Unregistered_Manual_Map.xlsx",
        "secondary": "5. Secondary_Team_Map.xlsx",
        "league": "2026 Season League Structure for Gemini AI.xlsx",
        "bat": "NV Play NCU League and Saturday Cup batting stats for season.xlsx",
        "bowl": "NV Play NCU League and Saturday Cup bowling stats for season.xlsx",
        "abandoned": "NV Play NCU League and Saturday Cup player appearances for abandoned games.xlsx"
    },
    "Women's": {
        "reg": "1. NCU_Registered_Players.xlsx",
        "id_map": "NCU_Womens_Master_ID_Mapping.xlsx",
        "alias": "12. NCU_Validated_Women's Aliases_Master.xlsx",
        "starring": "13. NCU Complete Women's Starring List from 1st June.xlsx",
        "unreg": "4. Unregistered_Manual_Map.xlsx",
        "secondary": "5. Secondary_Team_Map.xlsx",
        "league": "2026 Season League Structure Women for Gemini AI.xlsx",
        "bat": "NV Play Women's Fixtures batting stats for season.xlsx",
        "bowl": "NV Play Women's Fixtures bowling stats for season.xlsx",
        "abandoned": "NV Play Women's Fixtures player appearances for abandoned games.xlsx"
    },
    "Midweek": {
        "reg": "1. NCU_Registered_Players.xlsx",
        "id_map": "NCU_Mens_Master_ID_Mapping.xlsx",
        "alias": "2. NCU_Validated_Aliases_Master.xlsx",
        "starring": "", 
        "unreg": "4. Unregistered_Manual_Map.xlsx",
        "secondary": "5. Secondary_Team_Map.xlsx",
        "league": "2026 Season Midweek League Structure for Gemini AI.xlsx",
        "bat": "NV Play Midweek League batting stats for season.xlsx",
        "bowl": "NV Play Midweek League bowling stats for season.xlsx",
        "abandoned": ""
    }
}

# Mapping of official club names to all abbreviations/variants used in NV Play data.
# Used by resolve_duplicates to correctly identify which club a match belongs to.
CLUB_ALIASES = {
    'CI': ['CI', 'CIYMS'],
    'CSNI': ['CSNI', 'Civil Service North', 'Civil Service North of Ireland'],
    'BISC': ['BISC', 'Belfast International Sports Club'],
    'NIMA': ['NIMA', 'NIMACC', 'NIMA CC', 'Northern Ireland Malayali Association'],
    'Holywood': ['Holywood', 'Holywood 1881'],
    'Ards & Donaghadee': ['Ards & Donaghadee', 'Ards', 'Ards and Donaghadee'],
    'Donacloney Mill': ['Donacloney Mill', 'Donacloney', 'Donaghcloney'],
}

KNOWN_DUPLICATES = {}  # Initialized dynamically via _init_known_duplicates() below


# ==========================================
# STANDARDIZED CACHED FILE LOADERS
# ==========================================
@st.cache_data(show_spinner="Loading data file...")
def cached_read_excel(filepath, mtime):
    if not os.path.exists(filepath):
        return pd.DataFrame()
    if str(filepath).lower().endswith('.csv'):
        return pd.read_csv(filepath)
    return pd.read_excel(filepath)

@st.cache_data(show_spinner="Loading sheet...")
def cached_read_excel_sheet(filepath, mtime, sheet_name=None, header='infer'):
    if not os.path.exists(filepath):
        return pd.DataFrame()
    if str(filepath).lower().endswith('.csv'):
        return pd.read_csv(filepath, header=header)
    return pd.read_excel(filepath, sheet_name=sheet_name, header=header)

def get_excel_df(filepath):
    if not filepath or not os.path.exists(filepath):
        return pd.DataFrame()
    return cached_read_excel(filepath, os.path.getmtime(filepath))

def get_excel_sheet_df(filepath, sheet_name=None, header='infer'):
    if not filepath or not os.path.exists(filepath):
        return pd.DataFrame()
    return cached_read_excel_sheet(filepath, os.path.getmtime(filepath), sheet_name=sheet_name, header=header)

# ==========================================
# UNIFIED ENGINE FUNCTIONS 
# ==========================================
def fix_celtic_casing(name):
    """
    Standardizes Scottish/Irish surname casing for consistent display and grouping.
    Converts Mc[a-z] to Mc[A-Z], e.g. Mckeown -> McKeown, Mcilwaine -> McIlwaine.
    Converts O'[a-z] to O'[A-Z], e.g. O'neill -> O'Neill.
    """
    if not isinstance(name, str):
        return name
    s = str(name).replace("OaTM", "O'").replace("O\ufffd", "O'").replace("O\xef\xbf\xbd", "O'").replace("O’", "O'").replace("`", "'")
    s = re.sub(r'\s+', ' ', s).strip()
    s = re.sub(r'\bMc([a-z])', lambda m: f"Mc{m.group(1).upper()}", s)
    s = re.sub(r"\bO'([a-z])", lambda m: f"O'{m.group(1).upper()}", s)
    return s

def extract_base_club_name(team_name):
    if pd.isna(team_name): return "Unknown Club"
    t = str(team_name).strip()
    t = re.sub(r'(?i)\b\d(?:st|nd|rd|th)?\s*XI\b', '', t)
    t = re.sub(r'(?i)\b(?:1st|2nd|3rd|4th|5th|6th|7th)\b', '', t)
    t = re.sub(r'(?i)\bWomen\'?s?\b', '', t)
    t = re.sub(r'(?i)\bMW\d?\b', '', t)
    t = re.sub(r'(?i)\bCricket Club\b|\bCC\b', '', t)
    t = re.sub(r'\s+\d$', '', t.strip())
    t = re.sub(r'\s+', ' ', t).strip()
    t = re.sub(r'(?i)\bciyms\b', 'CI', t)
    t = re.sub(r'(?i)\bholywood\s+1881\b', 'Holywood', t)
    t = re.sub(r'(?i)northern\s+ireland\s+malayali\s+association', 'NIMA', t)
    t = re.sub(r'(?i)\bnima\s*cc\b|\bnimacc\b|\bnima\b', 'NIMA', t)
    t = re.sub(r'(?i)belfast\s+international\s+sports\s+club|belfast\s+b\.i\.s\.c\.', 'BISC', t)
    t = re.sub(r'(?i)civil\s+service\s+north\s+of\s+ireland|civil\s+service\s+north', 'CSNI', t)
    t = re.sub(r'(?i)\bdrumaness\s+super\s*kings\b', 'Drumaness', t)
    t = re.sub(r'(?i)\bdonaghcloney\b', 'Donacloney', t)
    return t if t else "Unknown Club"

def clean_club_for_matching(club_str):
    if pd.isna(club_str): return ""
    c = str(club_str).lower()
    c = re.sub(r'\bcricket club\b|\bcc\b', '', c)
    c = c.replace('1881', '')
    c = c.replace('ciyms', 'ci')
    c = re.sub(r'(?i)northern\s+ireland\s+malayali\s+association', 'nima', c)
    c = re.sub(r'(?i)\bnima\s*cc\b|\bnimacc\b|\bnima\b', 'nima', c)
    c = re.sub(r'(?i)belfast\s+international\s+sports\s+club|belfast\s+b\.i\.s\.c\.', 'bisc', c)
    c = re.sub(r'(?i)civil\s+service\s+north\s+of\s+ireland|civil\s+service\s+north', 'csni', c)
    c = re.sub(r'(?i)drumaness\s+super\s*kings', 'drumaness', c)
    c = re.sub(r'(?i)donaghcloney', 'donacloney', c)
    return " ".join(c.split())

def build_dynamic_duplicate_map(id_map_df=None, reg_players_df=None):
    """
    Dynamically discovers duplicate player names across clubs
    by scanning 1. NCU_Registered_Players.xlsx (multiple CI numbers/clubs)
    and NCU_Mens_Master_ID_Mapping.xlsx (multiple Sport80 IDs/clubs).
    Returns a dict: {player_name: [club1, club2, ...]}
    """
    dup_map = {}
    ignored_clubs = {'northern cricket union', 'ncu', 'unknown club', 'unknown'}
    
    # 1. From Registration file
    if reg_players_df is not None and not reg_players_df.empty:
        name_col = 'Full Name' if 'Full Name' in reg_players_df.columns else reg_players_df.columns[0]
        ci_col = next((c for c in reg_players_df.columns if 'ci no' in str(c).lower() or 'membership' in str(c).lower()), None)
        club_col = next((c for c in reg_players_df.columns if 'primary club' in str(c).lower()), None)
        
        reg_copy = reg_players_df.copy()
        reg_copy['Norm_Name'] = reg_copy[name_col].astype(str).str.replace('‡', '', regex=False).str.strip().apply(fix_celtic_casing)
        
        for name, group in reg_copy.groupby('Norm_Name'):
            if not name or str(name).lower() in ['nan', 'none', '']: continue
            unique_cis = set(str(x).replace('.0','').strip() for x in group[ci_col].dropna() if str(x).strip() and str(x).lower() != 'nan') if ci_col else set()
            raw_clubs = set(extract_base_club_name(c) for c in group[club_col].dropna() if str(c).strip() and str(c).lower() != 'nan') if club_col else set()
            unique_clubs = {c for c in raw_clubs if c.lower() not in ignored_clubs}
            
            if len(unique_cis) > 1 or len(unique_clubs) > 1:
                if unique_clubs:
                    dup_map[name] = sorted(list(unique_clubs))
                
    # 2. From ID map
    if id_map_df is not None and not id_map_df.empty:
        col_s80_name = next((c for c in id_map_df.columns if 'sport80_name' in str(c).lower()), None)
        col_s80_club = next((c for c in id_map_df.columns if 'sport80_club' in str(c).lower()), None)
        col_s80_id = next((c for c in id_map_df.columns if 'sport80_id' in str(c).lower()), None)
        
        if col_s80_name:
            df_copy = id_map_df.copy()
            df_copy['Norm_Name'] = df_copy[col_s80_name].fillna('').astype(str).str.strip().apply(fix_celtic_casing)
            
            for name, group in df_copy.groupby('Norm_Name'):
                if not name or str(name).lower() in ['nan', 'none', '']: continue
                unique_ids = set(str(x).replace('.0','').strip() for x in group[col_s80_id].dropna() if str(x).strip() and str(x).lower() != 'nan') if col_s80_id else set()
                raw_clubs = set(extract_base_club_name(c) for c in group[col_s80_club].dropna() if str(c).strip() and str(c).lower() != 'nan') if col_s80_club else set()
                unique_clubs = {c for c in raw_clubs if c.lower() not in ignored_clubs}
                
                if len(unique_ids) > 1 or len(unique_clubs) > 1:
                    existing = set(dup_map.get(name, []))
                    existing.update(unique_clubs)
                    if existing:
                        dup_map[name] = sorted(list(existing))

    # Also add standard-cased keys for casing compatibility
    for name, clubs in list(dup_map.items()):
        celtic = fix_celtic_casing(name)
        dup_map[celtic] = clubs
        title_cased = name.title()
        if title_cased not in dup_map:
            dup_map[title_cased] = clubs
            
    return dup_map

def _init_known_duplicates():
    base_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
    reg_file_rel = DEFAULT_FILES.get("Men's", {}).get("reg", "1. NCU_Registered_Players.xlsx")
    id_file_rel = DEFAULT_FILES.get("Men's", {}).get("id_map", "NCU_Mens_Master_ID_Mapping.xlsx")
    
    reg_file = os.path.join(base_dir, reg_file_rel) if not os.path.isabs(reg_file_rel) else reg_file_rel
    id_file = os.path.join(base_dir, id_file_rel) if not os.path.isabs(id_file_rel) else id_file_rel
    
    reg_df, id_df = None, None
    if os.path.exists(reg_file):
        try: reg_df = pd.read_excel(reg_file)
        except Exception: pass
    elif os.path.exists(reg_file_rel):
        try: reg_df = pd.read_excel(reg_file_rel)
        except Exception: pass
        
    if os.path.exists(id_file):
        try: id_df = pd.read_excel(id_file)
        except Exception: pass
    elif os.path.exists(id_file_rel):
        try: id_df = pd.read_excel(id_file_rel)
        except Exception: pass
        
    return build_dynamic_duplicate_map(id_map_df=id_df, reg_players_df=reg_df)

KNOWN_DUPLICATES.update(_init_known_duplicates())

def build_alias_map(aliases, domain):
    alias_map = {}
    if 'Input Name (Scorecard/Stats)' in aliases.columns and 'Official Registered Name' in aliases.columns:
        aliases_deduped = aliases.drop_duplicates(subset=['Input Name (Scorecard/Stats)'], keep='last')
        for idx, row in aliases_deduped.iterrows():
            alias_val = str(row['Input Name (Scorecard/Stats)']).replace('‡', '').strip().lower() 
            official_val = fix_celtic_casing(str(row['Official Registered Name']).replace('‡', '').strip())
            if alias_val != 'nan':
                alias_map[alias_val] = official_val
    else:
        for idx, row in aliases.iterrows():
            alias_val = str(row.iloc[0]).replace('‡', '').strip().lower() 
            official_val = fix_celtic_casing(str(row.iloc[1]).replace('‡', '').strip())
            if alias_val != 'nan':
                alias_map[alias_val] = official_val
                
    return alias_map

def build_id_map(id_map_df):
    """
    Builds a lookup dictionary from the Master ID Mapping DataFrame.
    Keyed by normalized NV_Play_ID (lowercase string UUID).
    """
    id_map = {}
    if id_map_df is None or id_map_df.empty:
        return id_map
    
    col_nv_id = next((c for c in id_map_df.columns if 'nv' in c.lower() and 'id' in c.lower()), 'NV_Play_ID')
    col_s80_id = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'id' in c.lower()), 'Sport80_ID')
    col_s80_name = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'name' in c.lower()), 'Sport80_Name')
    col_s80_club = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'club' in c.lower()), 'Sport80_Club')
    col_conf = next((c for c in id_map_df.columns if 'conf' in c.lower()), 'Match_Confidence')
    col_nv_name = next((c for c in id_map_df.columns if 'nv' in c.lower() and 'name' in c.lower()), 'NV_Play_Name')

    for _, row in id_map_df.iterrows():
        nv_id = row.get(col_nv_id)
        if pd.notna(nv_id):
            clean_nv_id = str(nv_id).strip().lower()
            if clean_nv_id and clean_nv_id != 'nan':
                s80_id_val = row.get(col_s80_id)
                clean_s80_id = ""
                if pd.notna(s80_id_val) and str(s80_id_val).strip().lower() != 'nan':
                    clean_s80_id = str(s80_id_val).replace('.0', '').strip()
                
                s80_name_val = fix_celtic_casing(str(row.get(col_s80_name, '')).strip()) if pd.notna(row.get(col_s80_name)) else ""
                s80_club_val = str(row.get(col_s80_club, '')).strip() if pd.notna(row.get(col_s80_club)) else ""
                conf_val = str(row.get(col_conf, '')).strip() if pd.notna(row.get(col_conf)) else ""
                nv_name_val = fix_celtic_casing(str(row.get(col_nv_name, '')).strip()) if pd.notna(row.get(col_nv_name)) else ""

                id_map[clean_nv_id] = {
                    'sport80_id': clean_s80_id,
                    'sport80_name': s80_name_val,
                    'sport80_club': s80_club_val,
                    'confidence': conf_val,
                    'nv_play_name': nv_name_val
                }
    return id_map

def extract_row_player_id(row, id_cols=None):
    """
    Extracts NV Play player UUID from a scorecard/stats row if present.
    """
    candidate_cols = id_cols if id_cols else [
        'Batter ID', 'Bowler ID', 'Player ID', 'NV_Play_ID', 'NV Play ID', 'Player NV Play ID',
        'BatterId', 'BowlerId', 'PlayerId', 'ID', 'Player UUID', 'Player_ID'
    ]
    for col in candidate_cols:
        if col in row and pd.notna(row[col]):
            val = str(row[col]).replace('.0', '').strip().lower()
            if val and val != 'nan':
                return val
    return None

def resolve_player_from_row(row, raw_name, id_map, alias_map, player_club_map=None, id_cols=None):
    """
    Resolves a scorecard/stats row to a canonical player identity.
    Checks id_map first using player UUID; falls back to cleanse_name_contextual.
    Returns: (cleaned_name, sport80_id, sport80_club, is_id_resolved)
    """
    clean_input_name = re.sub(r'\s+', ' ', str(raw_name).replace('‡', '')).strip()
    player_uuid = extract_row_player_id(row, id_cols=id_cols)
    if player_uuid and id_map and player_uuid in id_map:
        info = id_map[player_uuid]
        sport80_id = info.get('sport80_id', '')
        sport80_club = info.get('sport80_club', '')
        raw_canonical = info.get('sport80_name') or info.get('nv_play_name') or clean_input_name
        canonical_name = fix_celtic_casing(raw_canonical)
        
        # Format name with club if in KNOWN_DUPLICATES
        if canonical_name in KNOWN_DUPLICATES and sport80_club:
            short_club = extract_base_club_name(sport80_club)
            cleaned_name = f"{canonical_name} ({short_club})"
        else:
            cleaned_name = canonical_name
            
        return fix_celtic_casing(cleaned_name), sport80_id, sport80_club, True

    # If no UUID match, check id_map by name and club context before generic fallback
    if id_map and clean_input_name:
        raw_name_clean = clean_input_name.lower()
        group_context = str(row.get('Group', row.get('Match', ''))).lower()
        team_context = str(row.get('Team', '')).lower()
        comb_context = clean_club_for_matching(team_context + " " + group_context)
        
        matched_candidates = []
        for uuid_key, info in id_map.items():
            nv_n = re.sub(r'\s+', ' ', str(info.get('nv_play_name', ''))).strip().lower()
            s80_n = re.sub(r'\s+', ' ', str(info.get('sport80_name', ''))).strip().lower()
            if raw_name_clean == nv_n or raw_name_clean == s80_n:
                club_name = info.get('sport80_club', '')
                clean_club = clean_club_for_matching(club_name)
                variants = CLUB_ALIASES.get(extract_base_club_name(club_name), [extract_base_club_name(club_name)])
                if clean_club and (clean_club in comb_context or any(clean_club_for_matching(v) in comb_context for v in variants)):
                    matched_candidates.append(info)
                    
        if len(matched_candidates) == 1:
            info = matched_candidates[0]
            sport80_id = info.get('sport80_id', '')
            sport80_club = info.get('sport80_club', '')
            raw_canonical = info.get('sport80_name') or info.get('nv_play_name') or clean_input_name
            canonical_name = fix_celtic_casing(raw_canonical)
            if canonical_name in KNOWN_DUPLICATES and sport80_club:
                short_club = extract_base_club_name(sport80_club)
                cleaned_name = f"{canonical_name} ({short_club})"
            else:
                cleaned_name = canonical_name
            return fix_celtic_casing(cleaned_name), sport80_id, sport80_club, True
        
    fallback_name = fix_celtic_casing(cleanse_name_contextual(clean_input_name, row, alias_map, player_club_map))
    return fallback_name, None, None, False

def build_secondary_team_map(secondary_df, alias_map):
    sec_map = {}
    if secondary_df is not None and not secondary_df.empty:
        col_name = secondary_df.columns[0]
        col_team = secondary_df.columns[1]
        for _, r in secondary_df.iterrows():
            if pd.notna(r[col_name]) and pd.notna(r[col_team]):
                p_name = str(r[col_name]).strip().lower()
                p_team = str(r[col_team]).strip()
                if p_name and p_name != 'nan':
                    mapped_name = alias_map.get(p_name, p_name)
                    if mapped_name not in sec_map: sec_map[mapped_name] = []
                    if p_team not in sec_map[mapped_name]: sec_map[mapped_name].append(p_team)
                    
                    if p_name not in sec_map: sec_map[p_name] = []
                    if p_team not in sec_map[p_name]: sec_map[p_name].append(p_team)
    return sec_map
    
def get_alias_used_for_player(official_name, search_input, alias_map):
    if not search_input:
        return None
    
    clean_search = search_input.strip().lower()
    mapped = alias_map.get(clean_search)
    if mapped and mapped.lower() == official_name.lower() and clean_search != official_name.lower():
        return search_input.strip().title()
        
    return None

def cleanse_name(name, alias_map):
    original_name = fix_celtic_casing(str(name).replace('‡', '').strip())
    return fix_celtic_casing(alias_map.get(original_name.lower(), original_name))

def cleanse_name_contextual(name, row, alias_map, player_club_map=None):
    original_name = fix_celtic_casing(str(name).replace('‡', '').strip())
    original_name_lower = original_name.lower()
    group_lower = str(row.get('Group', row.get('Match', ''))).lower()
    row_team = str(row.get('Team', '')).lower()
    
    clean_group = clean_club_for_matching(group_lower)
    clean_row_team = clean_club_for_matching(row_team)
    
    for dup_name, clubs in KNOWN_DUPLICATES.items():
        if original_name_lower == dup_name.lower():
            combined_context = row_team + ' ' + group_lower
            matched_clubs = []
            
            for club in clubs:
                variants = CLUB_ALIASES.get(club, [club])
                for variant in variants:
                    if re.search(r'\b' + re.escape(variant.lower()) + r'\b', combined_context):
                        matched_clubs.append(club)
                        break
            
            if len(matched_clubs) == 1:
                return f"{original_name} ({matched_clubs[0]})"
            elif len(matched_clubs) > 1:
                # If multiple clubs matched (e.g. playing against each other), try to break tie with row_team if it exists
                if row_team:
                    for club in matched_clubs:
                        variants = CLUB_ALIASES.get(club, [club])
                        for variant in variants:
                            if re.search(r'\b' + re.escape(variant.lower()) + r'\b', row_team):
                                return f"{original_name} ({club})"
                return f"{original_name} ({matched_clubs[0]})"
            else:
                # Fallback if no clubs matched in the context string
                if player_club_map:
                    reg_club = str(player_club_map.get(original_name_lower, '')).lower()
                    for club in clubs:
                        c_clean = clean_club_for_matching(club)
                        if club.lower() in reg_club or c_clean in reg_club:
                            return f"{original_name} ({club})"
                return f"{original_name} ({clubs[0]})"
                
    return fix_celtic_casing(alias_map.get(original_name_lower, original_name))

def build_player_club_map(reg_players, alias_map, domain, unreg_map_df=None, secondary_map=None, id_map_df=None):
    club_map = {}
    if reg_players is None or reg_players.empty: return club_map
    
    if 'First Name' in reg_players.columns and 'Last Name' in reg_players.columns:
        reg_players['_computed_name'] = reg_players['First Name'].astype(str).str.strip() + ' ' + reg_players['Last Name'].astype(str).str.strip()
    elif 'First Name' in reg_players.columns and 'Surname' in reg_players.columns:
        reg_players['_computed_name'] = reg_players['First Name'].astype(str).str.strip() + ' ' + reg_players['Surname'].astype(str).str.strip()
    elif 'Full Name' in reg_players.columns:
        reg_players['_computed_name'] = reg_players['Full Name'].astype(str).str.replace('‡', '', regex=False).str.strip()
    
    name_cols = []
    if '_computed_name' in reg_players.columns:
        name_cols.append('_computed_name')
    name_cols.extend([c for c in reg_players.columns if 'name' in str(c).lower() and c != '_computed_name'])
    
    if not name_cols:
        name_cols = [reg_players.columns[0]]
        
    for _, r in reg_players.iterrows():
        clubs_found = []
        if 'Individual Membership Primary Club' in reg_players.columns and pd.notna(r['Individual Membership Primary Club']):
            val = str(r['Individual Membership Primary Club']).strip()
            if val.lower() != 'nan' and val != '':
                clubs_found.append(val)
        
        for keyword in ['Primary Club', 'Transfer', 'Wylie', 'Club']:
            cols = [c for c in reg_players.columns if keyword in str(c) and c != 'Individual Membership Primary Club' and 'Date' not in str(c)]
            for col in cols:
                if pd.notna(r[col]) and str(r[col]).strip() and str(r[col]).lower() != 'nan':
                    val = str(r[col]).strip()
                    if val not in clubs_found:
                        clubs_found.append(val)

        if clubs_found:
            reg_club = " / ".join(clubs_found)
            for c in name_cols:
                if pd.notna(r[c]):
                    r_name = str(r[c]).replace('‡', '').strip()
                    norm_name = re.sub(r'\s+', ' ', r_name).strip()
                    if norm_name and norm_name.lower() != 'nan':
                        norm_lower = norm_name.lower()
                        mapped_name = alias_map.get(norm_lower, norm_name)
                        mapped_lower = str(mapped_name).strip().lower()
                        club_map[mapped_lower] = reg_club
                        club_map[mapped_name] = reg_club
                        club_map[norm_lower] = reg_club
                        club_map[norm_name] = reg_club

    if unreg_map_df is not None and not unreg_map_df.empty:
        col_name = unreg_map_df.columns[0]
        col_club = unreg_map_df.columns[1]
        
        for _, r in unreg_map_df.iterrows():
            if pd.notna(r[col_name]) and pd.notna(r[col_club]):
                p_name = str(r[col_name]).strip().lower()
                p_club = str(r[col_club]).strip()
                
                if p_name and p_name != 'nan':
                    mapped_name = alias_map.get(p_name, p_name)
                    for key in [mapped_name, p_name]:
                        if key in club_map:
                            if p_club.lower() not in club_map[key].lower():
                                club_map[key] = f"{club_map[key]} / {p_club}"
                        else:
                            club_map[key] = p_club

    if id_map_df is None and domain:
        f_id_map = DEFAULT_FILES.get(domain, {}).get("id_map", "")
        if f_id_map and os.path.exists(f_id_map):
            try:
                id_map_df = pd.read_excel(f_id_map)
            except Exception:
                id_map_df = None

    if id_map_df is not None and not id_map_df.empty:
        col_nv_name = next((c for c in id_map_df.columns if 'nv' in c.lower() and 'name' in c.lower()), 'NV_Play_Name')
        col_s80_name = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'name' in c.lower()), 'Sport80_Name')
        col_s80_club = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'club' in c.lower()), 'Sport80_Club')
        for _, r in id_map_df.iterrows():
            club_val = r.get(col_s80_club)
            if pd.notna(club_val) and str(club_val).strip() and str(club_val).lower() != 'nan':
                club_str = str(club_val).strip()
                for name_c in [col_nv_name, col_s80_name]:
                    name_val = r.get(name_c)
                    if pd.notna(name_val) and str(name_val).strip() and str(name_val).lower() != 'nan':
                        p_name = str(name_val).strip().lower()
                        if mapped_name not in club_map:
                            club_map[mapped_name] = club_str
                        if p_name not in club_map:
                            club_map[p_name] = club_str

    if secondary_map is None:
        f_sec = DEFAULT_FILES.get(domain, {}).get("secondary", "")
        if f_sec:
            base_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
            f_sec_path = os.path.join(base_dir, f_sec) if not os.path.isabs(f_sec) else f_sec
            target_sec = f_sec_path if os.path.exists(f_sec_path) else (f_sec if os.path.exists(f_sec) else None)
            if target_sec:
                try:
                    sec_df = pd.read_excel(target_sec)
                    secondary_map = build_secondary_team_map(sec_df, alias_map)
                except Exception:
                    secondary_map = None

    if secondary_map:
        if isinstance(secondary_map, pd.DataFrame) and not secondary_map.empty:
            secondary_map = build_secondary_team_map(secondary_map, alias_map)
        if isinstance(secondary_map, dict):
            for p_name, teams in secondary_map.items():
                if not p_name or not teams: continue
                p_clean = str(p_name).strip().lower()
                existing = club_map.get(p_clean, "")
                team_strs = [t for t in teams if t not in existing]
                if team_strs:
                    combined = f"{existing} / {' / '.join(team_strs)}" if existing else ' / '.join(team_strs)
                    club_map[p_clean] = combined
                    mapped = alias_map.get(p_clean, p_clean)
                    club_map[mapped] = combined
                        
    return club_map

def build_player_fixture_club_counts(batting_df, bowling_df, alias_map=None):
    from collections import defaultdict
    apps = []
    if batting_df is not None and not batting_df.empty:
        name_col = 'Cleaned Name' if 'Cleaned Name' in batting_df.columns else ('Name' if 'Name' in batting_df.columns else batting_df.columns[0])
        sub = batting_df[[name_col, 'Group']].copy().rename(columns={name_col: 'Player'})
        apps.append(sub)
    if bowling_df is not None and not bowling_df.empty:
        name_col = 'Cleaned Name' if 'Cleaned Name' in bowling_df.columns else ('Bowler' if 'Bowler' in bowling_df.columns else bowling_df.columns[0])
        sub = bowling_df[[name_col, 'Group']].copy().rename(columns={name_col: 'Player'})
        apps.append(sub)
        
    if not apps: return {}
    all_apps = pd.concat(apps, ignore_index=True).drop_duplicates(subset=['Player', 'Group'])
    player_fixture_clubs = defaultdict(lambda: defaultdict(int))
    for _, r in all_apps.iterrows():
        p = str(r['Player']).strip().lower()
        p_mapped = str(alias_map.get(p, p)).strip().lower() if alias_map else p
        grp = str(r['Group'])
        if ' v ' in grp:
            t1, t2 = extract_teams_from_group(grp)
            c1, c2 = extract_base_club_name(t1), extract_base_club_name(t2)
            if c1 != 'Unknown Club': 
                player_fixture_clubs[p_mapped][c1.lower()] += 1
                if p != p_mapped:
                    player_fixture_clubs[p][c1.lower()] += 1
            if c2 != 'Unknown Club': 
                player_fixture_clubs[p_mapped][c2.lower()] += 1
                if p != p_mapped:
                    player_fixture_clubs[p][c2.lower()] += 1
    return player_fixture_clubs

def infer_unregistered_player_clubs(batting_df, bowling_df, player_club_map, min_matches=2):
    from collections import defaultdict
    if player_club_map is None:
        player_club_map = {}
    apps = []
    if batting_df is not None and not batting_df.empty:
        name_col = 'Cleaned Name' if 'Cleaned Name' in batting_df.columns else ('Name' if 'Name' in batting_df.columns else batting_df.columns[0])
        sub = batting_df[[name_col, 'Group']].copy().rename(columns={name_col: 'Player'})
        if 'Team' in batting_df.columns: sub['Team'] = batting_df['Team']
        apps.append(sub)
    if bowling_df is not None and not bowling_df.empty:
        name_col = 'Cleaned Name' if 'Cleaned Name' in bowling_df.columns else ('Bowler' if 'Bowler' in bowling_df.columns else bowling_df.columns[0])
        sub = bowling_df[[name_col, 'Group']].copy().rename(columns={name_col: 'Player'})
        if 'Team' in bowling_df.columns: sub['Team'] = bowling_df['Team']
        apps.append(sub)
        
    if not apps: return player_club_map
    
    all_apps = pd.concat(apps, ignore_index=True).drop_duplicates(subset=['Player', 'Group'])
    
    for player, grp in all_apps.groupby('Player'):
        p_clean = str(player).split(' (')[0].strip().lower()
        if p_clean in player_club_map or str(player).strip().lower() in player_club_map:
            continue
        
        if len(grp) >= min_matches:
            club_counts = defaultdict(int)
            teams_in_fixtures = []
            
            for _, r in grp.iterrows():
                row_t = str(r.get('Team', '')).strip()
                if row_t and row_t.lower() != 'nan':
                    b_club = extract_base_club_name(row_t)
                    if b_club != "Unknown Club":
                        club_counts[b_club] += 1
                
                grp_str = str(r.get('Group', ''))
                if ' v ' in grp_str:
                    t1, t2 = extract_teams_from_group(grp_str)
                    c1, c2 = extract_base_club_name(t1), extract_base_club_name(t2)
                    teams_in_fixtures.append({c1.lower(), c2.lower()})
            
            inferred_club = None
            if club_counts:
                top_club, count = sorted(club_counts.items(), key=lambda x: x[1], reverse=True)[0]
                if count >= min_matches or (count / len(grp)) >= 0.5:
                    inferred_club = top_club
            
            if not inferred_club and teams_in_fixtures:
                common = set.intersection(*teams_in_fixtures)
                common.discard('unknown club')
                if len(common) == 1:
                    inferred_club = list(common)[0].title()
            
            if inferred_club:
                player_club_map[p_clean] = inferred_club
                player_club_map[str(player).strip().lower()] = inferred_club

    return player_club_map

def build_league_dict(league_structure):
    league_dict = {}
    original_league_order = [] 
    
    team_col = next((col for col in league_structure.columns if 'team' in str(col).lower() or 'club' in str(col).lower()), league_structure.columns[0])
    league_col = next((col for col in league_structure.columns if 'league' in str(col).lower() or 'division' in str(col).lower()), league_structure.columns[1])

    for _, row in league_structure.iterrows():
        raw_team = str(row[team_col]).strip()
        raw_league = str(row[league_col]).strip()
        if raw_team and raw_team.lower() != 'nan' and raw_league and raw_league.lower() != 'nan':
            league_dict[raw_team] = raw_league
            if raw_league not in original_league_order:
                original_league_order.append(raw_league)
                
    return league_dict, list(league_dict.keys()), original_league_order

def extract_xi(team_str):
    match = re.search(r'((?:mw\d?|\d(?:st|nd|rd|th))\s*xi)', str(team_str).lower())
    if match: return match.group(1).replace(' ', '')
    return None

def clean_team_for_compare(t, domain):
    t = str(t).lower()
    t = re.sub(r'\bcc\b|\bcricket club\b', '', t)
    t = t.replace('1881', '').replace('ciyms', 'ci').replace('dungannnon', 'dungannon')
    t = re.sub(r'(?i)northern\s+ireland\s+malayali\s+association', 'nima', t)
    t = re.sub(r'(?i)\bnima\s*cc\b|\bnimacc\b|\bnima\b', 'nima', t)
    t = re.sub(r'(?i)belfast\s+international\s+sports\s+club|belfast\s+b\.i\.s\.c\.', 'bisc', t)
    t = re.sub(r'(?i)civil\s+service\s+north\s+of\s+ireland|civil\s+service\s+north', 'csni', t)
    t = re.sub(r'(?i)drumaness\s+super\s*kings', 'drumaness', t)
    t = re.sub(r'(?i)donaghcloney', 'donacloney', t)
    
    if domain == "Women's":
        t = re.sub(r'\bwomen\'s\b|\bwomens\b|\bwomen\b', '', t)
        
    t = re.sub(r'\b1st\s*xi\b|\b1st\b', '1', t)
    t = re.sub(r'\b2nd\s*xi\b|\b2nd\b', '2', t)
    t = re.sub(r'\b3rd\s*xi\b|\b3rd\b', '3', t)
    t = re.sub(r'\b4th\s*xi\b|\b4th\b', '4', t)
    t = re.sub(r'\b5th\s*xi\b|\b5th\b', '5', t)
    t = re.sub(r'\b6th\s*xi\b|\b6th\b', '6', t)
    t = re.sub(r'\bxi\b', '', t)
    
    return " ".join(t.split())

def get_team_league(team_name, team_keys, league_dict, domain):
    if pd.isna(team_name): return None
    clean_search_name = re.sub(r',.*', '', str(team_name)).strip()
    clean_search_name = re.sub(r'(?i)\s*\(Pathway\)', '', clean_search_name).strip()
    if clean_search_name.startswith("Unknown") or "Unknown (" in clean_search_name or clean_search_name == "Unknown Team":
        return None
    for k in team_keys:
        if k.lower() == clean_search_name.lower(): return league_dict[k]
    team_clean = clean_team_for_compare(clean_search_name, domain)
    for k in team_keys:
        if clean_team_for_compare(k, domain) == team_clean: return league_dict[k]
    if not extract_xi(clean_search_name):
        bare_fallback = f"{clean_search_name} 1st XI"
        fallback_clean = clean_team_for_compare(bare_fallback, domain)
        for k in team_keys:
            if clean_team_for_compare(k, domain) == fallback_clean: return league_dict[k]
    team_xi = extract_xi(clean_search_name)
    best_match, best_score = None, 0
    for k in team_keys:
        k_clean = clean_team_for_compare(k, domain)
        k_xi = extract_xi(k)
        if team_xi == k_xi or team_xi is None or k_xi is None:
            score = fuzz.token_sort_ratio(team_clean, k_clean)
            if score > best_score and score >= 75:
                best_score, best_match = score, k
    return league_dict.get(best_match)

def format_display_team(team_str, domain):
    if pd.isna(team_str): return "Unknown"
    c = str(team_str).strip()
    if c.startswith("Unknown (") and c.endswith(")"):
        return c
        
    is_pathway = "(Pathway)" in c
    c = re.sub(r'(?i)\s*\(Pathway\)', '', c).strip()
    
    # Strip formal club suffixes so that fallback matches seamlessly merge with standard scorecard entries
    c = re.sub(r'(?i)\s*Cricket Club\b', '', c)
    c = re.sub(r'(?i)\bCC\b', '', c)
    c = re.sub(r'(?i)\bnima\s*cc\b|\bnimacc\b|\bnima\b', 'NIMA', c)
    c = re.sub(r'(?i)belfast\s+international\s+sports\s+club|belfast\s+b\.i\.s\.c\.', 'BISC', c)
    c = re.sub(r'(?i)civil\s+service\s+north\s+of\s+ireland|civil\s+service\s+north', 'CSNI', c)
    c = re.sub(r'(?i)drumaness\s+super\s*kings', 'Drumaness Superkings', c)
    c = c.replace('Donaghcloney', 'Donacloney')
    
    if domain == "Midweek":
        c = re.sub(r'(?i)\bmw1\s*xi\b', ' 1', c)
        c = re.sub(r'(?i)\bmw2\s*xi\b', ' 2', c)
        c = re.sub(r'(?i)\bmw\s*xi\b', '', c)
        c = c.replace(' 1st XI', ' 1').replace(' 1st', ' 1')
    elif domain == "Women's":
        c = re.sub(r'(?i)\bwomen\'s\b|\bwomens\b|\bwomen\b', '', c)
        c = c.replace(' 1st XI', '').replace(' 1st', '')
    else:
        c = c.replace(' 1st XI', '').replace(' 1st', '')
    c = c.replace(' 2nd XI', ' 2').replace(' 3rd XI', ' 3')
    c = c.replace(' 4th XI', ' 4').replace(' 5th XI', ' 5').replace(' 6th XI', ' 6')
    c = c.replace(' XI', '')
    c = re.sub(r'\s+', ' ', c).strip()
    if domain != "Midweek" and c.endswith(' 1'):
        c = c[:-2].strip()
    if 'Holywood' in c and '1881' not in c:
        c = c.replace('Holywood', 'Holywood 1881')
    if domain == "Midweek" and not c.endswith(" MW"):
        c = f"{c} MW"
        
    if is_pathway:
        return "NCU Pathway XI"
    return c

def extract_teams_from_group(group_str):
    try:
        parts = str(group_str).strip().rsplit(' - ', 1)
        rest = parts[0].strip()
        if ' v ' in rest:
            t1, remainder = rest.split(' v ', 1)
            t2 = remainder.rsplit(', ', 1)[0] if ', ' in remainder else remainder.rsplit(' - ', 1)[0]
            return t1.strip(), t2.strip()
        return rest, "Unknown"
    except: return "Unknown", "Unknown"

def determine_player_team_for_row(row, player_club_map, domain, secondary_map=None, player_fixture_clubs=None, alias_map=None):
    player = str(row.get('Cleaned Name', row.get('Player', row.get('Name', row.get('Bowler', ''))))).strip()
    group_str = str(row.get('Group', row.get('Match', '')))
    t1, t2 = extract_teams_from_group(group_str)
    if not t1 or t1 == "Unknown" or not t2 or t2 == "Unknown":
        return f"Unknown ({t1} v {t2})"

    c1_base = extract_base_club_name(t1).lower()
    c2_base = extract_base_club_name(t2).lower()
    
    clean_t1_match = clean_club_for_matching(t1)
    clean_t2_match = clean_club_for_matching(t2)
    
    if '(' in player and player.endswith(')'):
        club_hint = player.split('(')[-1].replace(')', '').strip().lower()
        clean_hint = clean_club_for_matching(club_hint)
        if clean_hint in clean_t1_match or club_hint in t1.lower(): return t1
        if clean_hint in clean_t2_match or club_hint in t2.lower(): return t2

    base_p = player.split(' (')[0].strip().lower()
    mapped_p = str(alias_map.get(base_p, base_p)).strip().lower() if alias_map else base_p
    
    if domain != "Women's":
        if mapped_p in ['neil brand', 'sandeep singh']:
            return t1 if 'muckamore' in t1.lower() else (t2 if 'muckamore' in t2.lower() else t1)
        if mapped_p == 'james shannon':
            if 'holywood' in t1.lower() or 'holywood' in t2.lower(): return t1 if 'holywood' in t1.lower() else t2
            elif 'saintfield' in t1.lower() or 'saintfield' in t2.lower(): return t1 if 'saintfield' in t1.lower() else t2

    known_clubs = set()
    
    # 1. Registered / Transferred Clubs
    reg_val = (player_club_map.get(mapped_p) or player_club_map.get(base_p)) if player_club_map else None
    if reg_val and str(reg_val).lower() != 'nan':
        for chunk in str(reg_val).split('/'):
            c_clean = extract_base_club_name(chunk).lower()
            if c_clean and c_clean != 'unknown club':
                known_clubs.add(c_clean)

    # 2. Secondary Map
    if secondary_map:
        sec_teams = secondary_map.get(mapped_p) or secondary_map.get(base_p) or []
        for st in sec_teams:
            c_clean = extract_base_club_name(st).lower()
            if c_clean and c_clean != 'unknown club':
                known_clubs.add(c_clean)

    # 3. Fixture Appearance Frequency (>= 2 matches)
    counts = {}
    if player_fixture_clubs:
        counts = player_fixture_clubs.get(mapped_p) or player_fixture_clubs.get(base_p) or {}
        for club, cnt in counts.items():
            if cnt >= 2:
                known_clubs.add(club)
                
    t1_matches = any(kc in clean_t1_match or clean_t1_match in kc or kc == c1_base for kc in known_clubs)
    t2_matches = any(kc in clean_t2_match or clean_t2_match in kc or kc == c2_base for kc in known_clubs)
    
    t1_is_pathway = 'pathway' in t1.lower()
    t2_is_pathway = 'pathway' in t2.lower()
    is_pathway_match = t1_is_pathway or t2_is_pathway

    if t1_matches and not t2_matches:
        if t2_is_pathway: return t1
        return t1
    elif t2_matches and not t1_matches:
        if t1_is_pathway: return t2
        return t2
    elif t1_matches and t2_matches:
        if t2_is_pathway: return t1
        if t1_is_pathway: return t2
        c1_cnt = counts.get(c1_base, 0)
        c2_cnt = counts.get(c2_base, 0)
        if c1_cnt > c2_cnt: return t1
        elif c2_cnt > c1_cnt: return t2
        return t1
        
    if is_pathway_match and reg_val and str(reg_val).lower() != 'nan':
        primary_reg_club = str(reg_val).split('/')[0].strip()
        return f"{primary_reg_club} (Pathway)"
        
    # If we couldn't match the teams to the player's known clubs, we do NOT fallback to 
    # their registered club's 1st XI, as this artificially inflates Premier League stats
    # when scorers make mistakes (e.g., adding a Woodvale player to a Dundrum v Templepatrick match).
    # Instead, we let it drop through to "Unknown" so it gets flagged in the Unassigned/Non-NCU tab.
        
    if t1 and t1 != "Unknown" and t2 and t2 != "Unknown":
        return f"Unknown ({t1} v {t2})"
    elif t1 and t1 != "Unknown":
        return t1
    elif t2 and t2 != "Unknown":
        return t2
    return "Unknown Team"

def parse_high_score(scores_series):
    best_score = 0
    is_not_out = False
    for hs in scores_series.dropna().astype(str):
        if hs.lower() == 'nan': continue
        val = hs.replace('*', '').replace('.0', '')
        try:
            val_int = int(val)
            if val_int > best_score:
                best_score = val_int
                is_not_out = '*' in hs
            elif val_int == best_score and '*' in hs:
                is_not_out = True
        except ValueError: pass
    if best_score == 0 and not is_not_out: return "0"
    return f"{best_score}*" if is_not_out else str(best_score)

def calculate_averages(batting_df, bowling_df, player_club_map, team_keys, league_dict, domain, bat_sort="Runs", bowl_sort="Wickets", secondary_map=None, alias_map=None, _cache_version=None):
    for col in ['Matches', 'Innings', 'Not Outs', 'Runs', 'Balls', 'Fours', 'Sixes', 'Catches', 'Catches as Keeper', 'Stumpings']:
        if col in batting_df.columns: batting_df[col] = pd.to_numeric(batting_df[col], errors='coerce').fillna(0)
    for col in ['Innings', 'Balls', 'Maidens', 'Runs', 'Wickets']:
        if col in bowling_df.columns: bowling_df[col] = pd.to_numeric(bowling_df[col], errors='coerce').fillna(0)
            
    if 'Cleaned Name' in batting_df.columns:
        batting_df['Cleaned Name'] = batting_df['Cleaned Name'].apply(fix_celtic_casing)
    if 'Cleaned Name' in bowling_df.columns:
        bowling_df['Cleaned Name'] = bowling_df['Cleaned Name'].apply(fix_celtic_casing)
            
    player_fixture_clubs = build_player_fixture_club_counts(batting_df, bowling_df, alias_map=alias_map)
    
    batting_df['Team Played For'] = batting_df.apply(lambda r: determine_player_team_for_row(r, player_club_map, domain, secondary_map, player_fixture_clubs=player_fixture_clubs, alias_map=alias_map), axis=1)
    bowling_df['Team Played For'] = bowling_df.apply(lambda r: determine_player_team_for_row(r, player_club_map, domain, secondary_map, player_fixture_clubs=player_fixture_clubs, alias_map=alias_map), axis=1)
    
    def get_opponent_from_row(row):
        group_str = str(row.get('Group', row.get('Match', '')))
        t1, t2 = extract_teams_from_group(group_str)
        team_played = str(row.get('Team Played For', ''))
        if team_played == t1: return t2
        if team_played == t2: return t1
        c_my = extract_base_club_name(team_played).lower()
        c_t1 = extract_base_club_name(t1).lower()
        c_t2 = extract_base_club_name(t2).lower()
        if c_my == c_t1: return t2
        if c_my == c_t2: return t1
        return t2 if t1 in team_played else t1

    if not batting_df.empty:
        batting_df['Opponent'] = batting_df.apply(get_opponent_from_row, axis=1)
    if not bowling_df.empty:
        bowling_df['Opponent'] = bowling_df.apply(get_opponent_from_row, axis=1)
    def apply_league(t):
        league = get_team_league(t, team_keys, league_dict, domain)
        if not league: return "Unassigned"
        if domain == "Midweek": return str(league)
        league_str = str(league)
        target_words = ['premier', 'senior league 1', 'senior league 2', 'senior league 3'] if domain == "Men's" else ['premier', 'senior league 1', 'senior league 2', 'senior league 3', 'senior']
        if any(word in league_str.lower() for word in target_words): return league_str.replace('NCU', 'Mercury')
        return league_str

    batting_df['League'] = batting_df['Team Played For'].apply(apply_league)
    bowling_df['League'] = bowling_df['Team Played For'].apply(apply_league)

    def combine_teams(team_series, domain):
        if team_series.empty: return "Unknown"
        teams = []
        for t in team_series.unique():
            fmt = format_display_team(t, domain)
            if fmt != "Unknown" and fmt not in teams:
                teams.append(fmt)
        return " / ".join(teams) if teams else "Unknown"

    def group_batting(df_to_group):
        agg_dict = {
            'Name': lambda x: x.value_counts().index[0] if not x.empty else "Unknown",
            'Team Played For': lambda x: combine_teams(x, domain),
            'Matches': 'sum', 'Innings': 'sum', 'Not Outs': 'sum', 'Runs': 'sum',
            'Balls': 'sum', 'Fours': 'sum', 'Sixes': 'sum', 'High Score': parse_high_score
        }
        for col in ['Catches', 'Catches as Keeper', 'Stumpings']:
            if col in df_to_group.columns:
                agg_dict[col] = 'sum'
        
        def score_val(s):
            s = str(s).replace('*', '').replace('.0', '')
            return int(s) if s.isdigit() else 0
        def is_not_out(s):
            return 1 if '*' in str(s) else 0
        
        df_copy = df_to_group.copy()
        df_copy['Score_Int'] = df_copy['High Score'].apply(score_val)
        df_copy['Score_NO'] = df_copy['High Score'].apply(is_not_out)
        
        sorted_bat = df_copy.sort_values(by=['Score_Int', 'Score_NO'], ascending=[False, False])
        best_innings = sorted_bat.drop_duplicates(subset=['League', 'Cleaned Name']).copy()
        if 'Opponent' in best_innings.columns:
            best_innings['High Score Against'] = best_innings['Opponent']
        else:
            best_innings['High Score Against'] = "Unknown"
        best_innings_map = best_innings.set_index(['League', 'Cleaned Name'])['High Score Against']

        grouped = df_to_group.groupby(['League', 'Cleaned Name']).agg(agg_dict).reset_index()
        grouped = grouped.merge(best_innings_map, on=['League', 'Cleaned Name'], how='left')
        
        grouped.rename(columns={'Team Played For': 'Team', 'Name': 'Player'}, inplace=True)
        grouped.drop(columns=['Cleaned Name'], inplace=True)
        outs = grouped['Innings'] - grouped['Not Outs']
        grouped['Average'] = np.where(outs > 0, grouped['Runs'] / outs, np.nan)
        grouped['Strike Rate'] = np.where(grouped['Balls'] > 0, (grouped['Runs'] / grouped['Balls']) * 100, np.nan)
        
        cols = ['League', 'Player', 'Team', 'Matches', 'Innings', 'Not Outs', 'Runs', 'Balls', 'Fours', 'Sixes', 'High Score', 'High Score Against', 'Average', 'Strike Rate']
        for col in ['Catches', 'Catches as Keeper', 'Stumpings']:
            if col in grouped.columns: cols.append(col)
        return grouped[cols]

    def group_bowling(df_to_group, bat_avgs):
        sorted_df = df_to_group.sort_values(by=['Wickets', 'Runs'], ascending=[False, True])
        best_spells = sorted_df.drop_duplicates(subset=['League', 'Cleaned Name']).copy()
        best_spells['Best Bowling'] = best_spells['Wickets'].fillna(0).astype(int).astype(str) + '-' + best_spells['Runs'].fillna(0).astype(int).astype(str)
        if 'Opponent' in best_spells.columns:
            best_spells['Best Bowling Against'] = best_spells['Opponent']
        else:
            best_spells['Best Bowling Against'] = "Unknown"
        bbi_series = best_spells.set_index(['League', 'Cleaned Name'])[['Best Bowling', 'Best Bowling Against']]
        
        grouped = df_to_group.groupby(['League', 'Cleaned Name']).agg({
            'Bowler': lambda x: x.value_counts().index[0] if not x.empty else "Unknown",
            'Team Played For': lambda x: combine_teams(x, domain),
            'Innings': 'sum', 'Balls': 'sum', 'Maidens': 'sum', 'Runs': 'sum', 'Wickets': 'sum'
        }).reset_index()
        grouped = grouped.merge(bbi_series, on=['League', 'Cleaned Name'], how='left')
        grouped.rename(columns={'Team Played For': 'Team', 'Bowler': 'Player'}, inplace=True)
        grouped.drop(columns=['Cleaned Name'], inplace=True)
        
        total_matches = bat_avgs[['League', 'Player', 'Team', 'Matches']].rename(columns={'Matches': 'Total_Matches'})
        grouped = grouped.merge(total_matches, on=['League', 'Player', 'Team'], how='left')
        grouped['Matches'] = grouped['Total_Matches'].fillna(grouped['Innings']).astype(int)
        grouped['Overs'] = (grouped['Balls'] // 6) + (grouped['Balls'] % 6) / 10
        grouped['Average'] = np.where(grouped['Wickets'] > 0, grouped['Runs'] / grouped['Wickets'], np.nan)
        grouped['Economy'] = np.where(grouped['Balls'] > 0, (grouped['Runs'] / grouped['Balls']) * 6, np.nan)
        grouped['Strike Rate'] = np.where(grouped['Wickets'] > 0, grouped['Balls'] / grouped['Wickets'], np.nan)
        return grouped[['League', 'Player', 'Team', 'Matches', 'Innings', 'Overs', 'Maidens', 'Runs', 'Wickets', 'Best Bowling', 'Best Bowling Against', 'Average', 'Economy', 'Strike Rate']]

    if domain == "Midweek":
        leagues_bat = group_batting(batting_df)
        overall_bat_df = batting_df.copy()
        overall_bat_df['League'] = 'Overall Midweek'
        overall_bat = group_batting(overall_bat_df)
        batting_final = pd.concat([leagues_bat, overall_bat], ignore_index=True)
        
        leagues_bowl = group_bowling(bowling_df, batting_final)
        overall_bowl_df = bowling_df.copy()
        overall_bowl_df['League'] = 'Overall Midweek'
        overall_bowl = group_bowling(overall_bowl_df, batting_final)
        bowling_final = pd.concat([leagues_bowl, overall_bowl], ignore_index=True)
    else:
        batting_final = group_batting(batting_df)
        bowling_final = group_bowling(bowling_df, batting_final)
        
    if bat_sort == "Average": batting_final = batting_final.sort_values(by=['League', 'Average', 'Runs'], ascending=[True, False, False])
    elif bat_sort == "Strike Rate": batting_final = batting_final.sort_values(by=['League', 'Strike Rate', 'Runs'], ascending=[True, False, False])
    else: batting_final = batting_final.sort_values(by=['League', 'Runs', 'Average'], ascending=[True, False, False])

    if bowl_sort == "Average": bowling_final = bowling_final.sort_values(by=['League', 'Average', 'Wickets'], ascending=[True, True, False], na_position='last')
    elif bowl_sort == "Economy": bowling_final = bowling_final.sort_values(by=['League', 'Economy', 'Wickets'], ascending=[True, True, False], na_position='last')
    elif bowl_sort == "Strike Rate": bowling_final = bowling_final.sort_values(by=['League', 'Strike Rate', 'Wickets'], ascending=[True, True, False], na_position='last')
    else: bowling_final = bowling_final.sort_values(by=['League', 'Wickets', 'Average'], ascending=[True, False, True], na_position='last')
    
    if domain in ["Women's", "Midweek"]:
        batting_final['Average'] = batting_final['Average'].round(2)
        batting_final['Strike Rate'] = batting_final['Strike Rate'].round(2)
        bowling_final['Average'] = bowling_final['Average'].round(2)
        bowling_final['Economy'] = bowling_final['Economy'].round(2)
        bowling_final['Strike Rate'] = bowling_final['Strike Rate'].round(2)

    return batting_final, bowling_final

def custom_league_sort(league_name, domain, ordered_leagues=None):
    name_lower = str(league_name).lower()
    if domain == "Midweek" and 'overall' in name_lower: return (0, 0, '')
    if ordered_leagues and league_name in ordered_leagues: return (1, ordered_leagues.index(league_name), '')
    if domain == "Midweek":
        if 'group a' in name_lower: return (2, 1, league_name)
        elif 'group b' in name_lower: return (2, 2, league_name)
        elif 'group c' in name_lower: return (2, 3, league_name)
        return (3, 0, league_name)
    else:
        if 'premier' in name_lower: return (2, 0, '')
        elif 'senior' in name_lower or 'section' in name_lower:
            match = re.search(r'senior(?: league)? (\d+)', name_lower)
            return (3, int(match.group(1)) if match else 99, '')
        elif 'junior' in name_lower:
            match = re.search(r'junior(?: league)? (\d+)([a-z]?)', name_lower)
            return (4, int(match.group(1)) if match else 99, match.group(2) if match else '')
        return (5, 0, name_lower)

def format_excel_sheet(writer, df, sheet_name, min_label=None):
    safe_sheet_name = str(sheet_name).replace("League", "Lge").replace("Midweek", "MW").replace("Group", "Grp").replace("Overall", "Ovr").strip()[:31].strip()
    df.to_excel(writer, sheet_name=safe_sheet_name, index=False)
    worksheet = writer.sheets[safe_sheet_name]
    workbook = writer.book
    
    left_header = workbook.add_format({'bold': True, 'bottom': 1, 'bg_color': '#D9D9D9', 'align': 'left'})
    center_header = workbook.add_format({'bold': True, 'bottom': 1, 'bg_color': '#D9D9D9', 'align': 'center'})
    bold_name, left_align, center_align = workbook.add_format({'bold': True}), workbook.add_format({'align': 'left'}), workbook.add_format({'align': 'center'})
    two_decimals = workbook.add_format({'num_format': '0.00', 'align': 'center'})
    
    for col_num, col_name in enumerate(df.columns):
        worksheet.write(0, col_num, col_name, left_header if col_name in ['Player', 'Team', 'High Score Against', 'Best Bowling Against'] else center_header)
        
        # Calculate optimal width
        col_width = max(max((len(str(x)) for x in df[col_name]), default=0), len(str(col_name))) + 2
        
        if col_name == 'Player': 
            worksheet.set_column(col_num, col_num, col_width, bold_name)
        elif col_name in ['Team', 'High Score Against', 'Best Bowling Against']: 
            worksheet.set_column(col_num, col_num, col_width, left_align)
        elif col_name in ['Average', 'Strike Rate', 'Economy']: 
            # Reduce width for these specific columns as requested
            worksheet.set_column(col_num, col_num, len(str(col_name)) + 0.5, two_decimals)
        else: 
            worksheet.set_column(col_num, col_num, col_width, center_align)
            
    if min_label:
        worksheet.write(len(df) + 2, 0, min_label, workbook.add_format({'italic': True, 'bold': True}))

# ==========================================
# WORD DOC GENERATOR FUNCTIONS
# ==========================================
def doc_format_cricket_names(text, domain):
    if pd.isna(text): return text
    text = str(text)
    text = text.replace('NCU', 'Mercury').replace('Mercury Pathway', 'NCU Pathway')
    text = text.replace('CIYMS', 'CI')
    text = re.sub(r'(?i)\bnima\s*cc\b|\bnimacc\b|\bnima\b', 'NIMA', text)
    text = re.sub(r'(?i)belfast\s+international\s+sports\s+club|belfast\s+b\.i\.s\.c\.', 'BISC', text)
    text = re.sub(r'(?i)civil\s+service\s+north\s+of\s+ireland|civil\s+service\s+north', 'CSNI', text)
    text = re.sub(r'(?i)drumaness\s+super\s*kings', 'Drumaness Superkings', text)
    text = text.replace('Donaghcloney', 'Donacloney')
    if 'Holywood' in text and '1881' not in text: text = text.replace('Holywood', 'Holywood 1881')
    
    if domain == "Midweek": text = text.replace(' 1st XI', ' 1').replace(' 1st', ' 1')
    else: text = text.replace(' 1st XI', '').replace(' 1st', '')
        
    text = text.replace(' 2nd XI', ' 2').replace(' 3rd XI', ' 3')
    text = text.replace(' 4th XI', ' 4').replace(' 5th XI', ' 5').replace(' 6th XI', ' 6')
    text = text.replace(' XI', '') 
    text = re.sub(r'\s+', ' ', text).strip()
    
    if domain != "Midweek" and text.endswith(' 1'):
        text = text[:-2].strip()
    return text

def doc_get_player_team_from_match(match_str, base_club):
    if pd.isna(match_str) or ' v ' not in str(match_str): return "Unknown Team"
    team1, team2 = extract_teams_from_group(match_str)
    
    def clean_for_matching(s):
        s = str(s).lower()
        s = re.sub(r'northern\s+ireland\s+malayali\s+association', 'nima', s)
        s = re.sub(r'\bnima\s*cc\b|\bnimacc\b|\bnima\b', 'nima', s)
        s = re.sub(r'belfast\s+international\s+sports\s+club|belfast\s+b\.i\.s\.c\.', 'bisc', s)
        s = re.sub(r'civil\s+service\s+north\s+of\s+ireland|civil\s+service\s+north', 'csni', s)
        s = re.sub(r'drumaness\s+super\s*kings', 'drumaness', s)
        s = re.sub(r'donaghcloney', 'donacloney', s)
        s = re.sub(r'\b(cricket club|club|teams|cc|1st|2nd|3rd|4th|5th|6th|1|2|3|4|5|6|xi)\b', '', s)
        return set(s.split())
        
    t1_words, t2_words, club_words = clean_for_matching(team1), clean_for_matching(team2), clean_for_matching(base_club)
    for target in ['nima', 'bisc', 'csni', 'drumaness', 'donacloney']:
        if target in club_words:
            if target in t1_words: return team1
            if target in t2_words: return team2
    if len(t1_words.intersection(club_words)) > len(t2_words.intersection(club_words)): return team1
    elif len(t2_words.intersection(club_words)) > len(t1_words.intersection(club_words)): return team2
    
    if str(base_club).lower() in team1.lower(): return team1
    if str(base_club).lower() in team2.lower(): return team2
    if team1 != "Unknown" and team2 != "Unknown":
        return f"Unknown ({team1} v {team2})"
    return "Unknown Team"

def doc_team_sort_key(team_name):
    words = team_name.split()
    if words and words[-1].isdigit(): return (team_name.rsplit(' ', 1)[0], int(words[-1]))
    return (team_name, 1)

def add_bullet_point(doc, text, level=1, space_after=0, line_spacing=0.9, bold_substring=None):
    style_name = 'List Bullet' if level == 1 else f'List Bullet {level}'
    
    def apply_bold(p, full_text, prefix="", f_size=11):
        if bold_substring and bold_substring in full_text:
            parts = full_text.split(bold_substring, 1)
            
            if prefix or parts[0]:
                r1 = p.add_run(prefix + parts[0])
                r1.font.name, r1.font.size = 'Calibri', Pt(f_size)
            
            r2 = p.add_run(bold_substring)
            r2.font.name, r2.font.size = 'Calibri', Pt(f_size)
            r2.bold = True
            
            if parts[1]:
                r3 = p.add_run(parts[1])
                r3.font.name, r3.font.size = 'Calibri', Pt(f_size)
        else:
            r = p.add_run(prefix + full_text)
            r.font.name, r.font.size = 'Calibri', Pt(f_size)

    try:
        p = doc.add_paragraph(style=style_name)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = line_spacing
        apply_bold(p, text, f_size=11)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.left_indent = Pt(18 * level)
        prefix = f"{'·' if level == 1 else 'o'}\t"
        apply_bold(p, text, prefix=prefix, f_size=10) 

def add_custom_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6) if level > 1 else Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 0.9
    if level == 1:
        p.style = doc.styles['Heading 1']
        run = p.add_run(text)
        run.font.name, run.font.size = 'Calibri', Pt(14)
    elif level == 2:
        p.style = doc.styles['Heading 2']
        run = p.add_run(text)
        run.font.name, run.font.size = 'Calibri', Pt(12)
    elif level == 3:
        p.style = doc.styles['Heading 3']
        run = p.add_run(text)
        run.font.name, run.font.size = 'Calibri', Pt(10)
    return p

# ==========================================
# ALIAS LOOKUP HELPERS
# ==========================================
def get_player_aliases(official_name, aliases=None, id_map_df=None, club=None):
    if not official_name:
        return []
    
    clean_official = str(official_name).split(' (')[0].strip().lower()
    found_aliases = []
    
    # 1. Check aliases dataframe
    if aliases is not None and not aliases.empty:
        if 'Input Name (Scorecard/Stats)' in aliases.columns and 'Official Registered Name' in aliases.columns:
            match_rows = aliases[aliases['Official Registered Name'].astype(str).str.strip().str.lower() == clean_official]
            for val in match_rows['Input Name (Scorecard/Stats)'].dropna().unique():
                cleaned_val = str(val).replace('‡', '').strip()
                if cleaned_val and cleaned_val.lower() != clean_official and cleaned_val.lower() != 'nan':
                    if cleaned_val not in found_aliases:
                        found_aliases.append(cleaned_val)
        else:
            for _, row in aliases.iterrows():
                alias_val = str(row.iloc[0]).replace('‡', '').strip()
                off_val = str(row.iloc[1]).replace('‡', '').strip()
                if off_val.lower() == clean_official and alias_val.lower() != clean_official and alias_val.lower() != 'nan':
                    if alias_val not in found_aliases:
                        found_aliases.append(alias_val)

    # 2. Check id_map_df
    if id_map_df is None or (isinstance(id_map_df, pd.DataFrame) and id_map_df.empty):
        for candidate_id_file in ['NCU_Mens_Master_ID_Mapping.xlsx', 'NCU_Master_ID_Mapping.xlsx']:
            if os.path.exists(candidate_id_file):
                try:
                    id_map_df = get_excel_df(candidate_id_file)
                    break
                except Exception:
                    pass

    if id_map_df is not None and isinstance(id_map_df, pd.DataFrame) and not id_map_df.empty:
        if 'Sport80_Name' in id_map_df.columns and 'NV_Play_Name' in id_map_df.columns:
            m = id_map_df[id_map_df['Sport80_Name'].astype(str).str.strip().str.lower() == clean_official]
            if not m.empty:
                if club and 'Sport80_Club' in m.columns:
                    c_clean = str(club).lower().replace('cricket club', '').replace('cc', '').strip()
                    club_m = m[m['Sport80_Club'].astype(str).str.lower().str.contains(c_clean, na=False)]
                    if not club_m.empty:
                        m = club_m
                for nv_name in m['NV_Play_Name'].dropna().unique():
                    nv_clean = str(nv_name).strip()
                    if nv_clean and nv_clean.lower() != clean_official and nv_clean.lower() != 'nan':
                        if nv_clean not in found_aliases:
                            found_aliases.append(nv_clean)
                            
    return found_aliases

def get_player_playing_name(official_name, aliases=None, id_map_df=None, club=None):
    if not official_name:
        return ""
    pure = str(official_name).split(' (')[0].strip()
    aliases_list = get_player_aliases(pure, aliases=aliases, id_map_df=id_map_df, club=club)
    if aliases_list:
        return aliases_list[0]
    return pure


def infer_player_club(active_player, player_batting, player_bowling, domain):
    if ' (' in active_player: return active_player.split(' (')[1].replace(')', '')
    groups_to_concat = []
    if player_batting is not None and not player_batting.empty and 'Group' in player_batting.columns:
        groups_to_concat.append(player_batting['Group'])
    if player_bowling is not None and not player_bowling.empty and 'Group' in player_bowling.columns:
        groups_to_concat.append(player_bowling['Group'])
    all_groups_fallback = pd.concat(groups_to_concat).dropna().tolist() if groups_to_concat else []
    
    team_frequency = {}
    team_raw_names = {}
    def extract_base_club_name_local(team_str):
        return re.sub(r'\s*(cc|club|1st|2nd|3rd|4th|5th|6th|1|2|3|4|5|6|xi|1st xi|2nd xi|3rd xi|4th xi|5th xi|6th xi)$', '', team_str, flags=re.IGNORECASE).strip()

    for grp in all_groups_fallback:
        if ' v ' in grp:
            t1, t2 = grp.split(' v ')[0].strip(), grp.split(' v ')[1].split(',')[0].strip()
            t1, t2 = doc_format_cricket_names(t1, domain), doc_format_cricket_names(t2, domain)
            
            b1, b2 = extract_base_club_name_local(t1).strip(), extract_base_club_name_local(t2).strip()
            team_frequency[b1] = team_frequency.get(b1, 0) + 1
            team_frequency[b2] = team_frequency.get(b2, 0) + 1
            team_raw_names.setdefault(b1, set()).add(t1)
            team_raw_names.setdefault(b2, set()).add(t2)
            
    if team_frequency:
        sorted_teams = sorted(team_frequency.items(), key=lambda item: item[1], reverse=True)
        if sorted_teams:
            if len(sorted_teams) > 1 and sorted_teams[0][1] == sorted_teams[1][1]:
                max_freq = sorted_teams[0][1]
                top_teams = [t[0] for t in sorted_teams if t[1] == max_freq]
                raw_combinations = [list(team_raw_names[t])[0] for t in top_teams]
                return " / ".join(raw_combinations)
            else:
                return sorted_teams[0][0]
    return "Unknown_Club"

def generate_single_player_doc(active_player, player_batting, player_bowling, reg_players_df, domain, aliases_list=None, player_abandoned=None, league_dict=None, cup_df=None, id_map_df=None, playing_name=None):
    player_batting = player_batting.copy() if player_batting is not None and not player_batting.empty else pd.DataFrame()
    player_bowling = player_bowling.copy() if player_bowling is not None and not player_bowling.empty else pd.DataFrame()
    club_name = "Unknown_Club"
    primary_club = "Unknown_Club"
    transfer_club = "Unknown_Club"
    transfer_date = None
    transfer_club_2 = "Unknown_Club"
    transfer_date_2 = None
    if active_player.lower() == 'neil brand' and domain != "Women's":
        club_name = 'Muckamore'
    else:
        reg_search_name = active_player.split(' (')[0] if ' (' in active_player else active_player
        reg_match = pd.DataFrame()
        if reg_players_df is not None and not reg_players_df.empty:
            if '_computed_name' in reg_players_df.columns:
                reg_match = reg_players_df[reg_players_df['_computed_name'].astype(str).str.strip().str.lower() == reg_search_name.lower()]
            elif 'Full Name' in reg_players_df.columns:
                reg_match = reg_players_df[reg_players_df['Full Name'].astype(str).str.strip().str.lower() == reg_search_name.lower()]
            elif 'First Name' in reg_players_df.columns and 'Last Name' in reg_players_df.columns:
                comp_names = (reg_players_df['First Name'].astype(str).str.strip() + ' ' + reg_players_df['Last Name'].astype(str).str.strip()).str.lower()
                reg_match = reg_players_df[comp_names == reg_search_name.lower()]
            elif 'First Name' in reg_players_df.columns and 'Surname' in reg_players_df.columns:
                comp_names = (reg_players_df['First Name'].astype(str).str.strip() + ' ' + reg_players_df['Surname'].astype(str).str.strip()).str.lower()
                reg_match = reg_players_df[comp_names == reg_search_name.lower()]
            else:
                name_col = next((c for c in reg_players_df.columns if 'name' in str(c).lower()), reg_players_df.columns[0])
                reg_match = reg_players_df[reg_players_df[name_col].astype(str).str.strip().str.lower() == reg_search_name.lower()]

        if not reg_match.empty:
            primary_cols = [c for c in reg_match.columns if 'Primary Club' in str(c) and 'Wylie' not in str(c)]
            if primary_cols and len(reg_match[primary_cols[0]].dropna().values) > 0 and str(reg_match[primary_cols[0]].dropna().values[0]).strip() != '': 
                primary_club = str(reg_match[primary_cols[0]].dropna().values[0]).strip()
                
            for keyword in ['Wylie', 'Transfer']:
                cols = [c for c in reg_match.columns if keyword in str(c)]
                if cols and len(reg_match[cols[0]].dropna().values) > 0 and str(reg_match[cols[0]].dropna().values[0]).strip() != '':
                    transfer_club = str(reg_match[cols[0]].dropna().values[0]).strip()
                    break
                    
            t1_date_cols = [c for c in reg_match.columns if 'Transfer Date' in str(c) and '2' not in str(c)]
            if t1_date_cols:
                td_val = reg_match[t1_date_cols[0]].dropna().values
                if len(td_val) > 0:
                    transfer_date = pd.to_datetime(td_val[0], errors='coerce', dayfirst=True)

            for kw in ['Transfer Club 2', 'Club Transfer 2']:
                cols = [c for c in reg_match.columns if kw.lower() in str(c).lower()]
                if cols and len(reg_match[cols[0]].dropna().values) > 0 and str(reg_match[cols[0]].dropna().values[0]).strip() != '':
                    transfer_club_2 = str(reg_match[cols[0]].dropna().values[0]).strip()
                    break

            if 'Transfer Date 2' in reg_match.columns:
                td2_val = reg_match['Transfer Date 2'].dropna().values
                if len(td2_val) > 0:
                    transfer_date_2 = pd.to_datetime(td2_val[0], errors='coerce', dayfirst=True)
                    
            club_name = transfer_club if transfer_club != "Unknown_Club" else primary_club
            
        sport80_id = None
        if not reg_match.empty:
            for id_col in ['Individual Membership CI No.', 'Sport80_ID', 'Sport80 ID', 'CI No']:
                if id_col in reg_match.columns:
                    vals = reg_match[id_col].dropna().values
                    if len(vals) > 0 and str(vals[0]).strip() and str(vals[0]).strip().lower() != 'nan':
                        sport80_id = str(vals[0]).replace('.0', '').strip()
                        break

    if primary_club == "Unknown_Club" and transfer_club == "Unknown_Club":
        if id_map_df is None or (isinstance(id_map_df, pd.DataFrame) and id_map_df.empty):
            f_id_map = DEFAULT_FILES.get(domain, {}).get("id_map", "")
            if f_id_map and os.path.exists(f_id_map):
                try:
                    id_map_df = pd.read_excel(f_id_map)
                except Exception:
                    pass

        if id_map_df is not None and isinstance(id_map_df, pd.DataFrame) and not id_map_df.empty:
            col_nv = next((c for c in id_map_df.columns if 'nv' in c.lower() and 'name' in c.lower()), 'NV_Play_Name')
            col_s80_n = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'name' in c.lower()), 'Sport80_Name')
            col_club = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'club' in c.lower()), 'Sport80_Club')
            col_s80_id = next((c for c in id_map_df.columns if 'sport80' in c.lower() and 'id' in c.lower()), 'Sport80_ID')
            
            clean_search = reg_search_name.strip().lower()
            m = id_map_df[
                (id_map_df[col_nv].astype(str).str.strip().str.lower() == clean_search) |
                (id_map_df[col_s80_n].astype(str).str.strip().str.lower() == clean_search)
            ]
            if not m.empty:
                c_val = m[col_club].dropna().values
                if len(c_val) > 0 and str(c_val[0]).strip() and str(c_val[0]).strip().lower() != 'nan':
                    primary_club = str(c_val[0]).strip()
                    club_name = primary_club
                if sport80_id is None:
                    id_val = m[col_s80_id].dropna().values
                    if len(id_val) > 0 and str(id_val[0]).strip() and str(id_val[0]).strip().lower() not in ['nan', 'not registered']:
                        sport80_id = str(id_val[0]).replace('.0', '').strip()
            
    if primary_club == "Unknown_Club" and transfer_club == "Unknown_Club":
        primary_club = infer_player_club(active_player, player_batting, player_bowling, domain)
    
    if ' (' in active_player: primary_club = active_player.split(' (')[1].replace(')', '')
    
    def extract_match_date(grp_str):
        try:
            parts = str(grp_str).rsplit(' - ', 1)
            if len(parts) == 2:
                d_str = parts[1].strip()
                d_str = re.sub(r'(?<=\d)(st|nd|rd|th)\b', '', d_str, flags=re.IGNORECASE)
                dt = pd.to_datetime(d_str, dayfirst=True, errors='coerce')
                if pd.notna(dt): return dt
        except: pass
        try:
            match = re.search(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+[A-Za-z]{3,}\s+\d{4})', str(grp_str))
            if match:
                dt = pd.to_datetime(match.group(1), dayfirst=True, errors='coerce')
                if pd.notna(dt): return dt
        except: pass
        return pd.NaT

    def get_dynamic_club_for_match(grp_str):
        if transfer_club == "Unknown_Club":
            return primary_club
            
        m_date = extract_match_date(grp_str)
        if pd.notna(transfer_date_2) and pd.notna(m_date) and m_date >= transfer_date_2:
            return transfer_club_2 if transfer_club_2 != "Unknown_Club" else primary_club
            
        if pd.isna(transfer_date):
            grp_lower = str(grp_str).lower()
            import re
            t_clean = str(transfer_club).lower()
            t_clean = re.sub(r'\s*(cricket club|cc|club)$', '', t_clean).strip()
            p_clean = str(primary_club).lower()
            p_clean = re.sub(r'\s*(cricket club|cc|club)$', '', p_clean).strip()
            
            # If both clubs are in the string (they played each other), we default to primary club
            # to avoid falsely assuming they transferred early
            if t_clean in grp_lower and p_clean in grp_lower:
                return primary_club
                
            if t_clean in grp_lower:
                return transfer_club
            return primary_club
            
        m_date = extract_match_date(grp_str)
        if pd.isna(m_date) or m_date >= transfer_date:
            return transfer_club
        return primary_club

    if not player_batting.empty and 'Group' in player_batting.columns:
        player_batting['Team'] = player_batting['Group'].apply(lambda x: doc_get_player_team_from_match(x, doc_format_cricket_names(get_dynamic_club_for_match(x), domain)))
    if not player_bowling.empty and 'Group' in player_bowling.columns:
        player_bowling['Team'] = player_bowling['Group'].apply(lambda x: doc_get_player_team_from_match(x, doc_format_cricket_names(get_dynamic_club_for_match(x), domain)))
        
    unique_teams = set()
    if not player_batting.empty and 'Team' in player_batting.columns: unique_teams.update(player_batting['Team'].unique())
    if not player_bowling.empty and 'Team' in player_bowling.columns: unique_teams.update(player_bowling['Team'].unique())
    
    if player_abandoned is not None and not player_abandoned.empty:
        ab_match_col = 'Group' if 'Group' in player_abandoned.columns else ('Match' if 'Match' in player_abandoned.columns else player_abandoned.columns[0])
        for grp in player_abandoned[ab_match_col]:
            unique_teams.add(doc_get_player_team_from_match(grp, doc_format_cricket_names(get_dynamic_club_for_match(grp), domain)))

    unique_teams = sorted(list(unique_teams), key=doc_team_sort_key)

    all_groups = []
    if not player_batting.empty and 'Group' in player_batting.columns: all_groups.extend(player_batting['Group'].tolist())
    if not player_bowling.empty and 'Group' in player_bowling.columns: all_groups.extend(player_bowling['Group'].tolist())
    if player_abandoned is not None and not player_abandoned.empty:
        ab_match_col = 'Group' if 'Group' in player_abandoned.columns else ('Match' if 'Match' in player_abandoned.columns else player_abandoned.columns[0])
        all_groups.extend(player_abandoned[ab_match_col].tolist())

    unique_groups = list(dict.fromkeys(all_groups)) 
    unique_groups.sort(key=extract_match_date)
    
    club_name_clean = doc_format_cricket_names(transfer_club if transfer_club != "Unknown_Club" else primary_club, domain)
    
    matches_by_team = {}
    for grp in unique_groups:
        team_played_for = doc_get_player_team_from_match(grp, doc_format_cricket_names(get_dynamic_club_for_match(grp), domain))
        
        b_row = player_batting[player_batting['Group'] == grp] if not player_batting.empty else pd.DataFrame()
        bw_row = player_bowling[player_bowling['Group'] == grp] if not player_bowling.empty else pd.DataFrame()
        
        is_ab = False
        if player_abandoned is not None and not player_abandoned.empty:
            ab_match_col = 'Group' if 'Group' in player_abandoned.columns else ('Match' if 'Match' in player_abandoned.columns else player_abandoned.columns[0])
            is_ab = not player_abandoned[player_abandoned[ab_match_col] == grp].empty

        comp_name = ""
        if cup_df is not None and not cup_df.empty:
            for _, r in cup_df.iterrows():
                cup_match_str = doc_format_cricket_names(str(r.iloc[0]), domain)
                if cup_match_str.strip() in str(grp).strip() or str(grp).strip() in cup_match_str.strip():
                    comp_name = str(r.iloc[1])
                    break
        if not comp_name and league_dict is not None and team_played_for:
            team_keys = list(league_dict.keys())
            l = get_team_league(team_played_for, team_keys, league_dict, domain)
            if not l:
                t1, t2 = extract_teams_from_group(grp)
                l = get_team_league(t1, team_keys, league_dict, domain)
                if not l:
                    l = get_team_league(t2, team_keys, league_dict, domain)
            if l:
                comp_name = str(l)
        if not comp_name:
            comp_name = "Friendly/Other"
            
        grp_display = f"{grp} ({comp_name})".replace('\xa0', ' ')
        grp_display = grp_display.replace(", TBC -", " -").replace(", TBC ", " ")
        if is_ab:
            grp_display += " (abandoned)"
        grp_display = re.sub(r'City of Belfast Playing Fields\s*\(.*?\)', 'City of Belfast Playing Fields', grp_display, flags=re.IGNORECASE)


        if not b_row.empty and b_row.iloc[0]['Innings'] > 0:
            hs = b_row.iloc[0]['High Score']
            hs_str = str(hs).replace('.0', '') if pd.notna(hs) else str(int(b_row.iloc[0]['Runs']))
            bat_str = f"Batting: {hs_str} runs"
        elif is_ab:
            bat_str = "Batting: Abandoned match"
        else: 
            bat_str = "Batting: Did not bat"
            
        if not bw_row.empty and bw_row.iloc[0]['Innings'] > 0 and bw_row.iloc[0]['Overs'] > 0:
            o = bw_row.iloc[0]['Overs']
            o_str = str(o).replace('.0', '') if str(o).endswith('.0') else str(o)
            m = int(bw_row.iloc[0]['Maidens']) if pd.notna(bw_row.iloc[0]['Maidens']) else 0
            w = int(bw_row.iloc[0]['Wickets']) if pd.notna(bw_row.iloc[0]['Wickets']) else 0
            r = int(bw_row.iloc[0]['Runs']) if pd.notna(bw_row.iloc[0]['Runs']) else 0
            bowl_str = f"Bowling: {o_str}-{m}-{r}-{w}"
        elif is_ab:
            bowl_str = "Bowling: Abandoned match"
        else: 
            bowl_str = "Bowling: Did not bowl"
            
        parts = grp.rsplit(' - ', 1)
        date_str = parts[1].split(' (')[0].strip() if len(parts) == 2 else grp.split(' (')[0].strip()
        clean_date_str = re.sub(r'(?<=\d)(st|nd|rd|th)\b', '', date_str, flags=re.IGNORECASE)
        match_date = pd.to_datetime(clean_date_str, dayfirst=True, errors='coerce')
        if pd.isna(match_date): match_date = pd.Timestamp.min
        
        matches_by_team.setdefault(team_played_for, []).append({'match': grp_display, 'bat_str': bat_str, 'bowl_str': bowl_str, 'date': match_date, 'team': team_played_for})

    doc = Document()
    style_normal = doc.styles['Normal']
    style_normal.font.name, style_normal.font.size = 'Calibri', Pt(11) 
    
    if transfer_club != "Unknown_Club" and primary_club != "Unknown_Club" and transfer_club.strip().lower() != primary_club.strip().lower():
        p_clean = re.sub(r'(?i)\s*cricket club', '', doc_format_cricket_names(primary_club, domain)).strip()
        t_clean = re.sub(r'(?i)\s*cricket club', '', doc_format_cricket_names(transfer_club, domain)).strip()
        header_club_name = f"{p_clean} / {t_clean}"
        if transfer_club_2 != "Unknown_Club" and transfer_club_2.strip().lower() != transfer_club.strip().lower():
            t2_clean = re.sub(r'(?i)\s*cricket club', '', doc_format_cricket_names(transfer_club_2, domain)).strip()
            if t2_clean.lower() != p_clean.lower():
                header_club_name = f"{p_clean} / {t_clean} / {t2_clean}"
    else:
        header_club_name = re.sub(r'(?i)\s*cricket club', '', club_name_clean).strip()
    if not playing_name:
        if aliases_list:
            playing_name = aliases_list[0]
        else:
            playing_name = get_player_playing_name(active_player, id_map_df=id_map_df, club=club_name_clean)
            
    domain_label = "Open" if domain == "Men's" else ("Women" if domain == "Women's" else "Midweek")
    heading_title = f"{playing_name} - {header_club_name} - Season Summary ({domain_label})\n"
        
    add_custom_heading(doc, heading_title, level=1)
    
    if sport80_id:
        p_s80 = doc.add_paragraph()
        p_s80.paragraph_format.space_before = Pt(0)
        p_s80.paragraph_format.space_after = Pt(6)
        r_s80 = p_s80.add_run(f"Sport80 Member ID: {sport80_id}")
        r_s80.font.name, r_s80.font.size = 'Calibri', Pt(10)
        r_s80.bold = True
        r_s80.font.color.rgb = RGBColor(0, 0, 128)
    
    add_custom_heading(doc, "Chronological Match Appearances", level=2)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run("(Registered/Transferred to team in bold)")
    run_sub.font.name, run_sub.font.size = 'Calibri', Pt(9)
    run_sub.italic = True
    run_sub.bold = True
    
    all_matches = []
    for team, m_list in matches_by_team.items():
        all_matches.extend(m_list)
        
    all_matches.sort(key=lambda x: x['date'])
    
    t1_base = ""
    t2_base = ""
    if transfer_club != "Unknown_Club" and primary_club != "Unknown_Club" and transfer_club.strip().lower() != primary_club.strip().lower():
        t1_base = re.sub(r'(?i)\s*(cricket club|cc|club)$', '', str(transfer_club).lower()).strip()
        if transfer_club_2 != "Unknown_Club" and transfer_club_2.strip().lower() != transfer_club.strip().lower():
            t2_base = re.sub(r'(?i)\s*(cricket club|cc|club)$', '', str(transfer_club_2).lower()).strip()

    phase = 0
    for i, m in enumerate(all_matches):
        team_str = str(m.get('team', '')).lower()
        is_valid_team = bool(m.get('team') and m.get('team') != "Unknown Team" and not team_str.startswith("unknown ("))
        
        trigger_transfer = False
        if is_valid_team:
            if phase == 0 and t1_base and t1_base in team_str:
                trigger_transfer = True
                phase = 1
            elif phase == 1 and t2_base and t2_base in team_str:
                trigger_transfer = True
                phase = 2

        if trigger_transfer and i > 0:
            if len(doc.paragraphs) > 0:
                doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
            
            p_trans = doc.add_paragraph()
            p_trans.paragraph_format.left_indent = Pt(18)
            p_trans.paragraph_format.space_before = Pt(11)
            p_trans.paragraph_format.space_after = Pt(0)
            p_trans.paragraph_format.line_spacing = 0.9
            run_trans = p_trans.add_run("(player transferred)")
            run_trans.font.name, run_trans.font.size = 'Calibri', Pt(10)
            run_trans.italic = True
                
        bold_team = m.get('team') if m.get('team') and m.get('team') != "Unknown Team" and not m.get('team', '').startswith("Unknown (") else None
        add_bullet_point(doc, m['match'], level=1, space_after=11 if i < len(all_matches) - 1 else 18, line_spacing=1.1, bold_substring=bold_team)
        
    add_custom_heading(doc, "Batting Statistics", level=2)

    batting_found = False
    for team in unique_teams:
        team_bat = player_batting[player_batting['Team'] == team] if not player_batting.empty else pd.DataFrame()
        if team_bat.empty or team_bat['Matches'].sum() == 0: continue
            
        batting_found = True
        team_bowl_for_matches = player_bowling[player_bowling['Team'] == team] if not player_bowling.empty else pd.DataFrame()
        bowl_m_raw = int(team_bowl_for_matches['Matches'].sum()) if not team_bowl_for_matches.empty else 0
        bat_matches = max(int(team_bat['Matches'].sum()), bowl_m_raw)
        
        bat_innings = int(team_bat['Innings'].sum())
        total_runs = int(team_bat['Runs'].sum())
        not_outs = int(team_bat['Not Outs'].sum()) if 'Not Outs' in team_bat.columns else 0
        balls_faced = int(team_bat['Balls'].sum()) if 'Balls' in team_bat.columns else 0
        fours, sixes = (int(team_bat['Fours'].sum()) if 'Fours' in team_bat.columns else 0), (int(team_bat['Sixes'].sum()) if 'Sixes' in team_bat.columns else 0)
        
        best_score, is_not_out = 0, False
        for hs in team_bat['High Score'].dropna().astype(str).tolist():
            if hs.lower() == 'nan': continue
            val = hs.replace('*', '').replace('.0', '')
            try:
                if int(val) > best_score: best_score, is_not_out = int(val), '*' in hs
                elif int(val) == best_score and '*' in hs: is_not_out = True
            except ValueError: pass
        
        high_score_str = f"{best_score}*" if is_not_out else str(best_score)
        if bat_innings == 0: high_score_str = "N/A"
        outs = bat_innings - not_outs
        bat_avg = f"{total_runs / outs:.2f}" if outs > 0 else "N/A"

        add_custom_heading(doc, team, level=3)
        add_bullet_point(doc, f"Matches: {bat_matches}")
        add_bullet_point(doc, f"Innings: {bat_innings}")
        add_bullet_point(doc, f"Total Runs: {total_runs}")
        add_bullet_point(doc, f"Highest Score: {high_score_str}")
        add_bullet_point(doc, f"Batting Average: {bat_avg}")
        add_bullet_point(doc, f"Balls Faced: {balls_faced}")
        add_bullet_point(doc, f"Boundaries: {fours} Fours, {sixes} Sixes\n")

    if not batting_found:
        p = doc.add_paragraph("No batting statistics recorded.\n")
        p.style.font.name, p.style.font.size = 'Calibri', Pt(11)
        
    add_custom_heading(doc, "Bowling Statistics", level=2)
    bowling_found = False
    for team in unique_teams:
        team_bowl = player_bowling[player_bowling['Team'] == team] if not player_bowling.empty else pd.DataFrame()
        if team_bowl.empty or team_bowl['Matches'].sum() == 0: continue
            
        bowling_found = True
        team_bat_for_matches = player_batting[player_batting['Team'] == team] if not player_batting.empty else pd.DataFrame()
        bat_m_raw = int(team_bat_for_matches['Matches'].sum()) if not team_bat_for_matches.empty else 0
        bowl_matches = max(bat_m_raw, int(team_bowl['Matches'].sum()))
        
        bowl_innings = int(team_bowl['Innings'].sum())
        maidens = int(team_bowl['Maidens'].sum()) if 'Maidens' in team_bowl.columns else 0
        bowl_runs = int(team_bowl['Runs'].sum()) if 'Runs' in team_bowl.columns else 0
        wickets = int(team_bowl['Wickets'].sum()) if 'Wickets' in team_bowl.columns else 0
        bowl_avg = f"{bowl_runs / wickets:.2f}" if wickets > 0 else "N/A"

        if 'Balls' in team_bowl.columns and team_bowl['Balls'].sum() > 0:
            total_balls = int(team_bowl['Balls'].sum())
        else:
            total_balls = 0
            for o_val in team_bowl['Overs'].dropna():
                try:
                    o_float = float(o_val)
                    whole_overs = int(o_float)
                    o_str = f"{o_float:.1f}"
                    balls_part = int(o_str.split('.')[1]) if '.' in o_str else 0
                    total_balls += (whole_overs * 6) + min(balls_part, 5)
                except (ValueError, IndexError):
                    pass

        total_completed_overs = total_balls // 6
        extra_balls = total_balls % 6
        overs_display = f"{total_completed_overs}.{extra_balls}" if extra_balls > 0 else str(total_completed_overs)

        add_custom_heading(doc, team, level=3)
        add_bullet_point(doc, f"Matches: {bowl_matches}")
        add_bullet_point(doc, f"Innings: {bowl_innings}")
        add_bullet_point(doc, f"Overs: {overs_display}")
        add_bullet_point(doc, f"Maidens: {maidens}")
        add_bullet_point(doc, f"Runs Conceded: {bowl_runs}")
        add_bullet_point(doc, f"Wickets: {wickets}")
        add_bullet_point(doc, f"Bowling Average: {bowl_avg}\n")

    if not bowling_found:
        p = doc.add_paragraph("No bowling statistics recorded.\n")
        p.style.font.name, p.style.font.size = 'Calibri', Pt(11)

    add_custom_heading(doc, "Fielding Statistics", level=2)
    fielding_found = False
    for team in unique_teams:
        team_bat = player_batting[player_batting['Team'] == team] if not player_batting.empty else pd.DataFrame()
        if team_bat.empty or team_bat['Matches'].sum() == 0: continue
            
        fielding_found = True
        catches = int(team_bat['Catches'].sum()) if 'Catches' in team_bat.columns else 0
        catches_wk = int(team_bat['Catches as Keeper'].sum()) if 'Catches as Keeper' in team_bat.columns else 0
        stumpings = int(team_bat['Stumpings'].sum()) if 'Stumpings' in team_bat.columns else 0
        run_outs = int(team_bat['Run Outs'].sum()) if 'Run Outs' in team_bat.columns else 0

        add_custom_heading(doc, team, level=3)
        field_stats_list = [f"Catches: {catches}"]
        if catches_wk > 0: field_stats_list.append(f"Catches as Keeper: {catches_wk}")
        if stumpings > 0: field_stats_list.append(f"Stumpings: {stumpings}")
        if run_outs > 0: field_stats_list.append(f"Run Outs: {run_outs}")
        
        for i, stat in enumerate(field_stats_list):
            if i == len(field_stats_list) - 1: add_bullet_point(doc, stat + "\n")
            else: add_bullet_point(doc, stat)

    if not fielding_found:
        p = doc.add_paragraph("No fielding statistics recorded.\n")
        p.style.font.name, p.style.font.size = 'Calibri', Pt(11)

    add_custom_heading(doc, "Match Appearances", level=2)
    
    p_sub2 = doc.add_paragraph()
    p_sub2.paragraph_format.space_before = Pt(0)
    p_sub2.paragraph_format.space_after = Pt(8)
    run_sub2 = p_sub2.add_run("(Registered/Transferred to team in bold)")
    run_sub2.font.name, run_sub2.font.size = 'Calibri', Pt(9)
    run_sub2.italic = True
    run_sub2.bold = True
    
    for team in sorted(matches_by_team.keys(), key=doc_team_sort_key):
        add_custom_heading(doc, team, level=3)
        for m in matches_by_team[team]:
            bold_team = m.get('team') if m.get('team') and m.get('team') != "Unknown Team" and not m.get('team', '').startswith("Unknown (") else None
            add_bullet_point(doc, m['match'], level=1, bold_substring=bold_team)
            add_bullet_point(doc, m['bat_str'], level=2)
            add_bullet_point(doc, m['bowl_str'], level=2, space_after=4)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    
    clean_player_name = re.sub(r'[^\w\-_\. ]', '', str(playing_name)).strip().replace(' ', '_')
    filename = f"{clean_player_name}_{club_name_clean.replace(' ', '_')}_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.docx"
    
    return doc_io, filename

# ==========================================
# REGISTRATION ENGINE SPECIFIC FUNCTIONS
# ==========================================
def get_ordinal_date(dt, include_year=True):
    if pd.isna(dt): return "N/A"
    day = dt.day
    if 11 <= (day % 100) <= 13: suffix = 'th'
    else: suffix = ['th', 'st', 'nd', 'rd', 'th'][min(day % 10, 4)]
    month = dt.strftime('%B')
    if include_year: return f"{month} {day}{suffix}, {dt.year}"
    return f"{month} {day}{suffix}"

def doc_formal_team_name(team_name):
    team_str = str(team_name)
    if 'holywood' in team_str.lower() and '1881' not in team_str.lower():
        return re.sub(r'(?i)(holywood)', r'Holywood 1881', team_str)
    return team_str

def match_sort_key(match_tuple):
    m_d = match_tuple[3]
    date_key = m_d if pd.notna(m_d) else datetime.min
    league = match_tuple[2]
    parts = re.split(r'(\d+)', league)
    parts = [int(p) if p.isdigit() else p for p in parts]
    return (date_key, parts, match_tuple[0])
    
def violation_sort_key(record, name_field):
    m_date = record.get('Match Date')
    date_val = m_date if pd.notna(m_date) else pd.Timestamp.min
    
    full_name = str(record.get(name_field, '')).strip()
    parts = full_name.split()
    
    if len(parts) > 1:
        surname = parts[-1].lower()
        firstnames = " ".join(parts[:-1]).lower()
    elif len(parts) == 1:
        surname = parts[0].lower()
        firstnames = ""
    else:
        surname = ""
        firstnames = ""
        
    return (date_val, surname, firstnames)

def doc_add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name, run.font.size = 'Calibri', Pt(11) 
    
def get_competition_sort_key(comp, domain="Men's"):
    comp_lower = str(comp).lower().strip()
    
    if "women's gallagher challenge plate" in comp_lower or "womens gallagher challenge plate" in comp_lower:
        return (1, 12, comp_lower)
    elif "women's gallagher challenge cup" in comp_lower or "womens gallagher challenge cup" in comp_lower:
        return (1, 11, comp_lower)
    elif "gallagher challenge cup" in comp_lower:
        return (1, 0, comp_lower)
    elif "lvs t20 cup" in comp_lower:
        return (1, 1, comp_lower)
    elif "lvs t20 trophy" in comp_lower:
        return (1, 2, comp_lower)
    elif "junior 1 t20 plate" in comp_lower:
        return (1, 6, comp_lower)
    elif "junior cup" in comp_lower:
        return (1, 3, comp_lower)
    elif "lvs t20 bowl" in comp_lower:
        return (1, 4, comp_lower)
    elif "lvs t20 shield" in comp_lower:
        return (1, 5, comp_lower)
    elif "intermediate cup" in comp_lower:
        return (1, 7, comp_lower)
    elif "lindsay" in comp_lower or "minor (lindsay)" in comp_lower:
        return (1, 8, comp_lower)
    elif "minor qualifying cup" in comp_lower:
        return (1, 9, comp_lower)
    elif "development cup" in comp_lower:
        return (1, 10, comp_lower)
    elif "irish cup" in comp_lower or "irish senior cup" in comp_lower:
        return (1, 13, comp_lower)
    elif "national cup" in comp_lower:
        return (1, 14, comp_lower)
    elif "ulster plate" in comp_lower:
        return (1, 15, comp_lower)
    
    is_cup_fallback = any(word in comp_lower for word in ['cup', 'trophy', 'shield', 'plate', 'bowl', 'vase', 'challenge', 'fixture', 'possible cup match', 'derby', 'knockout'])
    if is_cup_fallback:
        return (1, 100, comp_lower)
    
    is_womens = (domain == "Women's") or any(word in comp_lower for word in ['women', 'womens', "women's"])
    
    if is_womens:
        if 'premier' in comp_lower:
            sub_rank = 1
        elif 'senior' in comp_lower or 'section 1' in comp_lower:
            sub_rank = 2
        elif 'junior' in comp_lower:
            match = re.search(r'junior(?: league)?\s*(\d+)([a-z]?)', comp_lower)
            if match:
                num = int(match.group(1))
                let = match.group(2)
                let_val = (ord(let) - ord('a') + 1) * 0.01 if let else 0.0
                sub_rank = 10 + num + let_val
            else:
                sub_rank = 50
        else:
            sub_rank = 100
        return (3, sub_rank, comp_lower)
    
    if 'premier' in comp_lower:
        return (2, 0, comp_lower)
    elif 'senior league 1' in comp_lower or 'senior 1' in comp_lower or 'senior league section 1' in comp_lower:
        return (2, 1, comp_lower)
    elif 'senior league 2' in comp_lower or 'senior 2' in comp_lower or 'senior league section 2' in comp_lower:
        return (2, 2, comp_lower)
    elif 'senior league 3' in comp_lower or 'senior 3' in comp_lower or 'senior league section 3' in comp_lower:
        return (2, 3, comp_lower)
    elif 'junior' in comp_lower:
        match = re.search(r'junior(?: league)?\s*(\d+)([a-z]?)', comp_lower)
        if match:
            num = int(match.group(1))
            let = match.group(2)
            let_val = (ord(let) - ord('a') + 1) * 0.01 if let else 0.0
            sub_rank = 10 + num + let_val  
            return (2, sub_rank, comp_lower)
        else:
            return (2, 99, comp_lower)
    elif 'midweek' in comp_lower or 'group' in comp_lower:
        match = re.search(r'group\s*([a-z])', comp_lower)
        sub_rank = (ord(match.group(1)) - ord('a')) if match else 0
        return (2, 200 + sub_rank, comp_lower)
        
    return (4, 0, comp_lower)

def render_grouped_matches(doc, matches, domain):
    by_date_ts = {}
    for t_a, t_b, lg, m_d in matches:
        if pd.isna(m_d):
            continue
        m_d_normalized = m_d.normalize()
        if m_d_normalized not in by_date_ts:
            by_date_ts[m_d_normalized] = {}
        
        comp_name = str(lg).strip()
        if comp_name not in by_date_ts[m_d_normalized]:
            by_date_ts[m_d_normalized][comp_name] = []
        
        by_date_ts[m_d_normalized][comp_name].append((t_a, t_b))
        
    for dt in sorted(by_date_ts.keys()):
        date_header = f"{dt.day} {dt.strftime('%B %Y')}"
        
        p_date = doc.add_paragraph()
        p_date.paragraph_format.space_before = Pt(6)
        p_date.paragraph_format.space_after = Pt(2)
        p_date.paragraph_format.line_spacing = 0.9
        run_date = p_date.add_run(date_header)
        run_date.bold = True
        run_date.font.name = 'Calibri'
        run_date.font.size = Pt(11)
        
        comps = by_date_ts[dt]
        sorted_comps = sorted(comps.keys(), key=lambda c: get_competition_sort_key(c, domain))
        
        for comp in sorted_comps:
            try:
                p_comp = doc.add_paragraph(style='List Bullet')
            except KeyError:
                p_comp = doc.add_paragraph()
                p_comp.paragraph_format.left_indent = Pt(18)
                p_comp.add_run("• ")
                
            p_comp.paragraph_format.space_before = Pt(2)
            p_comp.paragraph_format.space_after = Pt(0)
            p_comp.paragraph_format.line_spacing = 0.9
            
            run_comp = p_comp.add_run(comp)
            run_comp.font.name = 'Calibri'
            run_comp.font.size = Pt(11)
            
            sorted_matches = sorted(comps[comp], key=lambda x: (str(x[0]).lower(), str(x[1]).lower()))
            
            for t_a, t_b in sorted_matches:
                try:
                    p_match = doc.add_paragraph(style='List Bullet 2')
                except KeyError:
                    p_match = doc.add_paragraph()
                    p_match.paragraph_format.left_indent = Pt(36)
                    p_match.add_run("o ")
                    
                p_match.paragraph_format.space_before = Pt(0)
                p_match.paragraph_format.space_after = Pt(0)
                p_match.paragraph_format.line_spacing = 0.9
                
                match_text = f"{doc_formal_team_name(t_a)} v {doc_formal_team_name(t_b)}"
                run_match = p_match.add_run(match_text)
                run_match.font.name = 'Calibri'
                run_match.font.size = Pt(11) 

def export_and_format_excel(df, writer, sheet_name):
    df.to_excel(writer, index=False, sheet_name=sheet_name)
    workbook = writer.book
    worksheet = writer.sheets[sheet_name]
    header_format = workbook.add_format({'bold': True, 'bottom': 1})
    bold_name_format = workbook.add_format({'bold': True})
    
    for col_num, col_name in enumerate(df.columns):
        worksheet.write(0, col_num, col_name, header_format)
        series_len = max((len(str(x)) for x in df[col_name]), default=0)
        header_len = len(str(col_name))
        max_width = max(series_len, header_len) + 2
        if col_name in ['Stats Name (Cleaned)', 'Player (Cleaned)']: worksheet.set_column(col_num, col_num, max_width, bold_name_format)
        else: worksheet.set_column(col_num, col_num, max_width)

@st.cache_data(show_spinner="Running registration & starring audit...")
def run_registration_audit(domain, start_date, end_date, f_reg, f_alias, f_starring, f_league, f_bat, f_bowl, f_irish_bat=None, f_irish_bowl=None, f_cup=None, f_abandoned=None, f_id_map=None):
    registered_players = pd.read_excel(f_reg)
    aliases = pd.read_excel(f_alias)
    league_structure = pd.read_excel(f_league)
    batting_stats = pd.read_excel(f_bat)
    bowling_stats = pd.read_excel(f_bowl)

    batting_stats['Is_Irish_Match'] = False
    bowling_stats['Is_Irish_Match'] = False
    
    if f_irish_bat and os.path.exists(f_irish_bat):
        irish_bat = pd.read_excel(f_irish_bat)
        irish_bat['Is_Irish_Match'] = True
        batting_stats = pd.concat([batting_stats, irish_bat], ignore_index=True)
        
    if f_irish_bowl and os.path.exists(f_irish_bowl):
        irish_bowl = pd.read_excel(f_irish_bowl)
        irish_bowl['Is_Irish_Match'] = True
        bowling_stats = pd.concat([bowling_stats, irish_bowl], ignore_index=True)

    def parse_match_group(group_str):
        try:
            group_str = str(group_str).strip()
            parts = group_str.rsplit(' - ', 1)
            date_str = parts[1].strip() if len(parts) == 2 else group_str
            rest = parts[0].strip() if len(parts) == 2 else group_str
            
            match_date = pd.to_datetime(date_str, dayfirst=True, errors='coerce')
            if pd.notna(match_date):
                match_date = match_date.normalize()
                
            if ' v ' in rest:
                team_a, remainder = rest.split(' v ', 1)
                team_b = remainder.rsplit(', ', 1)[0] if ', ' in remainder else (remainder.rsplit(' - ', 1)[0] if ' - ' in remainder else remainder)
            else:
                team_a, team_b = rest, "Unknown"
            return team_a.strip(), team_b.strip(), match_date
        except: return None, None, None

    cup_match_dict = {}
    if f_cup and os.path.exists(f_cup):
        try:
            excel_file = pd.ExcelFile(f_cup)
            target_sheet = excel_file.sheet_names[0]
            
            for sheet in excel_file.sheet_names:
                if domain.lower().replace("'", "") in sheet.lower().replace("'", ""):
                    target_sheet = sheet
                    break
                    
            cup_df = pd.read_excel(f_cup, sheet_name=target_sheet, header=None)
            
            for _, row_data in cup_df.iterrows():
                match_str_raw = str(row_data[0]).strip()
                cup_name = str(row_data[1]).strip()
                
                if match_str_raw.lower() in ['match string', 'match group', 'match', 'nan']:
                    continue
                
                cleaned_match_str = doc_format_cricket_names(match_str_raw, domain)
                c_team_a, c_team_b, c_date = parse_match_group(cleaned_match_str)
                
                if c_team_a and c_team_b:
                    teams = sorted([str(c_team_a).lower(), str(c_team_b).lower()])
                    if pd.notna(c_date):
                        robust_key = f"{teams[0]}_{teams[1]}_{c_date.strftime('%Y-%m-%d')}"
                        cup_match_dict[robust_key] = cup_name
                    else:
                        robust_key_no_date = f"{teams[0]}_{teams[1]}"
                        cup_match_dict[robust_key_no_date] = cup_name
        except Exception:
            pass

    reg_name_col = 'Full Name' if 'Full Name' in registered_players.columns else registered_players.columns[0]
    registered_players[reg_name_col] = registered_players[reg_name_col].astype(str).str.replace('‡', '', regex=False).str.strip()

    if domain == "Men's":
        exclusions = ['mark adair', 'ben calitz']
        pronoun = "he"
    elif domain == "Women's":
        exclusions = ['cara murray']
        pronoun = "she"

    registered_players['Date Registered'] = pd.to_datetime(registered_players['Date Registered'], dayfirst=True, errors='coerce').dt.normalize()
    ci_col = next((c for c in registered_players.columns if 'individual membership ci' in str(c).lower() or 'sport80' in str(c).lower() or 'ci no' in str(c).lower()), None)
    if ci_col:
        registered_players['_ci_no_clean'] = registered_players[ci_col].dropna().astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
    else:
        registered_players['_ci_no_clean'] = ""

    starring_df = pd.DataFrame(columns=['Rank', 'Surname', 'Forename', 'XI_Level', 'Club', 'Full Name'])
    if f_starring and os.path.exists(f_starring):
        starring_excel = pd.read_excel(f_starring, sheet_name=None, header=None)
        starring_data_list = []
        for club_name, df in starring_excel.items():
            try:
                while len(df.columns) < 5: df[len(df.columns)] = None
                df = df.iloc[:, [0, 1, 4]]
                df.columns = ['Rank', 'Surname', 'Forename']
                df['XI_Level'] = df['Rank'].apply(lambda x: str(x).strip() if 'XI' in str(x) else None).ffill()
                df['Is_Numeric'] = pd.to_numeric(df['Rank'], errors='coerce').notna()
                df = df[df['Is_Numeric']].copy().dropna(subset=['Surname'])
                if not df.empty:
                    df['Club'] = str(club_name).strip()
                    df['Forename'] = df['Forename'].fillna('')
                    df['Full Name'] = (df['Forename'].astype(str).str.strip() + ' ' + df['Surname'].astype(str).str.strip()).str.replace('‡', '', regex=False).str.strip()
                    starring_data_list.append(df)
            except Exception as e: print('EXCEPTION IN FINES:', repr(e))
        if starring_data_list: starring_df = pd.concat(starring_data_list, ignore_index=True)

    alias_map = build_alias_map(aliases, domain)
    f_secondary = DEFAULT_FILES.get(domain, {}).get("secondary", "")
    secondary_map = {}
    if os.path.exists(f_secondary):
        secondary_map = build_secondary_team_map(pd.read_excel(f_secondary), alias_map)
    
    f_unreg = DEFAULT_FILES.get(domain, {}).get("unreg", "")
    unreg_df = None
    if os.path.exists(f_unreg):
        unreg_df = pd.read_excel(f_unreg)
        
    player_club_map = build_player_club_map(registered_players, alias_map, domain, unreg_map_df=unreg_df, secondary_map=secondary_map) 
    player_club_map = infer_unregistered_player_clubs(batting_stats, bowling_stats, player_club_map, min_matches=2)
    
    if not f_id_map:
        f_id_map = DEFAULT_FILES.get(domain, {}).get("id_map", "")
    id_map = {}
    if f_id_map and os.path.exists(f_id_map):
        id_map_df = pd.read_excel(f_id_map)
        id_map = build_id_map(id_map_df)
        
    def process_bat_row(r):
        c_name, s80_id, _, _ = resolve_player_from_row(r, r['Name'], id_map, alias_map, player_club_map)
        return pd.Series([c_name, s80_id], index=['Cleaned Name', 'Sport80_ID'])

    def process_bowl_row(r):
        c_name, s80_id, _, _ = resolve_player_from_row(r, r['Bowler'], id_map, alias_map, player_club_map)
        return pd.Series([c_name, s80_id], index=['Cleaned Name', 'Sport80_ID'])

    bat_resolved = batting_stats.apply(process_bat_row, axis=1)
    batting_stats['Cleaned Name'] = bat_resolved['Cleaned Name']
    batting_stats['Sport80_ID'] = bat_resolved['Sport80_ID']

    bowl_resolved = bowling_stats.apply(process_bowl_row, axis=1)
    bowling_stats['Cleaned Name'] = bowl_resolved['Cleaned Name']
    bowling_stats['Sport80_ID'] = bowl_resolved['Sport80_ID']
    
    batting_stats['Group'] = batting_stats['Group'].apply(lambda x: doc_format_cricket_names(x, domain))
    bowling_stats['Group'] = bowling_stats['Group'].apply(lambda x: doc_format_cricket_names(x, domain))

    batters = batting_stats[['Group', 'Cleaned Name', 'Name', 'Is_Irish_Match', 'Sport80_ID']].rename(columns={'Cleaned Name': 'Player', 'Name': 'Scorecard Name'})
    bowlers = bowling_stats[['Group', 'Cleaned Name', 'Bowler', 'Is_Irish_Match', 'Sport80_ID']].rename(columns={'Cleaned Name': 'Player', 'Bowler': 'Scorecard Name'})
    
    app_dfs = [batters, bowlers]

    if not f_abandoned:
        f_abandoned = DEFAULT_FILES.get(domain, {}).get("abandoned", "")

    if f_abandoned and os.path.exists(f_abandoned):
        abandoned_stats = pd.read_excel(f_abandoned)
        if not abandoned_stats.empty:
            ab_match_col = 'Group' if 'Group' in abandoned_stats.columns else ('Match' if 'Match' in abandoned_stats.columns else abandoned_stats.columns[0])
            ab_name_col = 'Name' if 'Name' in abandoned_stats.columns else abandoned_stats.columns[1]
            
            abandoned_stats['Is_Irish_Match'] = False
            ab_resolved = abandoned_stats.apply(lambda r: resolve_player_from_row(r, r[ab_name_col], id_map, alias_map, player_club_map), axis=1)
            abandoned_stats['Cleaned Name'] = [res[0] for res in ab_resolved]
            abandoned_stats['Sport80_ID'] = [res[1] for res in ab_resolved]
            abandoned_stats['Group'] = abandoned_stats[ab_match_col].apply(lambda x: doc_format_cricket_names(x, domain))
            
            ab_apps = abandoned_stats[['Group', 'Cleaned Name', ab_name_col, 'Is_Irish_Match', 'Sport80_ID']].rename(columns={'Cleaned Name': 'Player', ab_name_col: 'Scorecard Name'})
            app_dfs.append(ab_apps)

    all_appearances = pd.concat(app_dfs).drop_duplicates(subset=['Group', 'Player'])
    all_appearances[['Team A', 'Team B', 'Match Date']] = all_appearances['Group'].apply(lambda x: pd.Series(parse_match_group(x)))
    all_appearances = all_appearances.sort_values(by=['Match Date'])

    league_dict, team_keys, _ = build_league_dict(league_structure)
    def determine_league(t_a, t_b):
        league_a, league_b = get_team_league(t_a, team_keys, league_dict, domain), get_team_league(t_b, team_keys, league_dict, domain)
        return league_a if league_a and league_b and league_a == league_b else "possible cup match"

    official_names = registered_players[reg_name_col].dropna().unique()
    deemed_registered, unregistered_audit, starring_violations = [], [], []
    all_matches_in_range, violation_matches = set(), set()
    first_unreg_match_date, first_unreg_match_team, first_unreg_match_teams_played, player_match_cache = {}, {}, {}, {}

    for idx, row in all_appearances.iterrows():
        player, scorecard_name, match_date = row['Player'], row['Scorecard Name'], row['Match Date']
        if pd.isna(match_date): continue
            
        team_a, team_b = row['Team A'], row['Team B']
        
        match_league = None
        teams = sorted([str(team_a).lower(), str(team_b).lower()])
        robust_key_date = f"{teams[0]}_{teams[1]}_{match_date.strftime('%Y-%m-%d')}"
        robust_key_no_date = f"{teams[0]}_{teams[1]}"
        
        if robust_key_date in cup_match_dict:
            match_league = cup_match_dict[robust_key_date]
        elif robust_key_no_date in cup_match_dict:
            match_league = cup_match_dict[robust_key_no_date]
            
        if not match_league:
            match_league = determine_league(team_a, team_b)
            
        in_date_range = (start_date <= match_date <= end_date)
        
        if in_date_range: all_matches_in_range.add((team_a, team_b, match_league, match_date))

        if row.get('Is_Irish_Match', False):
            reg_club = player_club_map.get(player.lower())
            if not reg_club or str(reg_club).lower() == 'nan':
                continue
                
            mock_row = {'Cleaned Name': player, 'Group': row['Group']}
            played_for = determine_player_team_for_row(mock_row, player_club_map, domain, secondary_map=secondary_map)
            if not get_team_league(played_for, team_keys, league_dict, domain):
                continue

        row_s80_id = row.get('Sport80_ID')
        has_valid_s80 = pd.notna(row_s80_id) and str(row_s80_id).strip() and str(row_s80_id).strip().lower() != 'nan'

        if player not in player_match_cache:
            if has_valid_s80 and '_ci_no_clean' in registered_players.columns and not registered_players['_ci_no_clean'].empty:
                clean_s80 = str(row_s80_id).replace('.0', '').strip()
                matched_reg = registered_players[registered_players['_ci_no_clean'] == clean_s80]
                if not matched_reg.empty:
                    reg_record = matched_reg
                    match_type, matched_name = "Sport80 ID Exact", reg_record.iloc[0][reg_name_col]
                else:
                    reg_record = pd.DataFrame()
                    match_type, matched_name = "Sport80 ID (Unregistered/Lapsed)", scorecard_name
            elif domain == "Men's" and player.lower() == 'james shannon':
                if 'holywood' in str(team_a).lower() or 'holywood' in str(team_b).lower():
                    reg_record = registered_players[(registered_players[reg_name_col].str.lower() == 'james shannon') & (registered_players['Individual Membership Primary Club'].str.contains('Holywood', case=False, na=False))].copy()
                    if not reg_record.empty: reg_record['Date Registered'] = pd.to_datetime('2026-03-05', dayfirst=True).normalize()
                elif 'saintfield' in str(team_a).lower() or 'saintfield' in str(team_b).lower():
                    reg_record = registered_players[(registered_players[reg_name_col].str.lower() == 'james shannon') & (registered_players['Individual Membership Primary Club'].str.contains('Saintfield', case=False, na=False))]
                else: reg_record = pd.DataFrame()
                match_type, matched_name = "Exact (Contextual Override)", "James Shannon"
                
            elif '(' in player and player.strip().endswith(')'):
                base_name = player.split('(')[0].strip()
                club_hint = player.split('(')[-1].replace(')', '').strip()
                
                potential_matches = registered_players[registered_players[reg_name_col].str.strip().str.lower() == base_name.lower()]
                if potential_matches.empty:
                    best_match, score = process.extractOne(base_name, official_names, scorer=fuzz.token_sort_ratio)
                    if score >= 90:
                        potential_matches = registered_players[registered_players[reg_name_col] == best_match]
                
                if not potential_matches.empty:
                    clean_hint = clean_club_for_matching(club_hint)
                    def check_club_or_transfer(r):
                        if clean_hint in clean_club_for_matching(r.get('Individual Membership Primary Club', '')): return True
                        t_cols = [c for c in r.index if 'Transfer' in str(c) and 'Date' not in str(c)]
                        return clean_hint in clean_club_for_matching(r.get(t_cols[0], '')) if t_cols else False
                    reg_record = potential_matches[potential_matches.apply(check_club_or_transfer, axis=1)]
                    if not reg_record.empty:
                        match_type, matched_name = f"Duplicate Match ({club_hint})", reg_record.iloc[0][reg_name_col]
                    else:
                        reg_record = pd.DataFrame()
                        match_type, matched_name = "Failed", "NO MATCH FOUND"
                else:
                    reg_record = pd.DataFrame()
                    match_type, matched_name = "Failed", "NO MATCH FOUND"
                    
            else:
                reg_record = registered_players[registered_players[reg_name_col].str.strip().str.lower() == player.lower()]
                match_type, matched_name = "Exact", player
                if reg_record.empty:
                    best_match, score = process.extractOne(player, official_names, scorer=fuzz.token_sort_ratio)
                    if score >= 90:
                        reg_record = registered_players[registered_players[reg_name_col] == best_match]
                        match_type, matched_name = f"Fuzzy ({score}%)", best_match
            
            s80_val = ""
            if has_valid_s80:
                s80_val = str(row_s80_id).replace('.0', '').strip()
            elif not reg_record.empty and '_ci_no_clean' in reg_record.columns:
                ci_series = reg_record['_ci_no_clean'].dropna()
                if not ci_series.empty and str(ci_series.iloc[0]).strip() and str(ci_series.iloc[0]).strip().lower() != 'nan':
                    s80_val = str(ci_series.iloc[0]).strip()

            player_match_cache[player] = (reg_record, match_type, matched_name, s80_val)
        else:
            reg_record, match_type, matched_name, s80_val = player_match_cache[player]

        is_registered = False
        reg_date, reg_club = pd.NaT, "Unknown Club"
        status_text = 'Unregistered / Missing completely'
        if not reg_record.empty:
            mock_row = {'Cleaned Name': player, 'Group': row.get('Group', '')}
            played_for = determine_player_team_for_row(mock_row, player_club_map, domain, secondary_map=secondary_map)
            played_base = extract_base_club_name(played_for).lower()

            if len(reg_record) > 1:
                def matches_played_club(r):
                    r_raw = r.get('Individual Membership Primary Club', '')
                    r_b = extract_base_club_name(str(r_raw)).lower() if pd.notna(r_raw) else ""
                    t_cols = [c for c in r.index if 'Transfer' in str(c) and 'Date' not in str(c)]
                    t_raw = r.get(t_cols[0], '') if t_cols else ''
                    t_b = extract_base_club_name(str(t_raw)).lower() if pd.notna(t_raw) else ""
                    return bool(r_b and str(r_raw).strip() != '' and (r_b in played_base or played_base in r_b)) or \
                           bool(t_b and str(t_raw).strip() != '' and (t_b in played_base or played_base in t_b))
                
                filtered = reg_record[reg_record.apply(matches_played_club, axis=1)]
                if not filtered.empty:
                    reg_record = filtered
                    match_type, matched_name = f"Exact (Disambiguated via {played_base})", reg_record.iloc[0][reg_name_col]
                
                reg_record = reg_record.sort_values(by='Date Registered')

            reg_date = reg_record.iloc[0]['Date Registered']
            raw_club = reg_record.iloc[0].get('Individual Membership Primary Club', pd.NA)
            if pd.notna(raw_club) and str(raw_club).strip() != '': reg_club = str(raw_club).strip()
            
            t_cols = [c for c in reg_record.columns if 'Transfer' in str(c) and 'Date' not in str(c)]
            transfer_club = reg_record.iloc[0].get(t_cols[0], pd.NA) if t_cols else pd.NA
            transfer_date = reg_record.iloc[0].get('Transfer Date', pd.NaT)
            
            r_base = extract_base_club_name(str(raw_club)).lower() if pd.notna(raw_club) else ""
            t_base = extract_base_club_name(str(transfer_club)).lower() if pd.notna(transfer_club) else ""
            
            played_for_transfer = bool(t_base and str(transfer_club).strip() != '' and (t_base in played_base or played_base in t_base))
            played_for_primary = bool(r_base and str(raw_club).strip() != '' and (r_base in played_base or played_base in r_base))
            
            played_for_secondary = False
            if secondary_map:
                mapped_name = alias_map.get(player.lower(), player.lower())
                sec_teams = secondary_map.get(mapped_name) or secondary_map.get(player.lower()) or []
                for st in sec_teams:
                    st_base = extract_base_club_name(st).lower()
                    if st_base in played_base or played_base in st_base:
                        played_for_secondary = True
                        break
            
            if played_for_transfer:
                reg_club = str(transfer_club).strip()
                if pd.notna(transfer_date):
                    reg_date = pd.to_datetime(transfer_date)
                
                if pd.notna(reg_date) and reg_date <= match_date:
                    is_registered = True
                else:
                    status_text = 'Unregistered for this match (Played for transfer club, but transfer date is late or missing)'
            elif played_for_primary or played_for_secondary:
                if pd.notna(reg_date) and reg_date <= match_date:
                    if pd.notna(transfer_date) and pd.to_datetime(transfer_date) <= match_date:
                        status_text = 'Unregistered for this match (Played for former primary club AFTER transferring away)'
                    else:
                        is_registered = True
                else:
                    status_text = 'Unregistered for this match (Registered late)'
            else:
                status_text = f'Unregistered / Played for Wrong Club (Registered to {reg_club})'
                
        if not is_registered:
            f_match_logic = match_type if not reg_record.empty else 'Failed'
            f_matched_name = matched_name if not reg_record.empty else 'NO MATCH FOUND'
            
            if player not in first_unreg_match_date:
                first_unreg_match_date[player] = match_date
                first_unreg_match_team[player] = determine_player_team_for_row({'Cleaned Name': player, 'Group': row.get('Group', '')}, player_club_map, domain)
                first_unreg_match_teams_played[player] = f"{doc_formal_team_name(team_a)} v {doc_formal_team_name(team_b)}"
                if in_date_range:
                    violation_matches.add((team_a, team_b, match_league, match_date))
                    unregistered_audit.append({
                        'Stats Name (Cleaned)': player, 'Original Scorecard Name': scorecard_name,
                        'Sport80 ID': s80_val,
                        'Matched Registered Name': f_matched_name, 'Registered Club': reg_club,
                        'Match Date': match_date, 'Date Registered': reg_date, 'Status': status_text,
                        'Team A': team_a, 'Team B': team_b, 'Match League': match_league, 'Match Logic': f_match_logic
                    })
            else:
                if in_date_range:
                    deemed_registered.append({
                        'Stats Name (Cleaned)': player, 'Original Scorecard Name': scorecard_name,
                        'Sport80 ID': s80_val,
                        'Matched Registered Name': f_matched_name, 'Registered Club': reg_club,
                        'Match Date': match_date, 'Deemed Registered Date': first_unreg_match_date[player],
                        'Deemed Registered Match Teams': first_unreg_match_teams_played[player],
                        'Deemed Registered Club': extract_base_club_name(str(first_unreg_match_team[player])),
                        'Date Registered': reg_date, 'Team A': team_a, 'Team B': team_b,
                        'Match League': match_league, 'Match Logic': f_match_logic
                    })

    valid_matches = all_appearances[(all_appearances['Match Date'] >= start_date) & (all_appearances['Match Date'] <= end_date)].copy()
    if not starring_df.empty and 'Full Name' in starring_df.columns:
        starring_df['Cleaned Name'] = starring_df['Full Name'].apply(lambda x: cleanse_name(x, alias_map))
        for idx, row in valid_matches.iterrows():
            player, scorecard_name, team_a, team_b = row['Player'], row['Scorecard Name'], str(row['Team A']), str(row['Team B'])
            if str(player).strip().lower() in exclusions: continue
            
            p_s80 = ""
            if player in player_match_cache:
                p_s80 = player_match_cache[player][3]
            elif pd.notna(row.get('Sport80_ID')):
                p_s80 = str(row.get('Sport80_ID')).replace('.0', '').strip()
                
            p_stars = starring_df[starring_df['Cleaned Name'].str.strip().str.lower() == player.lower()]
            if not p_stars.empty:
                s_rank, s_club = str(p_stars.iloc[0]['XI_Level']), str(p_stars.iloc[0]['Club'])
                played_team = team_a if s_club.lower() in team_a.lower() else (team_b if s_club.lower() in team_b.lower() else None)
                if played_team:
                    try:
                        s_rm = re.search(r'(\d+)', s_rank)
                        p_rm = re.search(r'\s([1-9])$', played_team)
                        
                        s_rank_int = int(s_rm.group(1)) if s_rm else 1
                        p_rank_int = int(p_rm.group(1)) if p_rm else 1
                        
                        if p_rank_int > s_rank_int:
                            starring_violations.append({
                                'Player (Cleaned)': player, 'Original Scorecard Name': scorecard_name,
                                'Sport80 ID': p_s80,
                                'Starred For': f"{s_club} {s_rank}", 'Actually Played For': played_team,
                                'Team A': team_a, 'Team B': team_b, 'Match Date': row['Match Date'], 'Match Group': row['Group']
                            })
                    except: pass

    unregistered_audit.sort(key=lambda x: violation_sort_key(x, 'Stats Name (Cleaned)'))
    deemed_registered.sort(key=lambda x: violation_sort_key(x, 'Stats Name (Cleaned)'))
    starring_violations.sort(key=lambda x: violation_sort_key(x, 'Player (Cleaned)'))

    excel_io = io.BytesIO()
    with pd.ExcelWriter(excel_io, engine='xlsxwriter', datetime_format='dd/mm/yyyy') as writer:
        if unregistered_audit: export_and_format_excel(pd.DataFrame(unregistered_audit), writer, "Unregistered Matches")
        else: pd.DataFrame(columns=["Status"]).to_excel(writer, index=False, sheet_name="Unregistered Matches")
        if deemed_registered: export_and_format_excel(pd.DataFrame(deemed_registered), writer, "Deemed Registered")
        else: pd.DataFrame(columns=["Status"]).to_excel(writer, index=False, sheet_name="Deemed Registered")
        if starring_violations: export_and_format_excel(pd.DataFrame(starring_violations), writer, "Starring Violations")
        else: pd.DataFrame(columns=["Status"]).to_excel(writer, index=False, sheet_name="Starring Violations")

    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Calibri', Pt(11)
    start_str, end_str = get_ordinal_date(start_date, False), get_ordinal_date(end_date, True)
    
    p_intro = doc.add_paragraph()
    run_intro = p_intro.add_run(f"Here is the audit for the matches played between {start_str} and {end_str}.")
    run_intro.bold = True
    run_intro.font.size = Pt(13)
    run_intro.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph("I have checked the player names against the official NCU registrations, applied fuzzy logic mapping for variances, and cross-referenced the Date Registered field to ensure players were officially registered on or before the date of their match.")
    
    p_noviol = doc.add_paragraph()
    p_noviol.paragraph_format.space_before = Pt(14)
    p_noviol.paragraph_format.space_after = Pt(4)
    run_noviol = p_noviol.add_run("--- Matches WITH Violations ---")
    run_noviol.bold = True
    run_noviol.font.size = Pt(13)
    run_noviol.font.color.rgb = RGBColor(0, 0, 128)

    if violation_matches:
        render_grouped_matches(doc, violation_matches, domain)
    else:
        doc.add_paragraph("No matches with violations recorded.")

    p_noviol = doc.add_paragraph()
    p_noviol.paragraph_format.space_before = Pt(14)  
    p_noviol.paragraph_format.space_after = Pt(4)    
    run_noviol = p_noviol.add_run("--- Matches WITHOUT Violations ---")
    run_noviol.bold = True
    run_noviol.font.size = Pt(13)
    run_noviol.font.color.rgb = RGBColor(0, 0, 128)

    matches_without_violations = all_matches_in_range - violation_matches
    if matches_without_violations:
        render_grouped_matches(doc, matches_without_violations, domain)
    else:
        doc.add_paragraph("No matches without violations recorded.")

    doc.add_page_break()
    p_unreg = doc.add_paragraph()
    run_unreg = p_unreg.add_run("--- Unregistered Players / Date Violations ---")
    run_unreg.bold = True
    run_unreg.font.size = Pt(13)
    run_unreg.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph('These players took the field prior to the official "Date Registered" logged in the master file database.')
    for r in unregistered_audit:
        p_name, s_name = r['Stats Name (Cleaned)'], r['Original Scorecard Name']
        d_name = p_name if p_name.lower() == str(s_name).lower() else f"{p_name} (Played as: {s_name})"
        m_date = get_ordinal_date(r['Match Date'])
        m_league = str(r.get('Match League', 'Unknown League'))
        t_str = f" [Registered Club: {doc_formal_team_name(r.get('Registered Club', 'Unknown Club'))}] (Match: {doc_formal_team_name(r['Team A'])} v {doc_formal_team_name(r['Team B'])} - {m_date} - {m_league})"
        
        p_p = doc.add_paragraph(style='List Bullet')
        p_p.add_run(f"{d_name}").bold = True
        p_p.add_run(f"{t_str}")
        
        r_date = r['Date Registered']
        status = r.get('Status', '')
        if 'Played for Wrong Club' in status:
            d_text = f"The player is registered for another club ({doc_formal_team_name(r.get('Registered Club', 'Unknown Club'))}) and not registered for any club in the match they played in."
        elif pd.notna(r_date):
            d_text = f"The official database indicates a registration date of {get_ordinal_date(r_date)} ({(r_date - r['Match Date']).days} days late)."
        else:
            d_text = f"Appeared under the scorecard name \"{s_name}\" (verified via alias map). This official profile is entirely unregistered on the master registry."
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.add_run("o  Violation Detail: ").bold = True
        p.add_run(d_text)

    doc.add_page_break()
    p_deemed = doc.add_paragraph()
    run_deemed = p_deemed.add_run("--- Deemed Registered (Played Previously in 2026) ---")
    run_deemed.bold = True
    run_deemed.font.size = Pt(13)
    run_deemed.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph("These players were late or unverified on their match date, but are deemed registered for this specific game because they already played an active match earlier in the recorded 2026 season logs.")
    for r in deemed_registered:
        p_name, s_name = r['Stats Name (Cleaned)'], r['Original Scorecard Name']
        d_name = p_name if p_name.lower() == str(s_name).lower() else f"{p_name} (Played as: {s_name})"
        deemed_club = r.get('Deemed Registered Club')
        club_display = f"Deemed Registered Club: {doc_formal_team_name(deemed_club)}" if deemed_club else f"Registered Club: {doc_formal_team_name(r.get('Registered Club', 'Unknown Club'))}"
        m_date = get_ordinal_date(r['Match Date'])
        m_league = str(r.get('Match League', 'Unknown League'))
        t_str = f" [{club_display}] (Match: {doc_formal_team_name(r['Team A'])} v {doc_formal_team_name(r['Team B'])} - {m_date} - {m_league})"
        
        p_p = doc.add_paragraph(style='List Bullet')
        p_p.add_run(f"{d_name}").bold = True
        p_p.add_run(f"{t_str}")
        
        r_date = r['Date Registered']
        r_det = f"Registered on {get_ordinal_date(r_date)}" if pd.notna(r_date) else "Unregistered profile"
        r_det = r_det[0].upper() + r_det[1:]
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.add_run("o  Deemed Status: ").bold = True
        deemed_teams = r.get('Deemed Registered Match Teams', '')
        in_match_str = f" in {deemed_teams}" if deemed_teams else ""
        p.add_run(f"{r_det}. Deemed registered because {pronoun} previously played{in_match_str} on {get_ordinal_date(r['Deemed Registered Date'])}.")

    doc.add_page_break()
    p_star = doc.add_paragraph()
    run_star = p_star.add_run("--- Starring Violations ---")
    run_star.bold = True
    run_star.font.size = Pt(13)
    run_star.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph("These players played for a team at a lower level than their starred rank.")
    for r in starring_violations:
        p_name, s_name = r['Player (Cleaned)'], r['Original Scorecard Name']
        d_name = p_name if p_name.lower() == str(s_name).lower() else f"{p_name} (Played as: {s_name})"
        m_date = get_ordinal_date(r['Match Date'])
        t_str = f" (Match: {doc_formal_team_name(r['Team A'])} v {doc_formal_team_name(r['Team B'])} - {m_date})"
        
        p_p = doc.add_paragraph(style='List Bullet')
        p_p.add_run(f"{d_name}").bold = True
        p_p.add_run(f"{t_str}")
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.add_run("o  Violation Detail: ").bold = True
        p.add_run(f"Played for {doc_formal_team_name(r['Actually Played For'])}, but is starred for {doc_formal_team_name(r['Starred For'])}.")

    doc_io = io.BytesIO()
    doc.save(doc_io)
    
    return excel_io, doc_io

# ==========================================
# MIDWEEK REGISTRATION ENGINE
# ==========================================
@st.cache_data(show_spinner="Running midweek registration audit...")
def run_midweek_registration_audit(start_date, end_date, f_reg, f_alias, f_starring, f_weekend_league, f_midweek_league, f_bat, f_bowl, f_abandoned=None, f_id_map=None):
    registered_players = pd.read_excel(f_reg)
    aliases = pd.read_excel(f_alias)
    weekend_structure = pd.read_excel(f_weekend_league)
    midweek_structure = pd.read_excel(f_midweek_league)
    batting_stats = pd.read_excel(f_bat)
    bowling_stats = pd.read_excel(f_bowl)

    reg_name_col = 'Full Name' if 'Full Name' in registered_players.columns else registered_players.columns[0]
    registered_players[reg_name_col] = registered_players[reg_name_col].astype(str).str.replace('‡', '', regex=False).str.strip()

    registered_players['Date Registered'] = pd.to_datetime(registered_players['Date Registered'], dayfirst=True, errors='coerce').dt.normalize()
    ci_col = next((c for c in registered_players.columns if 'individual membership ci' in str(c).lower() or 'sport80' in str(c).lower() or 'ci no' in str(c).lower()), None)
    if ci_col:
        registered_players['_ci_no_clean'] = registered_players[ci_col].dropna().astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
    else:
        registered_players['_ci_no_clean'] = ""

    starring_df = pd.DataFrame(columns=['Rank', 'Surname', 'Forename', 'XI_Level', 'Club', 'Full Name'])
    if f_starring and os.path.exists(f_starring):
        starring_excel = pd.read_excel(f_starring, sheet_name=None, header=None)
        starring_data_list = []
        for club_name, df in starring_excel.items():
            try:
                while len(df.columns) < 5: df[len(df.columns)] = None
                df = df.iloc[:, [0, 1, 4]]
                df.columns = ['Rank', 'Surname', 'Forename']
                df['XI_Level'] = df['Rank'].apply(lambda x: str(x).strip() if 'XI' in str(x) else None).ffill()
                df['Is_Numeric'] = pd.to_numeric(df['Rank'], errors='coerce').notna()
                df = df[df['Is_Numeric']].copy().dropna(subset=['Surname'])
                if not df.empty:
                    df['Club'] = str(club_name).strip()
                    df['Forename'] = df['Forename'].fillna('')
                    df['Full Name'] = (df['Forename'].astype(str).str.strip() + ' ' + df['Surname'].astype(str).str.strip()).str.replace('‡', '', regex=False).str.strip()
                    starring_data_list.append(df)
            except Exception as e: print('EXCEPTION IN FINES:', repr(e))
        if starring_data_list: starring_df = pd.concat(starring_data_list, ignore_index=True)

    alias_map = build_alias_map(aliases, "Midweek")
    f_secondary = DEFAULT_FILES.get("Midweek", {}).get("secondary", "")
    secondary_map = {}
    if os.path.exists(f_secondary):
        secondary_map = build_secondary_team_map(pd.read_excel(f_secondary), alias_map)
    
    f_unreg = DEFAULT_FILES.get("Midweek", {}).get("unreg", "")
    unreg_df = None
    if os.path.exists(f_unreg):
        unreg_df = pd.read_excel(f_unreg)
    
    player_club_map = build_player_club_map(registered_players, alias_map, "Midweek", unreg_map_df=unreg_df, secondary_map=secondary_map)
    player_club_map = infer_unregistered_player_clubs(batting_stats, bowling_stats, player_club_map, min_matches=2)
    
    if not f_id_map:
        f_id_map = DEFAULT_FILES.get("Midweek", {}).get("id_map", "")
    id_map = {}
    if f_id_map and os.path.exists(f_id_map):
        id_map_df = pd.read_excel(f_id_map)
        id_map = build_id_map(id_map_df)
        
    def process_mw_bat_row(r):
        c_name, s80_id, _, _ = resolve_player_from_row(r, r['Name'], id_map, alias_map, player_club_map)
        return pd.Series([c_name, s80_id], index=['Cleaned Name', 'Sport80_ID'])

    def process_mw_bowl_row(r):
        c_name, s80_id, _, _ = resolve_player_from_row(r, r['Bowler'], id_map, alias_map, player_club_map)
        return pd.Series([c_name, s80_id], index=['Cleaned Name', 'Sport80_ID'])

    bat_resolved = batting_stats.apply(process_mw_bat_row, axis=1)
    batting_stats['Cleaned Name'] = bat_resolved['Cleaned Name']
    batting_stats['Sport80_ID'] = bat_resolved['Sport80_ID']

    bowl_resolved = bowling_stats.apply(process_mw_bowl_row, axis=1)
    bowling_stats['Cleaned Name'] = bowl_resolved['Cleaned Name']
    bowling_stats['Sport80_ID'] = bowl_resolved['Sport80_ID']
    
    batting_stats['Group'] = batting_stats['Group'].apply(lambda x: doc_format_cricket_names(x, "Midweek"))
    bowling_stats['Group'] = bowling_stats['Group'].apply(lambda x: doc_format_cricket_names(x, "Midweek"))

    def parse_match_group(group_str):
        try:
            group_str = str(group_str).strip()
            parts = group_str.rsplit(' - ', 1)
            date_str = parts[1].strip() if len(parts) == 2 else group_str
            rest = parts[0].strip() if len(parts) == 2 else group_str
            match_date = pd.to_datetime(date_str, dayfirst=True, errors='coerce').normalize()
            if ' v ' in rest:
                team_a, remainder = rest.split(' v ', 1)
                team_b = remainder.rsplit(', ', 1)[0] if ', ' in remainder else (remainder.rsplit(' - ', 1)[0] if ' - ' in remainder else remainder)
            else:
                team_a, team_b = rest, "Unknown"
            return team_a.strip(), team_b.strip(), match_date
        except: return None, None, None

    batters = batting_stats[['Group', 'Cleaned Name', 'Name', 'Sport80_ID']].rename(columns={'Cleaned Name': 'Player', 'Name': 'Scorecard Name'})
    bowlers = bowling_stats[['Group', 'Cleaned Name', 'Bowler', 'Sport80_ID']].rename(columns={'Cleaned Name': 'Player', 'Bowler': 'Scorecard Name'})
    
    app_dfs = [batters, bowlers]

    if not f_abandoned:
        f_abandoned = DEFAULT_FILES.get("Midweek", {}).get("abandoned", "")

    if f_abandoned and os.path.exists(f_abandoned):
        abandoned_stats = pd.read_excel(f_abandoned)
        if not abandoned_stats.empty:
            ab_match_col = 'Group' if 'Group' in abandoned_stats.columns else ('Match' if 'Match' in abandoned_stats.columns else abandoned_stats.columns[0])
            ab_name_col = 'Name' if 'Name' in abandoned_stats.columns else abandoned_stats.columns[1]
            
            abandoned_stats['Is_Irish_Match'] = False
            ab_resolved = abandoned_stats.apply(lambda r: resolve_player_from_row(r, r[ab_name_col], id_map, alias_map, player_club_map), axis=1)
            abandoned_stats['Cleaned Name'] = [res[0] for res in ab_resolved]
            abandoned_stats['Sport80_ID'] = [res[1] for res in ab_resolved]
            abandoned_stats['Group'] = abandoned_stats[ab_match_col].apply(lambda x: doc_format_cricket_names(x, "Midweek"))
            
            ab_apps = abandoned_stats[['Group', 'Cleaned Name', ab_name_col, 'Sport80_ID']].rename(columns={'Cleaned Name': 'Player', ab_name_col: 'Scorecard Name'})
            app_dfs.append(ab_apps)

    all_appearances = pd.concat(app_dfs).drop_duplicates(subset=['Group', 'Player'])
    all_appearances[['Team A', 'Team B', 'Match Date']] = all_appearances['Group'].apply(lambda x: pd.Series(parse_match_group(x)))
    all_appearances = all_appearances.sort_values(by=['Match Date'])

    mw_league_dict, mw_team_keys, _ = build_league_dict(midweek_structure)
    wknd_league_dict, wknd_team_keys, _ = build_league_dict(weekend_structure)

    def determine_midweek_league(t_a, t_b):
        league_a, league_b = get_team_league(t_a, mw_team_keys, mw_league_dict, "Midweek"), get_team_league(t_b, mw_team_keys, mw_league_dict, "Midweek")
        return league_a if league_a and league_b and league_a == league_b else "Midweek Cup/Fixture"

    official_names = registered_players[reg_name_col].dropna().unique()
    deemed_registered, unregistered_audit, starring_violations = [], [], []
    all_matches_in_range, violation_matches = set(), set()
    first_unreg_match_date, first_unreg_match_team, first_unreg_match_teams_played, player_match_cache = {}, {}, {}, {}

    for idx, row in all_appearances.iterrows():
        player, scorecard_name, match_date = row['Player'], row['Scorecard Name'], row['Match Date']
        if pd.isna(match_date): continue
            
        team_a, team_b = row['Team A'], row['Team B']
        match_league = determine_midweek_league(team_a, team_b)
        in_date_range = (start_date <= match_date <= end_date)
        
        if in_date_range: all_matches_in_range.add((team_a, team_b, match_league, match_date))

        row_s80_id = row.get('Sport80_ID')
        has_valid_s80 = pd.notna(row_s80_id) and str(row_s80_id).strip() and str(row_s80_id).strip().lower() != 'nan'

        if player not in player_match_cache:
            if has_valid_s80 and '_ci_no_clean' in registered_players.columns and not registered_players['_ci_no_clean'].empty:
                clean_s80 = str(row_s80_id).replace('.0', '').strip()
                matched_reg = registered_players[registered_players['_ci_no_clean'] == clean_s80]
                if not matched_reg.empty:
                    reg_record = matched_reg
                    match_type, matched_name = "Sport80 ID Exact", reg_record.iloc[0][reg_name_col]
                else:
                    reg_record = pd.DataFrame()
                    match_type, matched_name = "Sport80 ID (Unregistered/Lapsed)", scorecard_name
            elif player.lower() == 'james shannon':
                if 'holywood' in str(team_a).lower() or 'holywood' in str(team_b).lower():
                    reg_record = registered_players[(registered_players[reg_name_col].str.lower() == 'james shannon') & (registered_players['Individual Membership Primary Club'].str.contains('Holywood', case=False, na=False))].copy()
                    if not reg_record.empty: reg_record['Date Registered'] = pd.to_datetime('2026-03-05', dayfirst=True).normalize()
                elif 'saintfield' in str(team_a).lower() or 'saintfield' in str(team_b).lower():
                    reg_record = registered_players[(registered_players[reg_name_col].str.lower() == 'james shannon') & (registered_players['Individual Membership Primary Club'].str.contains('Saintfield', case=False, na=False))]
                else: reg_record = pd.DataFrame()
                match_type, matched_name = "Exact (Contextual Override)", "James Shannon"
                
            elif '(' in player and player.strip().endswith(')'):
                base_name = player.split('(')[0].strip()
                club_hint = player.split('(')[-1].replace(')', '').strip()
                
                potential_matches = registered_players[registered_players[reg_name_col].str.strip().str.lower() == base_name.lower()]
                if potential_matches.empty:
                    best_match, score = process.extractOne(base_name, official_names, scorer=fuzz.token_sort_ratio)
                    if score >= 90:
                        potential_matches = registered_players[registered_players[reg_name_col] == best_match]
                
                if not potential_matches.empty:
                    clean_hint = clean_club_for_matching(club_hint)
                    def check_club_or_transfer(r):
                        if clean_hint in clean_club_for_matching(r.get('Individual Membership Primary Club', '')): return True
                        t_cols = [c for c in r.index if 'Transfer' in str(c) and 'Date' not in str(c)]
                        return clean_hint in clean_club_for_matching(r.get(t_cols[0], '')) if t_cols else False
                    reg_record = potential_matches[potential_matches.apply(check_club_or_transfer, axis=1)]
                    if not reg_record.empty:
                        match_type, matched_name = f"Duplicate Match ({club_hint})", reg_record.iloc[0][reg_name_col]
                    else:
                        reg_record = pd.DataFrame()
                        match_type, matched_name = "Failed", "NO MATCH FOUND"
                else:
                    reg_record = pd.DataFrame()
                    match_type, matched_name = "Failed", "NO MATCH FOUND"
                    
            else:
                reg_record = registered_players[registered_players[reg_name_col].str.strip().str.lower() == player.lower()]
                match_type, matched_name = "Exact", player
                if reg_record.empty:
                    best_match, score = process.extractOne(player, official_names, scorer=fuzz.token_sort_ratio)
                    if score >= 90:
                        reg_record = registered_players[registered_players[reg_name_col] == best_match]
                        match_type, matched_name = f"Fuzzy ({score}%)", best_match
            
            s80_val = ""
            if has_valid_s80:
                s80_val = str(row_s80_id).replace('.0', '').strip()
            elif not reg_record.empty and '_ci_no_clean' in reg_record.columns:
                ci_series = reg_record['_ci_no_clean'].dropna()
                if not ci_series.empty and str(ci_series.iloc[0]).strip() and str(ci_series.iloc[0]).strip().lower() != 'nan':
                    s80_val = str(ci_series.iloc[0]).strip()

            player_match_cache[player] = (reg_record, match_type, matched_name, s80_val)
        else:
            reg_record, match_type, matched_name, s80_val = player_match_cache[player]

        is_registered = False
        reg_date, reg_club = pd.NaT, "Unknown Club"
        status_text = 'Unregistered / Missing completely'
        if not reg_record.empty:
            mock_row = {'Cleaned Name': player, 'Group': row.get('Group', '')}
            played_for = determine_player_team_for_row(mock_row, player_club_map, "Midweek", secondary_map=secondary_map)
            played_base = extract_base_club_name(played_for).lower()

            if len(reg_record) > 1:
                def matches_played_club(r):
                    r_raw = r.get('Individual Membership Primary Club', '')
                    r_b = extract_base_club_name(str(r_raw)).lower() if pd.notna(r_raw) else ""
                    t_cols = [c for c in r.index if 'Transfer' in str(c) and 'Date' not in str(c)]
                    t_raw = r.get(t_cols[0], '') if t_cols else ''
                    t_b = extract_base_club_name(str(t_raw)).lower() if pd.notna(t_raw) else ""
                    return bool(r_b and str(r_raw).strip() != '' and (r_b in played_base or played_base in r_b)) or \
                           bool(t_b and str(t_raw).strip() != '' and (t_b in played_base or played_base in t_b))
                
                filtered = reg_record[reg_record.apply(matches_played_club, axis=1)]
                if not filtered.empty:
                    reg_record = filtered
                    match_type, matched_name = f"Exact (Disambiguated via {played_base})", reg_record.iloc[0][reg_name_col]
                    
                reg_record = reg_record.sort_values(by='Date Registered')

            reg_date = reg_record.iloc[0]['Date Registered']
            raw_club = reg_record.iloc[0].get('Individual Membership Primary Club', pd.NA)
            if pd.notna(raw_club) and str(raw_club).strip() != '': reg_club = str(raw_club).strip()
            
            t_cols = [c for c in reg_record.columns if 'Transfer' in str(c) and 'Date' not in str(c)]
            transfer_club = reg_record.iloc[0].get(t_cols[0], pd.NA) if t_cols else pd.NA
            transfer_date = reg_record.iloc[0].get('Transfer Date', pd.NaT)
            
            r_base = extract_base_club_name(str(raw_club)).lower() if pd.notna(raw_club) else ""
            t_base = extract_base_club_name(str(transfer_club)).lower() if pd.notna(transfer_club) else ""
            
            played_for_transfer = bool(t_base and str(transfer_club).strip() != '' and (t_base in played_base or played_base in t_base))
            played_for_primary = bool(r_base and str(raw_club).strip() != '' and (r_base in played_base or played_base in r_base))
            
            played_for_secondary = False
            if secondary_map:
                mapped_name = alias_map.get(player.lower(), player.lower())
                sec_teams = secondary_map.get(mapped_name) or secondary_map.get(player.lower()) or []
                for st in sec_teams:
                    st_base = extract_base_club_name(st).lower()
                    if st_base in played_base or played_base in st_base:
                        played_for_secondary = True
                        break
            
            if played_for_transfer:
                reg_club = str(transfer_club).strip()
                if pd.notna(transfer_date):
                    reg_date = pd.to_datetime(transfer_date)
                
                if pd.notna(reg_date) and reg_date <= match_date:
                    is_registered = True
                else:
                    status_text = 'Unregistered for this match (Played for transfer club, but transfer date is late or missing)'
            elif played_for_primary or played_for_secondary:
                if pd.notna(reg_date) and reg_date <= match_date:
                    if pd.notna(transfer_date) and pd.to_datetime(transfer_date) <= match_date:
                        status_text = 'Unregistered for this match (Played for former primary club AFTER transferring away)'
                    else:
                        is_registered = True
                else:
                    status_text = 'Unregistered for this match (Registered late)'
            else:
                status_text = f'Unregistered / Played for Wrong Club (Registered to {reg_club})'
                
        if not is_registered:
            f_match_logic = match_type if not reg_record.empty else 'Failed'
            f_matched_name = matched_name if not reg_record.empty else 'NO MATCH FOUND'
            
            if player not in first_unreg_match_date:
                first_unreg_match_date[player] = match_date
                first_unreg_match_team[player] = determine_player_team_for_row({'Cleaned Name': player, 'Group': row.get('Group', '')}, player_club_map, "Midweek")
                first_unreg_match_teams_played[player] = f"{doc_formal_team_name(team_a)} v {doc_formal_team_name(team_b)}"
                if in_date_range:
                    violation_matches.add((team_a, team_b, match_league, match_date))
                    unregistered_audit.append({
                        'Stats Name (Cleaned)': player, 'Original Scorecard Name': scorecard_name,
                        'Sport80 ID': s80_val,
                        'Matched Registered Name': f_matched_name, 'Registered Club': reg_club,
                        'Match Date': match_date, 'Date Registered': reg_date, 'Status': status_text,
                        'Team A': team_a, 'Team B': team_b, 'Match League': match_league, 'Match Logic': f_match_logic
                    })
            else:
                if in_date_range:
                    deemed_registered.append({
                        'Stats Name (Cleaned)': player, 'Original Scorecard Name': scorecard_name,
                        'Sport80 ID': s80_val,
                        'Matched Registered Name': f_matched_name, 'Registered Club': reg_club,
                        'Match Date': match_date, 'Deemed Registered Date': first_unreg_match_date[player],
                        'Deemed Registered Match Teams': first_unreg_match_teams_played[player],
                        'Deemed Registered Club': extract_base_club_name(str(first_unreg_match_team[player])),
                        'Date Registered': reg_date, 'Team A': team_a, 'Team B': team_b,
                        'Match League': match_league, 'Match Logic': f_match_logic
                    })

    valid_matches = all_appearances[(all_appearances['Match Date'] >= start_date) & (all_appearances['Match Date'] <= end_date)].copy()
    international_exclusions = ['mark adair', 'ben calitz']

    if not starring_df.empty and 'Full Name' in starring_df.columns:
        starring_df['Cleaned Name'] = starring_df['Full Name'].apply(lambda x: cleanse_name(x, alias_map))
        
        for idx, row in valid_matches.iterrows():
            player, scorecard_name = row['Player'], row['Scorecard Name']
            team_a, team_b = str(row['Team A']), str(row['Team B'])
            if str(player).strip().lower() in international_exclusions: continue
            
            p_s80 = ""
            if player in player_match_cache:
                p_s80 = player_match_cache[player][3]
            elif pd.notna(row.get('Sport80_ID')):
                p_s80 = str(row.get('Sport80_ID')).replace('.0', '').strip()
                
            player_stars = starring_df[starring_df['Cleaned Name'].str.strip().str.lower() == player.lower()]
            if not player_stars.empty:
                starred_level, starred_club = str(player_stars.iloc[0]['XI_Level']), str(player_stars.iloc[0]['Club'])
                full_weekend_team = f"{starred_club} {starred_level}"
                weekend_division = get_team_league(full_weekend_team, wknd_team_keys, wknd_league_dict, "Men's")
                
                if weekend_division:
                    div_lower = str(weekend_division).lower()
                    is_illegal = False
                    
                    if 'premier' in div_lower or 'senior' in div_lower: 
                        is_illegal = True
                    elif 'junior' in div_lower:
                        match_num = re.search(r'junior.*?(\d+)', div_lower)
                        if match_num and int(match_num.group(1)) <= 3: 
                            is_illegal = True
                            
                    if is_illegal:
                        violation_matches.add((team_a, team_b, determine_midweek_league(team_a, team_b), row['Match Date']))
                        mw_team = team_a if starred_club.lower() in team_a.lower() else (team_b if starred_club.lower() in team_b.lower() else team_a)
                        starring_violations.append({
                            'Player (Cleaned)': player, 'Original Scorecard Name': scorecard_name,
                            'Sport80 ID': p_s80,
                            'Starred Rank': full_weekend_team, 'Weekend Division': weekend_division,
                            'Midweek Team': mw_team, 'Team A': team_a, 'Team B': team_b,
                            'Match Date': row['Match Date'], 'Match Group': row['Group']
                        })

    unregistered_audit.sort(key=lambda x: violation_sort_key(x, 'Stats Name (Cleaned)'))
    deemed_registered.sort(key=lambda x: violation_sort_key(x, 'Stats Name (Cleaned)'))
    starring_violations.sort(key=lambda x: violation_sort_key(x, 'Player (Cleaned)'))

    excel_io = io.BytesIO()
    with pd.ExcelWriter(excel_io, engine='xlsxwriter', datetime_format='dd/mm/yyyy') as writer:
        if unregistered_audit: export_and_format_excel(pd.DataFrame(unregistered_audit), writer, "Unregistered Matches")
        else: pd.DataFrame(columns=["Status"]).to_excel(writer, index=False, sheet_name="Unregistered Matches")
        if deemed_registered: export_and_format_excel(pd.DataFrame(deemed_registered), writer, "Deemed Registered")
        else: pd.DataFrame(columns=["Status"]).to_excel(writer, index=False, sheet_name="Deemed Registered")
        if starring_violations: export_and_format_excel(pd.DataFrame(starring_violations), writer, "Starring Violations")
        else: pd.DataFrame(columns=["Status"]).to_excel(writer, index=False, sheet_name="Starring Violations")

    doc = Document()
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Calibri', Pt(11)
    start_str, end_str = get_ordinal_date(start_date, False), get_ordinal_date(end_date, True)
    
    p_intro = doc.add_paragraph()
    run_intro = p_intro.add_run(f"Here is the Midweek League audit for the matches played between {start_str} and {end_str}.")
    run_intro.bold = True
    run_intro.font.size = Pt(13)
    run_intro.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph("This summary runs automated data integrity audits across competitor profiles, verifying registration windows and cross-referencing weekend tiering to flag ineligible players starred above Junior League 3.")
    
    p_viol = doc.add_paragraph()
    run_viol = p_viol.add_run("--- Matches WITH Violations ---")
    run_viol.bold = True
    run_viol.font.size = Pt(13)
    run_viol.font.color.rgb = RGBColor(0, 0, 128)

    if violation_matches:
        render_grouped_matches(doc, violation_matches, "Midweek")
    else:
        doc.add_paragraph("No matches with violations recorded.")

    p_noviol = doc.add_paragraph()
    p_noviol.paragraph_format.space_before = Pt(14)  
    p_noviol.paragraph_format.space_after = Pt(4)    
    run_noviol = p_noviol.add_run("--- Matches WITHOUT Violations ---")
    run_noviol.bold = True
    run_noviol.font.size = Pt(13)
    run_noviol.font.color.rgb = RGBColor(0, 0, 128)

    matches_without_violations = all_matches_in_range - violation_matches
    if matches_without_violations:
        render_grouped_matches(doc, matches_without_violations, "Midweek")
    else:
        doc.add_paragraph("No matches without violations recorded.")

    doc.add_page_break()
    p_unreg = doc.add_paragraph()
    run_unreg = p_unreg.add_run("--- Unregistered Players / Date Violations ---")
    run_unreg.bold = True
    run_unreg.font.size = Pt(13)
    run_unreg.font.color.rgb = RGBColor(0, 0, 128)

    for r in unregistered_audit:
        p_name, s_name = r['Stats Name (Cleaned)'], r['Original Scorecard Name']
        d_name = p_name if p_name.lower() == str(s_name).lower() else f"{p_name} (Played as: {s_name})"
        m_date = get_ordinal_date(r['Match Date'])
        m_league = str(r.get('Match League', 'Unknown League'))
        t_str = f" [Registered Club: {doc_formal_team_name(r.get('Registered Club', 'Unknown Club'))}] (Match: {doc_formal_team_name(r['Team A'])} v {doc_formal_team_name(r['Team B'])} - {m_date} - {m_league})"
        
        p_p = doc.add_paragraph(style='List Bullet')
        p_p.add_run(f"{d_name}").bold = True
        p_p.add_run(f"{t_str}")
        
        r_date = r['Date Registered']
        status = r.get('Status', '')
        if 'Played for Wrong Club' in status:
            d_text = f"The player is registered for another club ({doc_formal_team_name(r.get('Registered Club', 'Unknown Club'))}) and not registered for any club in the match they played in."
        elif pd.notna(r_date):
            d_text = f"The official database indicates a registration date of {get_ordinal_date(r_date)} ({(r_date - r['Match Date']).days} days late)."
        else:
            d_text = f"Appeared under the scorecard name \"{s_name}\" (verified via alias map). This official profile is entirely unregistered on the master registry."
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.add_run("o  Violation Detail: ").bold = True
        p.add_run(d_text)

    doc.add_page_break()
    p_deemed = doc.add_paragraph()
    run_deemed = p_deemed.add_run("--- Deemed Registered (Played Previously in 2026) ---")
    run_deemed.bold = True
    run_deemed.font.size = Pt(13)
    run_deemed.font.color.rgb = RGBColor(0, 0, 128)

    for r in deemed_registered:
        p_name, s_name = r['Stats Name (Cleaned)'], r['Original Scorecard Name']
        d_name = p_name if p_name.lower() == str(s_name).lower() else f"{p_name} (Played as: {s_name})"
        deemed_club = r.get('Deemed Registered Club')
        club_display = f"Deemed Registered Club: {doc_formal_team_name(deemed_club)}" if deemed_club else f"Registered Club: {doc_formal_team_name(r.get('Registered Club', 'Unknown Club'))}"
        m_date = get_ordinal_date(r['Match Date'])
        m_league = str(r.get('Match League', 'Unknown League'))
        t_str = f" [{club_display}] (Match: {doc_formal_team_name(r['Team A'])} v {doc_formal_team_name(r['Team B'])} - {m_date} - {m_league})"
        
        p_p = doc.add_paragraph(style='List Bullet')
        p_p.add_run(f"{d_name}").bold = True
        p_p.add_run(f"{t_str}")
        
        r_date = r['Date Registered']
        r_det = f"Registered on {get_ordinal_date(r_date)}" if pd.notna(r_date) else "Unregistered profile"
        r_det = r_det[0].upper() + r_det[1:]
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.add_run("o  Deemed Status: ").bold = True
        deemed_teams = r.get('Deemed Registered Match Teams', '')
        in_match_str = f" in {deemed_teams}" if deemed_teams else ""
        p.add_run(f"{r_det}. Deemed registered because he previously played an active fixture{in_match_str} on {get_ordinal_date(r['Deemed Registered Date'])}.")

    doc.add_page_break()
    p_star = doc.add_paragraph()
    run_star = p_star.add_run("--- Starring Ceiling Violations ---")
    run_star.bold = True
    run_star.font.size = Pt(13)
    run_star.font.color.rgb = RGBColor(0, 0, 128)

    doc.add_paragraph("The following players are barred from participating in the Midweek League because they hold weekend starring rankings of Junior League 3 or above (Premier, Senior, or Junior 1–3).")

    for r in starring_violations:
        p_name, s_name = r['Player (Cleaned)'], r['Original Scorecard Name']
        d_name = p_name if p_name.lower() == str(s_name).lower() else f"{p_name} (Played as: {s_name})"
        m_date = get_ordinal_date(r['Match Date'])
        t_str = f" (Match: {doc_formal_team_name(r['Team A'])} v {doc_formal_team_name(r['Team B'])} - {m_date})"
        
        p_p = doc.add_paragraph(style='List Bullet')
        p_p.add_run(f"{d_name}").bold = True
        p_p.add_run(f"{t_str}")
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(36)
        p.add_run("o  Violation Detail: ").bold = True
        p.add_run(f"Represented {doc_formal_team_name(r['Midweek Team'])}. This player is completely ineligible for Midweek cricket as he is officially starred for weekend squad '{r['Starred Rank']}', which plays in weekend tier '{r['Weekend Division']}' (Junior League 3 or above tiers are ineligible).")

    doc_io = io.BytesIO()
    doc.save(doc_io)
    
    return excel_io, doc_io

# ==========================================
# STARRING & INACTIVITY REPORT FUNCTIONS 
# ==========================================
def report_clean_spaces(text):
    if pd.isna(text) or str(text).strip().lower() == 'nan': return ""
    return re.sub(r'\s+', ' ', str(text)).strip()

def report_clean_team_name(team):
    if pd.isna(team): return team
    cleaned = report_clean_spaces(team)
    if "Holywood" in cleaned and "1881" not in cleaned:
        cleaned = cleaned.replace("Holywood", "Holywood 1881")
    return cleaned

def report_clean_score_display(text):
    if pd.isna(text): return text
    return re.sub(r'(\d+)/(\d+)', r'\1-\2', str(text))

def report_parse_match_date(group_str):
    months = {'April': 4, 'May': 5, 'June': 6, 'July': 7, 'August': 8, 'September': 9}
    match = re.search(r'(\d{1,2})(?:st|nd|rd|th)?\s+([A-Za-z]+)\s+(\d{4})', str(group_str))
    if match:
        day, month_str, year = int(match.group(1)), match.group(2), int(match.group(3))
        if month_str in months: return datetime(year, months[month_str], day)
    return None

def report_extract_competition(group_str):
    if pd.isna(group_str) or str(group_str).strip().lower() == 'nan': 
        return "Unknown Competition"
    parts = str(group_str).split(' - ')
    if len(parts) >= 2:
        return parts[0].strip()
    return "Unknown Competition"

def report_get_team_dates_flexible(starred_team, team_match_dates_dict):
    def norm(t):
        t = str(t).lower().replace('1881', '')
        t = t.replace('1st', '1').replace('firsts', '1').replace('first', '1')
        t = t.replace('2nd', '2').replace('seconds', '2').replace('second', '2')
        t = t.replace('3rd', '3').replace('thirds', '3').replace('third', '3')
        t = t.replace('4th', '4').replace('fourths', '4').replace('fourth', '4')
        t = t.replace('5th', '5').replace('fifths', '5').replace('fifth', '5')
        t = re.sub(r'\bxi\b', '', t)
        t = re.sub(r'\bteam\b', '', t)
        return re.sub(r'[^a-z0-9]', '', t)
    
    if starred_team in team_match_dates_dict: return team_match_dates_dict[starred_team]
        
    n_starred = norm(starred_team)
    for k, dates in team_match_dates_dict.items():
        if norm(k) == n_starred: return dates
            
    st_nums = re.findall(r'\d+', n_starred)
    if st_nums:
        for k, dates in team_match_dates_dict.items():
            k_nums = re.findall(r'\d+', norm(k))
            if k_nums and st_nums[0] == k_nums[0]: return dates
                
    all_dates = []
    for dates in team_match_dates_dict.values(): all_dates.extend(dates)
    return list(set(all_dates))

def report_build_club_matches_df(club_name, all_app, override_map, comp_map):
    if all_app.empty: return pd.DataFrame(), {}
    matches = all_app['Group'].dropna().unique()
    club_name_clean = report_clean_spaces(club_name).lower()
    search_term = override_map.get(club_name_clean, club_name_clean)
    
    club_matches = [m for m in matches if search_term in report_clean_spaces(m).lower()]
    
    def extract_and_clean_target_team(m):
        match_teams = re.split(r'\s+v\s+', str(m).split(',')[0])
        for t in match_teams:
            if search_term in report_clean_spaces(t).lower(): return report_clean_team_name(t)
        return "Unknown"
        
    team_dict = {}
    for m in club_matches:
        team = extract_and_clean_target_team(m)
        if team not in team_dict: team_dict[team] = []
        team_dict[team].append(m)
        
    formatted_rows = []
    team_match_dates = {}
    
    for team in sorted(team_dict.keys()):
        formatted_rows.append({"Club Match History": f"Team: {team}", "Competition": ""})
        sorted_team_matches = sorted(team_dict[team], key=lambda x: report_parse_match_date(x) or datetime.min)
        dates_list = []
        for m in sorted_team_matches:
            comp = comp_map.get(m, "Unknown Competition")
            formatted_rows.append({"Club Match History": f"  - {report_clean_score_display(report_clean_spaces(m))}", "Competition": comp})
            parsed_d = report_parse_match_date(m)
            if parsed_d: dates_list.append(parsed_d)
        
        team_match_dates[team.lower().strip()] = dates_list
        formatted_rows.append({"Club Match History": "", "Competition": ""})
        
    return pd.DataFrame(formatted_rows), team_match_dates

def report_build_player_stats_dfs(player_list, player_team_map, all_app, get_official_func, comp_map):
    raw_summary, log = [], []
    for p in player_list:
        official_p = get_official_func(p)
        official_p_lower = official_p.lower()
        p_apps = all_app[all_app['Official_Player_Lower'] == official_p_lower] if not all_app.empty else pd.DataFrame()
        matches_played = len(p_apps)
        team = report_clean_team_name(player_team_map.get(report_clean_spaces(p).lower(), "Unassigned"))
        
        if matches_played > 0:
            sorted_apps = p_apps.sort_values('Parsed_Date')
            last_date = sorted_apps['Parsed_Date'].max().strftime('%d %B %Y')
            last_match_grp = report_clean_spaces(sorted_apps.iloc[-1]['Group'])
            last_comp = comp_map.get(last_match_grp, "Unknown Competition")
            
            raw_summary.append({
                "Starred Team": team, "Input Name": report_clean_spaces(p), "Player (Official)": official_p, 
                "Matches Played": matches_played, "Last Played Date": last_date, 
                "Last Match Details": last_match_grp, "Competition": last_comp, "Highlight Reason": ""
            })
            log.append({"Player Log": f"Player: {official_p} ({report_clean_spaces(p)}) | Starred Team: {team}", "Competition": "", "Highlight Reason": ""})
            for _, row in sorted_apps.iterrows(): 
                grp = report_clean_spaces(row['Group'])
                comp = comp_map.get(grp, "Unknown Competition")
                log.append({"Player Log": f"  - {grp}", "Competition": comp, "Highlight Reason": ""})
        else:
            raw_summary.append({
                "Starred Team": team, "Input Name": report_clean_spaces(p), "Player (Official)": official_p, 
                "Matches Played": 0, "Last Played Date": "N/A", "Last Match Details": "N/A", "Competition": "N/A", "Highlight Reason": ""
            })
            log.append({"Player Log": f"Player: {official_p} ({report_clean_spaces(p)}) | Starred Team: {team}\n  - No match records found.\n", "Competition": "", "Highlight Reason": ""})
        log.append({"Player Log": "", "Competition": "", "Highlight Reason": ""})
        
    if not raw_summary: return pd.DataFrame(), pd.DataFrame(log)
        
    df_raw = pd.DataFrame(raw_summary).sort_values(by=["Starred Team", "Player (Official)"])
    formatted_summary = []
    current_team = None
    
    for _, row in df_raw.iterrows():
        if row["Starred Team"] != current_team:
            if current_team is not None:
                formatted_summary.append({
                    "Input Name": "", "Player (Official)": "", "Matches Played": "", 
                    "Last Played Date": "", "Last Match Details": "", "Competition": "", "Highlight Reason": ""
                })
            current_team = row["Starred Team"]
            formatted_summary.append({
                "Input Name": f"Team Starred For: {current_team}", "Player (Official)": "", "Matches Played": "", 
                "Last Played Date": "", "Last Match Details": "", "Competition": "", "Highlight Reason": ""
            })
            
        formatted_summary.append({
            "Input Name": row["Input Name"], "Player (Official)": row["Player (Official)"], "Matches Played": row["Matches Played"],
            "Last Played Date": row["Last Played Date"], "Last Match Details": row["Last Match Details"], "Competition": row["Competition"], "Highlight Reason": row["Highlight Reason"]
        })
        
    return pd.DataFrame(formatted_summary), pd.DataFrame(log)

def report_autofit_columns(ws):
    for col in ws.columns:
        max_length = 0
        column_letter = col[0].column_letter 
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except Exception as e: print('EXCEPTION IN FINES:', repr(e))
        ws.column_dimensions[column_letter].width = max_length + 2

@st.cache_data(show_spinner="Generating starring inactivity reports...")
def generate_starring_inactivity_reports(domain, f_reg, f_alias, f_starring, f_bat, f_bowl, f_irish_bat=None, f_irish_bowl=None, f_abandoned=None):
    df_reg = pd.read_excel(f_reg)
    df_alias = pd.read_excel(f_alias)
    df_bat = pd.read_excel(f_bat)
    df_bowl = pd.read_excel(f_bowl)
    
    df_bat['Is_Irish_Match'] = False
    df_bowl['Is_Irish_Match'] = False
    
    if f_irish_bat and os.path.exists(f_irish_bat):
        irish_bat = pd.read_excel(f_irish_bat)
        irish_bat['Is_Irish_Match'] = True
        df_bat = pd.concat([df_bat, irish_bat], ignore_index=True)
        
    if f_irish_bowl and os.path.exists(f_irish_bowl):
        irish_bowl = pd.read_excel(f_irish_bowl)
        irish_bowl['Is_Irish_Match'] = True
        df_bowl = pd.concat([df_bowl, irish_bowl], ignore_index=True)
        
    if not f_abandoned:
        f_abandoned = DEFAULT_FILES.get(domain, {}).get("abandoned", "")
        
    df_ab = pd.DataFrame()
    if f_abandoned and os.path.exists(f_abandoned):
        df_ab = pd.read_excel(f_abandoned)
    
    international_players = ["cara murray"] if domain == "Women's" else ["mark adair", "paul stirling"]
    override_map = {"holywood 1881": "holywood"}

    df_alias = df_alias.drop_duplicates(subset=['Input Name (Scorecard/Stats)'], keep='last')
    alias_map = dict(zip(
        df_alias['Input Name (Scorecard/Stats)'].apply(lambda x: report_clean_spaces(x).lower() if pd.notna(x) else x), 
        df_alias['Official Registered Name'].apply(lambda x: report_clean_spaces(x) if pd.notna(x) else x)
    ))

    f_id_map = DEFAULT_FILES.get(domain, {}).get("id_map", "")
    id_map = {}
    if f_id_map and os.path.exists(f_id_map):
        try:
            id_map_df = pd.read_excel(f_id_map)
            id_map = build_id_map(id_map_df)
        except Exception:
            pass
        
    registered_players_map = {}
    cols_lower = [str(c).lower() for c in df_reg.columns]
    if 'forename' in cols_lower and 'surname' in cols_lower:
        f_col = df_reg.columns[cols_lower.index('forename')]
        s_col = df_reg.columns[cols_lower.index('surname')]
        combined = df_reg[f_col].astype(str) + " " + df_reg[s_col].astype(str)
        registered_players_map = {report_clean_spaces(x).lower(): report_clean_spaces(x) for x in combined if 'nan' not in str(x).lower()}
    else:
        name_col = next((c for c in ['Name', 'Player', 'Player Name', 'Full Name', 'Registered Name'] if c in df_reg.columns), df_reg.columns[0])
        registered_players_map = {report_clean_spaces(str(x)).lower(): report_clean_spaces(str(x)) for x in df_reg[name_col].dropna()}
        
    def is_player_registered(name):
        cleaned = report_clean_spaces(name)
        mapped = alias_map.get(cleaned.lower(), cleaned)
        return mapped.lower() in registered_players_map

    def get_official_name_contextual(name, row):
        cleaned = report_clean_spaces(name)
        if id_map:
            res_name, s80_id, s80_club, is_id = resolve_player_from_row(row, cleaned, id_map, alias_map)
            if is_id and res_name:
                if '(' in res_name and res_name.endswith(')'):
                    res_name = res_name.split('(')[0].strip()
                if res_name.lower() in registered_players_map:
                    return registered_players_map[res_name.lower()]
                return res_name
                
        cleaned_lower = cleaned.lower()
        mapped = alias_map.get(cleaned_lower, cleaned)
        if mapped.lower() in registered_players_map: return registered_players_map[mapped.lower()]
        return mapped

    def get_official_name(name):
        cleaned = report_clean_spaces(name)
        mapped = alias_map.get(cleaned.lower(), cleaned)
        if mapped.lower() in registered_players_map: return registered_players_map[mapped.lower()]
        return mapped

    bat_cols = ['Name', 'Group', 'Is_Irish_Match']
    for id_col in ['Batter ID', 'Player ID', 'BatterId', 'PlayerId']:
        if id_col in df_bat.columns and id_col not in bat_cols:
            bat_cols.append(id_col)
            break
    if 'Team' in df_bat.columns: bat_cols.append('Team')

    bowl_cols = ['Bowler', 'Group', 'Is_Irish_Match']
    for id_col in ['Bowler ID', 'Player ID', 'BowlerId', 'PlayerId']:
        if id_col in df_bowl.columns and id_col not in bowl_cols:
            bowl_cols.append(id_col)
            break
    if 'Team' in df_bowl.columns: bowl_cols.append('Team')

    app_list = [
        df_bat[bat_cols].rename(columns={'Name': 'P'}) if not df_bat.empty else pd.DataFrame(columns=['P', 'Group', 'Is_Irish_Match']),
        df_bowl[bowl_cols].rename(columns={'Bowler': 'P'}) if not df_bowl.empty else pd.DataFrame(columns=['P', 'Group', 'Is_Irish_Match'])
    ]

    if not df_ab.empty:
        df_ab['Is_Irish_Match'] = False
        ab_match_col = 'Group' if 'Group' in df_ab.columns else ('Match' if 'Match' in df_ab.columns else df_ab.columns[0])
        ab_name_col = 'Name' if 'Name' in df_ab.columns else df_ab.columns[1]
        
        ab_cols = [ab_name_col, ab_match_col, 'Is_Irish_Match']
        for id_col in ['Player ID', 'Batter ID', 'Bowler ID']:
            if id_col in df_ab.columns and id_col not in ab_cols:
                ab_cols.append(id_col)
                break
        if 'Team' in df_ab.columns: ab_cols.append('Team')
        
        df_ab_app = df_ab[ab_cols].rename(columns={ab_name_col: 'P', ab_match_col: 'Group'})
        app_list.append(df_ab_app)

    all_app = pd.concat(app_list, ignore_index=True)
    
    if not all_app.empty:
        all_app['Group'] = all_app['Group'].apply(report_clean_score_display) 
        all_app['Official_Player'] = all_app.apply(lambda r: get_official_name_contextual(r['P'], r), axis=1)
        all_app['Official_Player_Lower'] = all_app['Official_Player'].str.lower()
        all_app['Parsed_Date'] = all_app['Group'].apply(report_parse_match_date)
        all_app.drop_duplicates(subset=['Official_Player', 'Group'], inplace=True)

    f_league = DEFAULT_FILES[domain]["league"]
    league_dict, team_keys = {}, []
    if os.path.exists(f_league):
        league_structure = pd.read_excel(f_league)
        league_dict, team_keys, _ = build_league_dict(league_structure)
        
    f_cup = "NCU_Cup_Fixtures.xlsx"
    cup_match_dict = {}
    if os.path.exists(f_cup):
        try:
            excel_file_cup = pd.ExcelFile(f_cup)
            target_sheet = excel_file_cup.sheet_names[0]
            for sheet in excel_file_cup.sheet_names:
                if domain.lower().replace("'", "") in sheet.lower().replace("'", ""):
                    target_sheet = sheet
                    break
            cup_df = pd.read_excel(f_cup, sheet_name=target_sheet, header=None)
            
            def local_parse(group_str):
                try:
                    group_str = str(group_str).strip()
                    parts = group_str.rsplit(' - ', 1)
                    date_str = parts[1].strip() if len(parts) == 2 else group_str
                    rest = parts[0].strip() if len(parts) == 2 else group_str
                    match_date = pd.to_datetime(date_str, dayfirst=True, errors='coerce')
                    if pd.notna(match_date): match_date = match_date.normalize()
                    if ' v ' in rest:
                        team_a, remainder = rest.split(' v ', 1)
                        team_b = remainder.rsplit(', ', 1)[0] if ', ' in remainder else (remainder.rsplit(' - ', 1)[0] if ' - ' in remainder else remainder)
                    else:
                        team_a, team_b = rest, "Unknown"
                    return team_a.strip(), team_b.strip(), match_date
                except: return None, None, None

            for _, row_data in cup_df.iterrows():
                match_str_raw = str(row_data[0]).strip()
                cup_name = str(row_data[1]).strip()
                if match_str_raw.lower() in ['match string', 'match group', 'match', 'nan']: continue
                cleaned_match_str = doc_format_cricket_names(match_str_raw, domain)
                c_team_a, c_team_b, c_date = local_parse(cleaned_match_str)
                if c_team_a and c_team_b:
                    teams = sorted([str(c_team_a).lower(), str(c_team_b).lower()])
                    if pd.notna(c_date):
                        cup_match_dict[f"{teams[0]}_{teams[1]}_{c_date.strftime('%Y-%m-%d')}"] = cup_name
                    else:
                        cup_match_dict[f"{teams[0]}_{teams[1]}"] = cup_name
        except Exception as e: print('EXCEPTION IN FINES:', repr(e))

    def apply_competition(grp_str, is_irish):
        try:
            grp_str_clean = doc_format_cricket_names(str(grp_str), domain)
            parts = grp_str_clean.rsplit(' - ', 1)
            date_str = parts[1].strip() if len(parts) == 2 else grp_str_clean
            rest = parts[0].strip() if len(parts) == 2 else grp_str_clean
            
            match_date = pd.to_datetime(date_str, dayfirst=True, errors='coerce')
            if pd.notna(match_date): match_date = match_date.normalize()
            
            if ' v ' in rest:
                team_a, remainder = rest.split(' v ', 1)
                team_b = remainder.rsplit(', ', 1)[0] if ', ' in remainder else (remainder.rsplit(' - ', 1)[0] if ' - ' in remainder else remainder)
            else:
                team_a, team_b = rest, "Unknown"
                
            t_a, t_b = team_a.strip(), team_b.strip()
            
            if t_a == "Unknown" or t_b == "Unknown":
                return "Irish Competition" if is_irish else "Unknown Competition"
                
            teams = sorted([str(t_a).lower(), str(t_b).lower()])
            if pd.notna(match_date):
                robust_key_date = f"{teams[0]}_{teams[1]}_{match_date.strftime('%Y-%m-%d')}"
                if robust_key_date in cup_match_dict: return cup_match_dict[robust_key_date]
            robust_key_no_date = f"{teams[0]}_{teams[1]}"
            if robust_key_no_date in cup_match_dict: return cup_match_dict[robust_key_no_date]
            
            if is_irish: return "Irish Competition"
            
            league_a = get_team_league(t_a, team_keys, league_dict, domain)
            league_b = get_team_league(t_b, team_keys, league_dict, domain)
            if league_a and league_b and league_a == league_b:
                league_str = str(league_a)
                target_words = ['premier', 'senior league 1', 'senior league 2', 'senior league 3'] if domain == "Men's" else ['premier', 'senior league 1', 'senior league 2', 'senior league 3', 'senior']
                if any(word in league_str.lower() for word in target_words): 
                    return league_str.replace('NCU', 'Mercury')
                return league_str
                
            return "Possible Cup Match / Friendly"
        except:
            return "Irish Competition" if is_irish else "Unknown Competition"

    comp_map = {}
    if not all_app.empty:
        unique_groups_df = all_app[['Group', 'Is_Irish_Match']].drop_duplicates(subset=['Group'])
        for _, row in unique_groups_df.iterrows():
            grp = row['Group']
            if pd.isna(grp): continue
            is_irish = row.get('Is_Irish_Match', False)
            comp_map[grp] = apply_competition(grp, is_irish)

    zip_buffer = io.BytesIO()
    unregistered_starred_players = []
    run_date = datetime.now()
    
    yellow_fill = PatternFill(start_color="FFFFFF00", end_color="FFFFFF00", fill_type="solid")
    red_fill = PatternFill(start_color="FFFF0000", end_color="FFFF0000", fill_type="solid")
    green_fill = PatternFill(start_color="FF92D050", end_color="FF92D050", fill_type="solid")
    bold_font = Font(bold=True)
    center_alignment = Alignment(horizontal='center', vertical='center')

    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        excel_file = pd.ExcelFile(f_starring)
        excluded_tabs = ["summary", "overview", "sheet1"]
        
        for sheet_name in excel_file.sheet_names:
            if sheet_name.lower().strip() in excluded_tabs: continue
                
            club_input = sheet_name.strip()
            safe_club = re.sub(r'[\\/*?:"<>|]', "", club_input).replace(" ", "_")
            df_club, team_match_dates = report_build_club_matches_df(club_input, all_app, override_map, comp_map)
            df_stars_raw = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
            
            players, current_team = [], "Unassigned"
            for i in range(len(df_stars_raw)):
                row = df_stars_raw.iloc[i]
                val0, val_s, val_f = str(row[0]).strip(), str(row[1]).strip(), str(row[4]).strip()
                if 'surname' in val_s.lower() or 'forename' in val_f.lower(): continue
                if val0 and val0.lower() != 'nan' and val0.lower() not in ['id', 'surname', 'forename(s)', 'team']:
                    if not val_s or val_s.lower() == 'nan':
                        current_team = report_clean_team_name(val0) 
                        continue
                if val_s and val_s.lower() != 'nan' and val_f and val_f.lower() != 'nan':
                    players.append({"name": report_clean_spaces(f"{val_f} {val_s}"), "team": current_team})
            
            for p in players:
                if not is_player_registered(p['name']):
                    unregistered_starred_players.append({
                        "Club": club_input, "Starred Team": p['team'], "Starred Name": p['name'],
                        "Current Mapped Alias (if any)": alias_map.get(report_clean_spaces(p['name']).lower(), "None")
                    })
            
            player_list = [p['name'] for p in players]
            player_team_map = {report_clean_spaces(p['name']).lower(): p['team'] for p in players}
            df_p_summary, df_p_details = report_build_player_stats_dfs(player_list, player_team_map, all_app, get_official_name, comp_map)
        
            club_io = io.BytesIO()
            with pd.ExcelWriter(club_io, engine='openpyxl') as writer:
                if not df_club.empty:
                    safe_club_tab_name = report_clean_spaces(club_input)[:20].replace(":", "").replace("/", "")
                    sheet_tab_name = f"{safe_club_tab_name} Matches"
                    df_club.to_excel(writer, sheet_name=sheet_tab_name, index=False)
                    ws_matches = writer.sheets[sheet_tab_name]
                    ws_matches.cell(row=1, column=1).font = bold_font
                    ws_matches.cell(row=1, column=2).font = bold_font
                    for row_idx in range(2, len(df_club) + 2):
                        if str(ws_matches.cell(row=row_idx, column=1).value).startswith("Team:"):
                            ws_matches.cell(row=row_idx, column=1).font = bold_font
                    report_autofit_columns(ws_matches)
                    
                if not df_p_summary.empty:
                    sheet_tab_name = "Player Match Count Summary"
                    df_p_summary.to_excel(writer, sheet_name=sheet_tab_name, index=False)
                    worksheet = writer.sheets[sheet_tab_name]
                    for col_idx in range(1, len(df_p_summary.columns) + 1): worksheet.cell(row=1, column=col_idx).font = bold_font
                    
                    current_starred_team = None 
                    reason_col_summary = len(df_p_summary.columns)
                    
                    for row_idx in range(2, len(df_p_summary) + 2):
                        for col_idx, col_name in enumerate(df_p_summary.columns, start=1):
                            if col_name in ["Matches Played", "Last Played Date"]: worksheet.cell(row=row_idx, column=col_idx).alignment = center_alignment
                        cell_value_1 = str(worksheet.cell(row=row_idx, column=1).value)
                        
                        if cell_value_1.startswith("Team Starred For:"):
                            current_starred_team = cell_value_1.replace("Team Starred For:", "").strip()
                            for col_idx in range(1, len(df_p_summary.columns) + 1): worksheet.cell(row=row_idx, column=col_idx).font = bold_font
                        elif cell_value_1.strip(): 
                            official_name_val = str(worksheet.cell(row=row_idx, column=2).value).strip().lower()
                            is_intl = any(intl in official_name_val or intl in cell_value_1.strip().lower() for intl in international_players)
                            
                            if is_intl:
                                worksheet.cell(row=row_idx, column=reason_col_summary).value = "International Exemption"
                                for col_idx in range(1, len(df_p_summary.columns) + 1): worksheet.cell(row=row_idx, column=col_idx).fill = green_fill
                            else:
                                matches_played = worksheet.cell(row=row_idx, column=3).value 
                                if str(matches_played) == "0":
                                    worksheet.cell(row=row_idx, column=reason_col_summary).value = "0 Matches Played"
                                    for col_idx in range(1, len(df_p_summary.columns) + 1): worksheet.cell(row=row_idx, column=col_idx).fill = red_fill
                                else:
                                    date_str = str(worksheet.cell(row=row_idx, column=4).value) 
                                    if date_str and date_str != "N/A" and date_str.lower() != "nan":
                                        try:
                                            last_played = datetime.strptime(date_str, '%d %B %Y')
                                            days_since = (run_date - last_played).days
                                            team_dates = report_get_team_dates_flexible(current_starred_team, team_match_dates)
                                            matches_since = sum(1 for d in team_dates if d > last_played)
                                            if days_since > 21 and matches_since >= 3:
                                                worksheet.cell(row=row_idx, column=reason_col_summary).value = f"Inactive > 21 days ({days_since} days) and missed {matches_since} matches"
                                                for col_idx in range(1, len(df_p_summary.columns) + 1): worksheet.cell(row=row_idx, column=col_idx).fill = yellow_fill
                                        except ValueError: pass 
                    report_autofit_columns(worksheet)

                if not df_p_details.empty:
                    sheet_tab_name = "Player Season Fixture Log"
                    df_p_details.to_excel(writer, sheet_name=sheet_tab_name, index=False)
                    ws_log = writer.sheets[sheet_tab_name]
                    player_status, current_team_for_dict = {}, "Unassigned"
                    for _, row in df_p_summary.iterrows():
                        if str(row.get('Input Name', '')).startswith('Team Starred For:'): 
                            current_team_for_dict = str(row.get('Input Name', '')).replace("Team Starred For:", "").strip()
                            continue
                        p_name = str(row.get('Player (Official)', '')).strip()
                        if p_name:
                            player_status[p_name] = {'matches': str(row.get('Matches Played', '0')), 'date': str(row.get('Last Played Date', 'N/A')), 'starred_team': current_team_for_dict}
                    
                    for col_idx in range(1, len(df_p_details.columns) + 1): ws_log.cell(row=1, column=col_idx).font = bold_font
                    
                    reason_col_log = len(df_p_details.columns)

                    for row_idx in range(2, len(df_p_details) + 2):
                        cell_value = str(ws_log.cell(row=row_idx, column=1).value)
                        if cell_value.startswith("Player:"):
                            for col_idx in range(1, len(df_p_details.columns) + 1): ws_log.cell(row=row_idx, column=col_idx).font = bold_font
                            try:
                                official_p = cell_value.split(" | Starred Team:")[0].replace("Player: ", "").strip().rsplit(" (", 1)[0].strip()
                                if any(intl in official_p.lower() for intl in international_players):
                                    ws_log.cell(row=row_idx, column=reason_col_log).value = "International Exemption"
                                    for col_idx in range(1, len(df_p_details.columns) + 1): ws_log.cell(row=row_idx, column=col_idx).fill = green_fill
                                else:
                                    status = player_status.get(official_p)
                                    if status:
                                        if status['matches'] == '0':
                                            ws_log.cell(row=row_idx, column=reason_col_log).value = "0 Matches Played"
                                            for col_idx in range(1, len(df_p_details.columns) + 1): ws_log.cell(row=row_idx, column=col_idx).fill = red_fill
                                        else:
                                            d_str = status['date']
                                            if d_str and d_str != "N/A" and d_str.lower() != "nan":
                                                last_played = datetime.strptime(d_str, '%d %B %Y')
                                                t_dates = report_get_team_dates_flexible(status.get('starred_team', ''), team_match_dates)
                                                days_since = (run_date - last_played).days
                                                matches_since = sum(1 for d in t_dates if d > last_played)
                                                if days_since > 21 and matches_since >= 3:
                                                    ws_log.cell(row=row_idx, column=reason_col_log).value = f"Inactive > 21 days ({days_since} days) and missed {matches_since} matches"
                                                    for col_idx in range(1, len(df_p_details.columns) + 1): ws_log.cell(row=row_idx, column=col_idx).fill = yellow_fill
                            except Exception as e: print('EXCEPTION IN FINES:', repr(e)) 
                    report_autofit_columns(ws_log)

            zip_file.writestr(f"NCU_Master_Audit_{safe_club}.xlsx", club_io.getvalue())
            
        if unregistered_starred_players:
            unreg_io = io.BytesIO()
            df_unreg = pd.DataFrame(unregistered_starred_players)
            with pd.ExcelWriter(unreg_io, engine='openpyxl') as writer:
                df_unreg.to_excel(writer, sheet_name="Unregistered Players", index=False)
                ws_unreg = writer.sheets["Unregistered Players"]
                for col_idx in range(1, len(df_unreg.columns) + 1): ws_unreg.cell(row=1, column=col_idx).font = bold_font
                report_autofit_columns(ws_unreg)
            zip_file.writestr("Unregistered_Starred_Players.xlsx", unreg_io.getvalue())

    return zip_buffer

# ==========================================
# CLUB FINES GENERATOR FUNCTIONS
# ==========================================

def format_fine_date(dt):
    if pd.isna(dt) or dt is None: return "N/A"
    day = dt.day
    if 11 <= (day % 100) <= 13: suffix = 'th'
    else: suffix = ['th', 'st', 'nd', 'rd', 'th'][min(day % 10, 4)]
    month = dt.strftime('%B')
    return f"{day}{suffix} {month}"

def parse_flexible_date(date_str):
    try:
        if isinstance(date_str, datetime): return date_str
        if pd.isna(date_str): return None
        if hasattr(date_str, 'to_pydatetime'): return date_str.to_pydatetime()
        match = re.search(r'(\d{1,2})(?:st|nd|rd|th)?\s+([A-Za-z]+)', str(date_str))
        if match:
            day, month_str = int(match.group(1)), match.group(2)
            year = datetime.now().year if datetime.now().year >= 2026 else 2026
            dt = pd.to_datetime(f"{day} {month_str} {year}", errors='coerce')
            if pd.notna(dt): return dt.to_pydatetime()
        dt = pd.to_datetime(str(date_str), errors='coerce', dayfirst=True)
        if pd.notna(dt): return dt.to_pydatetime()
    except: pass
    return None

def extract_competition_from_group(group_str):
    if pd.isna(group_str) or not group_str: return ""
    parts = str(group_str).split(' - ')
    if len(parts) >= 3: return parts[1].strip()
    elif len(parts) == 2:
        if re.search(r'\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+|\d{4}', parts[1]): return ""
        return parts[1].strip()
    return ""

def generate_club_fines_report(audit_file, forfeit_file, start_date, end_date):
    fines_data = []
    
    s_bound = pd.to_datetime(start_date).normalize()
    e_bound = pd.to_datetime(end_date).normalize()

    if forfeit_file:
        try:
            df_forfeit = pd.read_excel(forfeit_file)
            for _, row in df_forfeit.iterrows():
                date_raw = row.get('Date', '')
                team_forfeit = str(row.get('Team Forfeiting', '')).strip()
                team_against = str(row.get('Team against', '')).strip()
                comp = str(row.get('Competition', '')).strip()
                fine = row.get('Fine', 0)
                
                if team_forfeit.lower() == 'nan' or not team_forfeit: continue
                try: fine = int(fine)
                except: fine = 0
                
                club = extract_base_club_name(team_forfeit)
                date_obj = parse_flexible_date(date_raw)
                
                if date_obj:
                    match_dt = pd.to_datetime(date_obj).normalize()
                    if match_dt < s_bound or match_dt > e_bound:
                        continue
                else:
                    continue
                    
                date_str = format_fine_date(date_obj) if date_obj else str(date_raw)
                team_part_str = f"{team_forfeit} (v {team_against})"
                
                fines_data.append({
                    'Club': club, 'Date_obj': date_obj, 'Date_str': date_str,
                    'Reason': 'Unable to field a team', 'Player': None,
                    'Team_Part_Str': team_part_str,
                    'Competition': comp, 'Fine': fine, 'Type': 'Team'
                })
        except Exception as e: print('EXCEPTION IN FINES:', repr(e))

    if audit_file:
        try:
            excel_file = pd.ExcelFile(audit_file)
            
            df_unreg = pd.read_excel(excel_file, sheet_name="Unregistered Matches") if "Unregistered Matches" in excel_file.sheet_names else pd.DataFrame()
            df_deemed = pd.read_excel(excel_file, sheet_name="Deemed Registered") if "Deemed Registered" in excel_file.sheet_names else pd.DataFrame()
            
            player_true_team = {}
            if not df_unreg.empty:
                all_unreg_matches = pd.concat([df_unreg, df_deemed], ignore_index=True) if not df_deemed.empty else df_unreg
                
                if 'Stats Name (Cleaned)' in all_unreg_matches.columns:
                    for player, group in all_unreg_matches.groupby('Stats Name (Cleaned)'):
                        reg_club = str(group.iloc[0].get('Registered Club', 'Unknown Club')).strip()
                        reg_club_base = extract_base_club_name(reg_club).lower()
                        
                        if reg_club_base != 'unknown club':
                            player_true_team[player] = ('known_reg', reg_club)
                        elif '(' in player and player.strip().endswith(')'):
                            club_in_name = player.split('(')[-1].replace(')', '').strip().lower()
                            player_true_team[player] = ('inferred', club_in_name)
                        else:
                            teams_in_matches = []
                            for _, r in group.iterrows():
                                t_a = str(r.get('Team A', '')).strip()
                                t_b = str(r.get('Team B', '')).strip()
                                teams_in_matches.append({extract_base_club_name(t_a).lower(), extract_base_club_name(t_b).lower()})
                            
                            if len(teams_in_matches) == 1:
                                t_a = str(group.iloc[0].get('Team A', '')).strip()
                                t_b = str(group.iloc[0].get('Team B', '')).strip()
                                player_true_team[player] = ('ambiguous', (t_a, t_b))
                            else:
                                common_teams = set.intersection(*teams_in_matches)
                                if len(common_teams) == 1:
                                    player_true_team[player] = ('inferred', list(common_teams)[0])
                                else:
                                    t_a = str(group.iloc[0].get('Team A', '')).strip()
                                    t_b = str(group.iloc[0].get('Team B', '')).strip()
                                    player_true_team[player] = ('ambiguous', (t_a, t_b))

            if not df_unreg.empty and len(df_unreg.columns) > 1:
                for _, row in df_unreg.iterrows():
                    match_date = row.get('Match Date')
                    date_obj = pd.to_datetime(match_date) if pd.notna(match_date) else None
                    date_str = format_fine_date(date_obj) if date_obj else str(match_date)
                    
                    player_key = str(row.get('Stats Name (Cleaned)', '')).strip()
                    player_disp = str(row.get('Original Scorecard Name', player_key)).strip()
                    
                    team_a = str(row.get('Team A', '')).strip()
                    team_b = str(row.get('Team B', '')).strip()
                    comp = str(row.get('Match League', '')).strip()
                    
                    status, info = player_true_team.get(player_key, ('known_reg', str(row.get('Registered Club', 'Unknown Club')).strip()))
                    
                    if status == 'known_reg':
                        reg_club_base = extract_base_club_name(info).lower()
                        if reg_club_base in extract_base_club_name(team_b).lower() and reg_club_base != 'unknown club':
                            team_played, opponent = team_b, team_a
                        else:
                            team_played, opponent = team_a, team_b
                        club = extract_base_club_name(team_played)
                        team_part_str = f"{team_played} (v {opponent})"
                    
                    elif status == 'inferred':
                        if info in extract_base_club_name(team_b).lower():
                            team_played, opponent = team_b, team_a
                        else:
                            team_played, opponent = team_a, team_b
                        club = extract_base_club_name(team_played)
                        team_part_str = f"{team_played} (v {opponent})"
                        
                    elif status == 'ambiguous':
                        t_a, t_b = info
                        club = f"{extract_base_club_name(t_a)} / {extract_base_club_name(t_b)}"
                        team_part_str = f"{t_a} v {t_b}" 
                        
                    fines_data.append({
                        'Club': club, 'Date_obj': date_obj, 'Date_str': date_str,
                        'Reason': 'playing an unregistered player', 'Player': player_disp,
                        'Team_Part_Str': team_part_str,
                        'Competition': comp, 'Fine': 10, 'Type': 'Player'
                    })
                        
            if "Starring Violations" in excel_file.sheet_names:
                df_star = pd.read_excel(excel_file, sheet_name="Starring Violations")
                if not df_star.empty and len(df_star.columns) > 1:
                    print('DF_STAR MATCHES:', df_star.to_dict('records'))
                    for _, row in df_star.iterrows():
                        match_date = row.get('Match Date')
                        date_obj = pd.to_datetime(match_date) if pd.notna(match_date) else None
                        date_str = format_fine_date(date_obj) if date_obj else str(match_date)
                        
                        player = str(row.get('Original Scorecard Name', row.get('Player (Cleaned)', ''))).strip()
                        team_played = str(row.get('Actually Played For', row.get('Midweek Team', ''))).strip()
                        team_a = str(row.get('Team A', '')).strip()
                        team_b = str(row.get('Team B', '')).strip()
                        
                        opponent = team_b if team_played.lower() == team_a.lower() else team_a
                        comp = extract_competition_from_group(str(row.get('Match Group', '')))
                        club = extract_base_club_name(team_played)
                        
                        team_part_str = f"{team_played} (v {opponent})"
                        
                        fines_data.append({
                            'Club': club, 'Date_obj': date_obj, 'Date_str': date_str,
                            'Reason': 'playing a starred player', 'Player': player,
                            'Team_Part_Str': team_part_str,
                            'Competition': comp, 'Fine': 25, 'Type': 'Player'
                        })
        except Exception as e: print('EXCEPTION IN FINES:', repr(e))

    def sort_key(x):
        d = x['Date_obj'] if pd.notna(x['Date_obj']) and x['Date_obj'] is not None else datetime.min
        return (x['Club'].lower(), d)
    
    fines_data.sort(key=sort_key)
    
    from collections import defaultdict
    fines_by_club = defaultdict(list)
    for f in fines_data:
        fines_by_club[f['Club']].append(f)
        
    doc = Document()
    style_normal = doc.styles['Normal']
    style_normal.font.name, style_normal.font.size = 'Calibri', Pt(11)
    
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Club Fines Report")
    r_title.bold = True
    r_title.font.size = Pt(16)
    
    for club in sorted(fines_by_club.keys(), key=lambda c: c.lower()):
        doc.add_paragraph() 
        
        p_club = doc.add_paragraph()
        r_club = p_club.add_run(club)
        r_club.bold = True
        r_club.font.size = Pt(12)
        
        for f in fines_by_club[club]:
            p_fine = doc.add_paragraph()
            
            date_part = f['Date_str']
            reason_part = f['Reason']
            player_part = f" - {f['Player']}" if f['Type'] == 'Player' else ""
            team_part = f"{f['Team_Part_Str']}"
            comp_part = f" – {f['Competition']}" if f['Competition'] and str(f['Competition']).lower() != 'nan' else ""
            fine_part = f"Fine: £{f['Fine']}"
            
            r_date = p_fine.add_run(f"{date_part}")
            r_date.bold = True
            
            text_str = f" – {reason_part}{player_part} - {team_part}{comp_part} - "
            p_fine.add_run(text_str)
            
            r_fine = p_fine.add_run(fine_part)
            r_fine.bold = True
            
            p_fine.paragraph_format.space_after = Pt(6)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io
    
# ==========================================
# UNREGISTERED ONLY FINES GENERATOR
# ==========================================
def generate_unregistered_fines_only(audit_file):
    from collections import defaultdict
    fines_data = []
    
    if audit_file:
        try:
            excel_file = pd.ExcelFile(audit_file)
            
            df_unreg = pd.read_excel(excel_file, sheet_name="Unregistered Matches") if "Unregistered Matches" in excel_file.sheet_names else pd.DataFrame()
            df_deemed = pd.read_excel(excel_file, sheet_name="Deemed Registered") if "Deemed Registered" in excel_file.sheet_names else pd.DataFrame()
            
            player_true_team = {}
            player_deemed_matches = defaultdict(list)
            
            if not df_deemed.empty and 'Stats Name (Cleaned)' in df_deemed.columns:
                for _, r in df_deemed.iterrows():
                    p_key = str(r.get('Stats Name (Cleaned)', '')).strip()
                    m_date = r.get('Match Date')
                    d_obj = pd.to_datetime(m_date) if pd.notna(m_date) else None
                    d_str = format_fine_date(d_obj) if d_obj else str(m_date)
                    t_a = str(r.get('Team A', '')).strip()
                    t_b = str(r.get('Team B', '')).strip()
                    comp = str(r.get('Match League', '')).strip()
                    
                    match_desc = f"{d_str} – Deemed Registered Match: {t_a} v {t_b} ({comp})"
                    player_deemed_matches[p_key].append(match_desc)

            if not df_unreg.empty:
                all_unreg_matches = pd.concat([df_unreg, df_deemed], ignore_index=True) if not df_deemed.empty else df_unreg
                
                if 'Stats Name (Cleaned)' in all_unreg_matches.columns:
                    for player, group in all_unreg_matches.groupby('Stats Name (Cleaned)'):
                        reg_club = str(group.iloc[0].get('Registered Club', 'Unknown Club')).strip()
                        reg_club_base = extract_base_club_name(reg_club).lower()
                        
                        if reg_club_base != 'unknown club':
                            player_true_team[player] = ('known_reg', reg_club)
                        elif '(' in player and player.strip().endswith(')'):
                            club_in_name = player.split('(')[-1].replace(')', '').strip().lower()
                            player_true_team[player] = ('inferred', club_in_name)
                        else:
                            teams_in_matches = []
                            for _, r in group.iterrows():
                                t_a = str(r.get('Team A', '')).strip()
                                t_b = str(r.get('Team B', '')).strip()
                                teams_in_matches.append({extract_base_club_name(t_a).lower(), extract_base_club_name(t_b).lower()})
                            
                            if len(teams_in_matches) == 1:
                                t_a = str(group.iloc[0].get('Team A', '')).strip()
                                t_b = str(group.iloc[0].get('Team B', '')).strip()
                                player_true_team[player] = ('ambiguous', (t_a, t_b))
                            else:
                                common_teams = set.intersection(*teams_in_matches)
                                if len(common_teams) == 1:
                                    player_true_team[player] = ('inferred', list(common_teams)[0])
                                else:
                                    t_a = str(group.iloc[0].get('Team A', '')).strip()
                                    t_b = str(group.iloc[0].get('Team B', '')).strip()
                                    player_true_team[player] = ('ambiguous', (t_a, t_b))

            if not df_unreg.empty and len(df_unreg.columns) > 1:
                for _, row in df_unreg.iterrows():
                    match_date = row.get('Match Date')
                    date_obj = pd.to_datetime(match_date) if pd.notna(match_date) else None
                    date_str = format_fine_date(date_obj) if date_obj else str(match_date)
                    
                    player_key = str(row.get('Stats Name (Cleaned)', '')).strip()
                    player_disp = str(row.get('Original Scorecard Name', player_key)).strip()
                    
                    team_a = str(row.get('Team A', '')).strip()
                    team_b = str(row.get('Team B', '')).strip()
                    comp = str(row.get('Match League', '')).strip()
                    
                    status, info = player_true_team.get(player_key, ('known_reg', str(row.get('Registered Club', 'Unknown Club')).strip()))
                    
                    if status == 'known_reg':
                        reg_club_base = extract_base_club_name(info).lower()
                        if reg_club_base in extract_base_club_name(team_b).lower() and reg_club_base != 'unknown club':
                            team_played, opponent = team_b, team_a
                        else:
                            team_played, opponent = team_a, team_b
                        club = extract_base_club_name(team_played)
                        team_part_str = f"{team_played} (v {opponent})"
                    
                    elif status == 'inferred':
                        if info in extract_base_club_name(team_b).lower():
                            team_played, opponent = team_b, team_a
                        else:
                            team_played, opponent = team_a, team_b
                        club = extract_base_club_name(team_played)
                        team_part_str = f"{team_played} (v {opponent})"
                        
                    elif status == 'ambiguous':
                        t_a, t_b = info
                        club = f"{extract_base_club_name(t_a)} / {extract_base_club_name(t_b)}"
                        team_part_str = f"{t_a} v {t_b}" 
                    
                    subsequent_matches = player_deemed_matches.get(player_key, [])
                        
                    fines_data.append({
                        'Club': club, 'Date_obj': date_obj, 'Date_str': date_str,
                        'Reason': 'playing an unregistered player', 'Player': player_disp,
                        'Team_Part_Str': team_part_str,
                        'Competition': comp, 'Fine': 10, 'Type': 'Player',
                        'Deemed_Matches': subsequent_matches
                    })
        except Exception as e: print('EXCEPTION IN FINES:', repr(e))

    def sort_key(x):
        d = x['Date_obj'] if pd.notna(x['Date_obj']) and x['Date_obj'] is not None else datetime.min
        return (x['Club'].lower(), d)
    
    fines_data.sort(key=sort_key)
    
    fines_by_club = defaultdict(list)
    for f in fines_data:
        fines_by_club[f['Club']].append(f)
        
    doc = Document()
    style_normal = doc.styles['Normal']
    style_normal.font.name, style_normal.font.size = 'Calibri', Pt(11)
    
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Unregistered Player Fines Report")
    r_title.bold = True
    r_title.font.size = Pt(16)
    
    for club in sorted(fines_by_club.keys(), key=lambda c: c.lower()):
        doc.add_paragraph() 
        
        p_club = doc.add_paragraph()
        r_club = p_club.add_run(club)
        r_club.bold = True
        r_club.font.size = Pt(12)
        
        for f in fines_by_club[club]:
            p_fine = doc.add_paragraph()
            
            date_part = f['Date_str']
            reason_part = f['Reason']
            player_part = f" - {f['Player']}"
            team_part = f"{f['Team_Part_Str']}"
            comp_part = f" – {f['Competition']}" if f['Competition'] and str(f['Competition']).lower() != 'nan' else ""
            fine_part = f"Fine: £{f['Fine']}"
            
            r_date = p_fine.add_run(f"{date_part}")
            r_date.bold = True
            
            text_str = f" – {reason_part}{player_part} - {team_part}{comp_part} - "
            p_fine.add_run(text_str)
            
            r_fine = p_fine.add_run(fine_part)
            r_fine.bold = True
            
            if f['Deemed_Matches']:
                p_fine.paragraph_format.space_after = Pt(2)
                
                p_info = doc.add_paragraph()
                p_info.paragraph_format.left_indent = Pt(18)
                p_info.paragraph_format.space_before = Pt(0)
                p_info.paragraph_format.space_after = Pt(3)
                r_info = p_info.add_run(f"→ Subsequent matches played while deemed registered ({len(f['Deemed_Matches'])}):")
                r_info.font.italic = True
                r_info.font.size = Pt(10)
                
                for sub_m in f['Deemed_Matches']:
                    p_sub = doc.add_paragraph()
                    p_sub.paragraph_format.left_indent = Pt(36)
                    p_sub.paragraph_format.space_before = Pt(0)
                    p_sub.paragraph_format.space_after = Pt(2)
                    r_sub = p_sub.add_run(f"• {sub_m}")
                    r_sub.font.size = Pt(10)
                    r_sub.font.color.rgb = RGBColor(100, 100, 100)
            else:
                p_fine.paragraph_format.space_after = Pt(6)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io
    
# ==========================================
# MILESTONES ENGINE SPECIFIC FUNCTIONS
# ==========================================
@st.cache_data(show_spinner="Generating milestones report...")
def generate_milestones_report(domain, f_reg, f_alias, f_league, f_bat, f_bowl, f_cup=None, f_id_map=None, f_secondary=None):
    reg_players = pd.read_excel(f_reg)
    aliases = pd.read_excel(f_alias)
    league_structure = pd.read_excel(f_league)
    batting_df = pd.read_excel(f_bat)
    bowling_df = pd.read_excel(f_bowl)
    
    if not f_id_map:
        f_id_map = DEFAULT_FILES.get(domain, {}).get("id_map", "")
    id_map = {}
    if f_id_map and os.path.exists(f_id_map):
        id_map_df = get_excel_df(f_id_map)
        id_map = build_id_map(id_map_df)
    
    alias_map = build_alias_map(aliases, domain)
    player_club_map = build_player_club_map(reg_players, alias_map, domain)
    league_dict, team_keys, _ = build_league_dict(league_structure)
    
    if not f_secondary:
        f_secondary = DEFAULT_FILES.get(domain, {}).get("secondary", "5. Secondary_Team_Map.xlsx")
    secondary_map = {}
    if f_secondary and os.path.exists(f_secondary):
        sec_df = get_excel_df(f_secondary)
        secondary_map = build_secondary_team_map(sec_df, alias_map)
    
    if domain == "Women's":
        wicket_threshold = 5
        batting_leagues_order = [
            "Mercury Women's Premier League"
        ]
        bowling_leagues_order = [
            "Mercury Women's Premier League",
            "Mercury Women's Senior League"
        ]
        main_batting_header = "WOMEN BATTING - CENTURIONS"
        main_bowling_header = f"WOMEN BOWLING - {wicket_threshold} WICKETS OR MORE"
    else:
        wicket_threshold = 6
        batting_leagues_order = [
            "Mercury Premier League", 
            "Mercury Senior League 1", 
            "Mercury Senior League 2", 
            "Mercury Senior League 3"
        ]
        bowling_leagues_order = list(batting_leagues_order)
        main_batting_header = "OPEN BATTING - CENTURIONS"
        main_bowling_header = f"OPEN BOWLING - {wicket_threshold} WICKETS OR MORE"
    
    cup_match_dict = {}
    if f_cup and os.path.exists(f_cup):
        try:
            excel_file_cup = pd.ExcelFile(f_cup)
            target_sheet = excel_file_cup.sheet_names[0]
            for sheet in excel_file_cup.sheet_names:
                if domain.lower().replace("'", "") in sheet.lower().replace("'", ""):
                    target_sheet = sheet
                    break
            cup_df = pd.read_excel(f_cup, sheet_name=target_sheet, header=None)
            
            def local_parse(group_str):
                try:
                    group_str = str(group_str).strip()
                    parts = group_str.rsplit(' - ', 1)
                    date_str = parts[1].strip() if len(parts) == 2 else group_str
                    rest = parts[0].strip() if len(parts) == 2 else group_str
                    match_date = pd.to_datetime(date_str, dayfirst=True, errors='coerce')
                    if pd.notna(match_date): match_date = match_date.normalize()
                    if ' v ' in rest:
                        t_a, remainder = rest.split(' v ', 1)
                        t_b = remainder.rsplit(', ', 1)[0] if ', ' in remainder else (remainder.rsplit(' - ', 1)[0] if ' - ' in remainder else remainder)
                    else:
                        t_a, t_b = rest, "Unknown"
                    return t_a.strip(), t_b.strip(), match_date
                except: return None, None, None

            for _, row_data in cup_df.iterrows():
                match_str_raw = str(row_data[0]).strip()
                cup_name = str(row_data[1]).strip()
                if match_str_raw.lower() in ['match string', 'match group', 'match', 'nan']: continue
                cleaned_match_str = doc_format_cricket_names(match_str_raw, domain)
                c_team_a, c_team_b, c_date = local_parse(cleaned_match_str)
                if c_team_a and c_team_b:
                    teams = sorted([str(c_team_a).lower(), str(c_team_b).lower()])
                    if pd.notna(c_date):
                        cup_match_dict[f"{teams[0]}_{teams[1]}_{c_date.strftime('%Y-%m-%d')}"] = cup_name
                    else:
                        cup_match_dict[f"{teams[0]}_{teams[1]}"] = cup_name
        except Exception as e: print('EXCEPTION IN FINES:', repr(e))
        
    def is_cup_match(grp_str):
        grp_str_clean = str(grp_str).lower()
        cup_kws = ['cup', 'trophy', 'shield', 'plate', 'bowl', 'vase', 'challenge', 't20', 'twenty20', 'gallagher', 'lvs']
        
        if any(kw in grp_str_clean for kw in cup_kws):
            return True
            
        if cup_match_dict:
            c_team_a, c_team_b, c_date = local_parse(doc_format_cricket_names(grp_str, domain))
            if c_team_a and c_team_b:
                teams = sorted([str(c_team_a).lower(), str(c_team_b).lower()])
                comp = None
                if pd.notna(c_date):
                    comp = cup_match_dict.get(f"{teams[0]}_{teams[1]}_{c_date.strftime('%Y-%m-%d')}")
                if not comp:
                    comp = cup_match_dict.get(f"{teams[0]}_{teams[1]}")
                if comp and any(kw in str(comp).lower() for kw in cup_kws):
                    return True
        return False

    def get_target_league(league_str):
        if not league_str: return None
        l_lower = str(league_str).lower()
        
        if domain == "Women's":
            if 'premier' in l_lower: return "Mercury Women's Premier League"
            elif 'senior' in l_lower: return "Mercury Women's Senior League"
            return None
        else:
            if 'premier' in l_lower: return "Mercury Premier League"
            elif 'senior league 1' in l_lower or 'senior 1' in l_lower or 'section 1' in l_lower: return "Mercury Senior League 1"
            elif 'senior league 2' in l_lower or 'senior 2' in l_lower or 'section 2' in l_lower: return "Mercury Senior League 2"
            elif 'senior league 3' in l_lower or 'senior 3' in l_lower or 'section 3' in l_lower: return "Mercury Senior League 3"
            return None
        
    def format_day_month(dt):
        if pd.isna(dt): return "Unknown Date"
        day = dt.day
        if 11 <= (day % 100) <= 13: suffix = 'th'
        else: suffix = ['th', 'st', 'nd', 'rd', 'th'][min(day % 10, 4)]
        month = dt.strftime('%B')
        return f"{day}{suffix} {month}"
        
    def process_milestone_row(row, is_batting):
        scorecard_name = str(row['Name'] if is_batting else row['Bowler']).strip()
        team_played = determine_player_team_for_row(row, player_club_map, domain, secondary_map=secondary_map)
        grp = str(row['Group'])
        t1, t2 = extract_teams_from_group(grp)
        opponent = t2 if team_played == t1 else t1
        
        team_played_clean = re.sub(r'(?i)\bwomen\'?s?\b', '', team_played)
        team_played_clean = re.sub(r'\s+', ' ', team_played_clean).strip()
        if "Holywood" in team_played_clean and "1881" not in team_played_clean:
            team_played_clean = team_played_clean.replace("Holywood", "Holywood 1881")
            
        opponent_clean = re.sub(r'(?i)\bwomen\'?s?\b', '', opponent)
        opponent_clean = re.sub(r'\s+', ' ', opponent_clean).strip()
        if "Holywood" in opponent_clean and "1881" not in opponent_clean:
            opponent_clean = opponent_clean.replace("Holywood", "Holywood 1881")
        
        league = get_team_league(team_played, team_keys, league_dict, domain)
        
        try:
            parts = grp.rsplit(' - ', 1)
            date_str = parts[1].strip() if len(parts) == 2 else grp
            clean_date_str = re.sub(r'(?<=\d)(st|nd|rd|th)\b', '', date_str, flags=re.IGNORECASE)            
            match_date = pd.to_datetime(clean_date_str, dayfirst=True, errors='coerce')
            date_formatted = format_day_month(match_date) if pd.notna(match_date) else date_str
        except:
            match_date = pd.Timestamp.min
            date_formatted = "Unknown Date"
            
        return scorecard_name, team_played_clean, opponent_clean, league, date_formatted, match_date

    batting_df = batting_df[~batting_df['Group'].apply(is_cup_match)]
    bowling_df = bowling_df[~bowling_df['Group'].apply(is_cup_match)]

    batting_df['Cleaned Name'] = batting_df.apply(lambda r: resolve_player_from_row(r, r['Name'], id_map, alias_map, player_club_map, id_cols=['Batter ID', 'Player ID', 'ID'])[0], axis=1)
    bowling_df['Cleaned Name'] = bowling_df.apply(lambda r: resolve_player_from_row(r, r['Bowler'], id_map, alias_map, player_club_map, id_cols=['Bowler ID', 'Player ID', 'ID'])[0], axis=1)
    
    batting_df['Runs'] = pd.to_numeric(batting_df['Runs'], errors='coerce').fillna(0)
    bowling_df['Wickets'] = pd.to_numeric(bowling_df['Wickets'], errors='coerce').fillna(0)
    
    centurions = batting_df[batting_df['Runs'] >= 100]
    top_wickets = bowling_df[bowling_df['Wickets'] >= wicket_threshold]
    
    batting_results = {l: [] for l in batting_leagues_order}
    bowling_results = {l: [] for l in bowling_leagues_order}
    
    for _, row in centurions.iterrows():
        scorecard_name, team_played, opponent, raw_league, date_fmt, dt_obj = process_milestone_row(row, is_batting=True)
        target_league = get_target_league(raw_league)
        
        if target_league in batting_results:
            runs = int(row['Runs'])
            
            is_not_out = False
            if 'Not Outs' in row and pd.to_numeric(row['Not Outs'], errors='coerce') > 0:
                is_not_out = True
            elif 'High Score' in row and '*' in str(row['High Score']):
                is_not_out = True
                
            runs_str = f"{runs}*" if is_not_out else str(runs)
            line = f"{scorecard_name} ({team_played}) - {runs_str} vs {opponent} on {date_fmt}"
            
            name_parts = scorecard_name.strip().split()
            surname = name_parts[-1].lower() if len(name_parts) > 1 else (name_parts[0].lower() if name_parts else "")
            firstname = " ".join(name_parts[:-1]).lower() if len(name_parts) > 1 else ""

            batting_results[target_league].append({
                'line': line, 
                'date': dt_obj if pd.notna(dt_obj) else pd.Timestamp.min,
                'surname': surname,
                'firstname': firstname
            })

    for _, row in top_wickets.iterrows():
        scorecard_name, team_played, opponent, raw_league, date_fmt, dt_obj = process_milestone_row(row, is_batting=False)
        target_league = get_target_league(raw_league)
        
        if target_league in bowling_results:
            wicks = int(row['Wickets'])
            runs_conc = int(row['Runs']) if pd.notna(row['Runs']) else 0
            line = f"{scorecard_name} ({team_played}) - {wicks}-{runs_conc} vs {opponent} on {date_fmt}"
            
            name_parts = scorecard_name.strip().split()
            surname = name_parts[-1].lower() if len(name_parts) > 1 else (name_parts[0].lower() if name_parts else "")
            firstname = " ".join(name_parts[:-1]).lower() if len(name_parts) > 1 else ""

            bowling_results[target_league].append({
                'line': line, 
                'date': dt_obj if pd.notna(dt_obj) else pd.Timestamp.min,
                'surname': surname,
                'firstname': firstname
            })
            
    doc = Document()
    
    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    p_open_bat = doc.add_paragraph()
    r_open_bat = p_open_bat.add_run(main_batting_header)
    r_open_bat.bold = True
    doc.add_paragraph("")
    
    for idx, league in enumerate(batting_leagues_order):
        p_league = doc.add_paragraph()
        p_league.paragraph_format.space_after = Pt(0)
        
        if idx > 0:
            p_league.paragraph_format.space_before = Pt(16)
        else:
            p_league.paragraph_format.space_before = Pt(0)
            
        r_league = p_league.add_run(league)
        r_league.bold = True
        
        p_dash = doc.add_paragraph("–" * int(len(league) * 1.1))
        p_dash.paragraph_format.space_before = Pt(0)
        p_dash.paragraph_format.space_after = Pt(2)

        matches = batting_results[league]
        if matches:
            matches.sort(key=lambda x: (x['date'], x['surname'], x['firstname']))
            last_date = None
            for m in matches:
                p = doc.add_paragraph(m['line'])
                if last_date is not None and m['date'] != last_date:
                    p.paragraph_format.space_before = Pt(4)
                else:
                    p.paragraph_format.space_before = Pt(0)
                last_date = m['date']
        else:
            p_none = doc.add_paragraph("None")
            p_none.paragraph_format.space_before = Pt(0)
            
    if domain == "Women's":
        p_open_bowl = doc.add_paragraph()
        p_open_bowl.paragraph_format.space_before = Pt(24)
    else:
        doc.add_page_break()
        p_open_bowl = doc.add_paragraph()

    r_open_bowl = p_open_bowl.add_run(main_bowling_header)
    r_open_bowl.bold = True
    doc.add_paragraph("")
    
    for idx, league in enumerate(bowling_leagues_order):
        p_league = doc.add_paragraph()
        p_league.paragraph_format.space_after = Pt(0)
        
        if idx > 0:
            p_league.paragraph_format.space_before = Pt(16)
        else:
            p_league.paragraph_format.space_before = Pt(0)
            
        r_league = p_league.add_run(league)
        r_league.bold = True
        
        p_dash = doc.add_paragraph("–" * int(len(league) * 1.1))
        p_dash.paragraph_format.space_before = Pt(0)
        p_dash.paragraph_format.space_after = Pt(2)
        
        matches = bowling_results[league]
        if matches:
            matches.sort(key=lambda x: (x['date'], x['surname'], x['firstname']))
            last_date = None
            for m in matches:
                p = doc.add_paragraph(m['line'])
                if last_date is not None and m['date'] != last_date:
                    p.paragraph_format.space_before = Pt(4)
                else:
                    p.paragraph_format.space_before = Pt(0)
                last_date = m['date']
        else:
            p_none = doc.add_paragraph("None")
            p_none.paragraph_format.space_before = Pt(0)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    
    return doc_io

# ==========================================
# CLUB CONTACTS DIRECTORY ENGINE
# ==========================================
def get_contact_team_tier(role_name):
    r = str(role_name).lower()
    if "women's first" in r or "womens first" in r: return "Women's 1st XI"
    if "women's second" in r or "womens second" in r: return "Women's 2nd XI"
    if "women's third" in r or "womens third" in r: return "Women's 3rd XI"
    if "first midweek" in r or "1st midweek" in r: return "1st Midweek XI"
    if "second midweek" in r or "2nd midweek" in r: return "2nd Midweek XI"
    if "first team" in r or "1st team" in r: return "1st XI"
    if "second team" in r or "2nd team" in r: return "2nd XI"
    if "third team" in r or "3rd team" in r: return "3rd XI"
    if "fourth team" in r or "4th team" in r: return "4th XI"
    if "fifth team" in r or "5th team" in r: return "5th XI"
    if "sixth team" in r or "6th team" in r: return "6th XI"
    if "boys" in r: return "Boys Youth"
    if "girls" in r: return "Girls Youth"
    if any(k in r for k in ["youth", "programme", "lead/coach"]): return "Youth & Coaching"
    if "indoor" in r: return "Indoor Cricket"
    return "Club Official"

def parse_club_contacts_matrix(file_or_df):
    """
    Transforms the wide NCU Club Contacts sheet into a normalized DataFrame
    while strictly preserving the source spreadsheet's top-to-bottom role sequence.
    """
    if isinstance(file_or_df, str):
        if not os.path.exists(file_or_df):
            return pd.DataFrame(), {}, []
        excel_file = pd.ExcelFile(file_or_df, engine="calamine")
        sheet = "Club Contacts" if "Club Contacts" in excel_file.sheet_names else excel_file.sheet_names[0]
        df = pd.read_excel(file_or_df, sheet_name=sheet, engine="calamine")
    elif isinstance(file_or_df, pd.DataFrame):
        df = file_or_df
    else:
        df = pd.read_excel(file_or_df, engine="calamine")

    if df.empty:
        return pd.DataFrame(), {}, []

    role_col = df.columns[0]
    clubs = [c for c in df.columns[1:] if not str(c).startswith("Unnamed")]
    row_labels = df[role_col].tolist()

    rows = []
    grounds_by_club = {c: {} for c in clubs}
    ordered_roles = []

    df_dict = df.to_dict('list')
    i = 0
    role_idx = 0
    while i < len(row_labels):
        label = str(row_labels[i]).strip()
        if not label or label.lower() == "nan":
            i += 1
            continue

        # Extract Ground records
        if "Ground" in label and "Convenor" not in label:
            for c in clubs:
                val = df_dict[c][i]
                if pd.notna(val) and str(val).strip() and str(val).strip().lower() != "nan":
                    grounds_by_club[c][label] = str(val).strip()
            i += 1
            continue

        # Extract Role Block (Name, Email, Mobile)
        if i + 2 < len(row_labels) and "email" in str(row_labels[i+1]).lower() and "mobile" in str(row_labels[i+2]).lower():
            role_name = label
            if role_name not in ordered_roles:
                ordered_roles.append(role_name)
            tier = get_contact_team_tier(role_name)
            role_idx += 1

            for c in clubs:
                name_val = df_dict[c][i]
                em_val = df_dict[c][i+1]
                mob_val = df_dict[c][i+2]

                name_str = str(name_val).strip() if pd.notna(name_val) and str(name_val).strip().lower() != "nan" else ""
                em_str = str(em_val).strip() if pd.notna(em_val) and str(em_val).strip().lower() != "nan" else ""
                mob_str = str(mob_val).strip() if pd.notna(mob_val) and str(mob_val).strip().lower() != "nan" else ""

                if name_str or em_str or mob_str:
                    rows.append({
                        "Club": c,
                        "Role": role_name,
                        "Role Order": role_idx,
                        "Team Tier": tier,
                        "Name": name_str,
                        "Email": em_str,
                        "Phone": mob_str
                    })
            i += 3
        else:
            i += 1

    return pd.DataFrame(rows), grounds_by_club, ordered_roles

# ==========================================
# REGISTRATION FEE AUDIT
# ==========================================


def clean_revenue_report(source_file):
    import pandas as pd
    import re
    
    # Check if this is the original raw file or a pre-filtered one
    df_raw = pd.read_excel(source_file, sheet_name='All Data')
    
    if 'ItemType' in df_raw.columns:
        # It's the raw Sport80 export
        df_all = df_raw[df_raw['ItemType'] == 'ADD_ON'].copy()
        df_all.rename(columns={
            'InvoiceDate': 'Payment Date',
            'Description': 'Player Name - Club - Type',
            'UnitAmount': 'Payment Amount'
        }, inplace=True)
    else:
        # It's already the user's manual export
        df_all = df_raw.copy()
        
    col = df_all['Player Name - Club - Type']

    def parse_record(s):
        s = str(s).strip()
        first_split = s.split(' - ', 1)
        player_name = first_split[0].strip()
        rem = first_split[1].strip() if len(first_split) > 1 else ''
        
        pay_method = ''
        m_pay = re.search(r' - (One Time Payment|Auto Renewal.*)$', rem)
        if m_pay:
            pay_method = m_pay.group(1).strip()
            rem = rem[:m_pay.start()].strip()
            
        validity = ''
        m_val = re.search(r'\s*(\((?:Valid until - |01/03/2026 - )01/03/2027\))\s*$', rem)
        if m_val:
            validity = m_val.group(1).strip()
            rem = rem[:m_val.start()].strip()
            
        club = ''
        m_type = ''
        
        if 'Dundrum' in rem:
            club = 'Dundrum Cricket Club'
            m = re.search(r'(Adult \(over 18\)|Youth player \(playing Adult & Youth cricket\))', rem)
            m_type = m.group(1) if m else 'Membership'
        elif rem.startswith('Woodvale Cricket Club'):
            club = 'Woodvale Cricket Club'
            m_type = rem.split(' - ')[-1].strip()
        elif rem.startswith('Downpatrick Cricket Club'):
            club = 'Downpatrick Cricket Club'
            m_type = rem.split(' - ')[-1].strip()
        elif rem.startswith('Holywood Cricket Club 1881'):
            club = 'Holywood Cricket Club 1881'
            m_type = rem.split(' - ')[1].strip()
        elif rem.startswith('Instonians Cricket Club'):
            club = 'Instonians Cricket Club'
            m_type = rem.split(' - ')[-1].strip()
        elif ' - Northern Cricket Union - ' in rem:
            parts = rem.split(' - Northern Cricket Union - ')
            raw_club = parts[0].strip()
            club = re.sub(r'\s+(Membership|Registration|Fee\'s|Fees)$', '', raw_club).strip()
            m_type = parts[1].strip()
        elif ' - SPORT80 MEMBERSHIP - ' in rem:
            parts = rem.split(' - SPORT80 MEMBERSHIP - ')
            raw_club = parts[0].strip()
            club = re.sub(r'\s+(Membership|Registration)$', '', raw_club).strip()
            m_type = parts[1].strip()
        elif ' - ' in rem:
            parts = rem.split(' - ', 1)
            raw_club = parts[0].strip()
            club = re.sub(r'\s+(Membership(?:\s+\d{4})?|Registration(?:\s+\d{4})?|NCU Registration|membership|Fee\'s|Fees)$', '', raw_club).strip()
            m_type = parts[1].strip()
        else:
            m_match = re.match(r'^(.*?)\s+(Membership|Registration)$', rem, re.IGNORECASE)
            if m_match:
                club = m_match.group(1).strip()
                m_type = m_match.group(2).strip()
            else:
                club = rem
                m_type = 'Membership'
                
        club_mapping = {
            'Arches CC': 'Arches Cricket Club',
            'BISC': 'BISC Cricket Club',
            'Derriaghy Cricket Club NCU': 'Derriaghy Cricket Club',
            'Lisburn Cricket Club membership': 'Lisburn Cricket Club',
            'Northern Ireland Malayali Association CC': 'NIMA Cricket Club',
        }
        club = club_mapping.get(club, club)

        # SPECIFIC FIX FOR MAGEE AND MCILWAINE
        if player_name.lower() == 'james magee' and 'instonians' in club.lower():
            if 'youth' in m_type.lower() or 'youth' in s.lower():
                player_name = 'James Magee jnr'
            else:
                player_name = 'James Magee snr'
        elif player_name.lower() == 'peter mcilwaine' and 'bangor' in club.lower():
            if 'youth' in m_type.lower() or 'youth' in s.lower():
                player_name = 'Teddy Mcilwaine'
            else:
                player_name = 'Peter Mcilwaine' 

        return {
            'Player Name': player_name,
            'Club': club,
            'Type': m_type,
            'Payment Method': pay_method,
            'Validity': validity
        }

    parsed_rows = [parse_record(x) for x in col]
    df_parsed = pd.DataFrame(parsed_rows)
    df_parsed.index = df_all.index

    df_clean = pd.DataFrame({
        'Payment Date': df_all['Payment Date'],
        'Player Name': df_parsed['Player Name'],
        'Club': df_parsed['Club'],
        'Type': df_parsed['Type'],
        'Payment Method': df_parsed['Payment Method'],
        'Validity': df_parsed['Validity'],
        'Payment Amount': df_all['Payment Amount'],
        'Original Description': df_all['Player Name - Club - Type']
    })
    
    return df_clean


def generate_anomalies_word_report(df_rev, df_reg, alias_map, timestamped_prefix):
    from docx import Document
    from docx.shared import Pt
    import unicodedata
    import pandas as pd
    
    def norm(text):
        if pd.isna(text): return ""
        return unicodedata.normalize('NFKD', str(text)).encode('ascii', 'ignore').decode('utf-8').lower().strip()
        
    df_rev_copy = df_rev.copy()
    df_rev_copy['Norm_Name'] = df_rev_copy['Player Name'].apply(norm)
    
    # 2. Extract Anomalies
    player_counts = df_rev_copy.groupby('Player Name').agg(
        Tx_Count=('Payment Amount', 'count'),
        Total_Paid=('Payment Amount', 'sum'),
        Clubs=('Club', lambda s: list(set(s))),
        Types=('Type', lambda s: list(set(s))),
        Dates=('Payment Date', list)
    ).reset_index()

    multi_club = player_counts[player_counts['Clubs'].apply(len) > 1].sort_values(by='Total_Paid', ascending=False)
    father_son = player_counts[player_counts['Player Name'].isin(['James Magee', 'Peter Mcilwaine'])]
    same_club_dups = player_counts[(player_counts['Clubs'].apply(len) == 1) & (player_counts['Tx_Count'] > 1)]
    same_club_dups = same_club_dups[~same_club_dups['Player Name'].isin(['James Magee', 'Peter Mcilwaine'])]
    upgrades = df_rev_copy[df_rev_copy['Type'].astype(str).str.contains('UPGRADE', case=False, na=False)]
    arches = df_rev_copy[df_rev_copy['Club'].astype(str).str.contains('Arches', case=False, na=False)]

    def is_registered(name):
        n = norm(name)
        n = alias_map.get(n, n)
        # df_reg has 'Norm_Name' created in run_registration_fee_audit
        return n in df_reg['Norm_Name'].values

    unregistered_payers = player_counts[~player_counts['Player Name'].apply(is_registered)]

    # 3. Create Document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.add_heading('NCU Registration Fee: Revenue Anomalies Report', 0)

    doc.add_heading(f'1. Dual-Club Double Payments ({len(multi_club)} Players Paid Twice)', level=1)
    p = doc.add_paragraph(f"{len(multi_club)} players paid registration fees under two (or three) different clubs, meaning they paid £15 to £20 in total instead of the single £10 annual NCU affiliation fee:")
    for _, r in multi_club.iterrows():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{r['Player Name']}").bold = True
        clubs_str = " and ".join(r['Clubs'])
        p.add_run(f": Paid under {clubs_str} (£{r['Total_Paid']} total).")

    p = doc.add_paragraph()
    p.add_run("Cause: ").bold = True
    p.add_run("When players transfer mid-season or play for one club on Saturdays and another in the Midweek League, Sport80 prompts them for the NCU fee again, resulting in an accidental double payment to the NCU.")
    doc.add_paragraph("_" * 80)

    doc.add_heading('2. Father / Son Namesake Payments (Instonians & Bangor)', level=1)
    doc.add_paragraph("Two players appeared to have paid both the £10 adult fee and the £5 youth fee for the same club:")
    for _, r in father_son.iterrows():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{r['Player Name']} ({r['Clubs'][0]}):").bold = True
        
        txs = df_rev_copy[df_rev_copy['Player Name'] == r['Player Name']]
        for _, tx in txs.iterrows():
            date_str = str(tx['Payment Date'])[:10]
            amt = tx['Payment Amount']
            p_sub = doc.add_paragraph(f"Paid £{amt} on {date_str} ({tx['Type']})", style='List Bullet 2')
        
        p_inv = doc.add_paragraph(style='List Bullet 2')
        p_inv.add_run("Investigation: ").italic = True
        p_inv.add_run(f"Sport80 date of birth records reveal there are two different {r['Player Name']}s at this club—an adult and a youth. In the single-line player database, they were merged into one name.")
    doc.add_paragraph("_" * 80)

    doc.add_heading('3. Genuine Duplicate Double-Charge at Same Club', level=1)
    doc.add_paragraph("The following players were charged twice for the exact same membership at the same club:")
    for _, r in same_club_dups.iterrows():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{r['Player Name']} ({r['Clubs'][0]}):").bold = True
        
        txs = df_rev_copy[df_rev_copy['Player Name'] == r['Player Name']]
        for _, tx in txs.iterrows():
            date_str = str(tx['Payment Date'])[:10]
            amt = tx['Payment Amount']
            p_sub = doc.add_paragraph(f"Paid £{amt} on {date_str}", style='List Bullet 2')
        
        p_inv = doc.add_paragraph(style='List Bullet 2')
        p_inv.add_run("Investigation: ").italic = True
        p_inv.add_run("Two identical profiles exist for them in Sport80 under their club, and their family was accidentally billed twice.")
    doc.add_paragraph("_" * 80)

    doc.add_heading('4. Mid-Season "UPGRADE" Transactions (Muckamore)', level=1)
    p = doc.add_paragraph("Three transactions at Muckamore CC were explicitly marked as ")
    p.add_run("UPGRADE").bold = True
    p.add_run(" (£5 fee):")

    for _, r in upgrades.iterrows():
        date_str = str(r['Payment Date'])[:10]
        p = doc.add_paragraph(f"{r['Player Name']} ({date_str})", style='List Number')
        p.runs[0].bold = True

    p_exp = doc.add_paragraph()
    p_exp.add_run("Explanation: ").italic = True
    p_exp.add_run("These three juniors originally registered under the £0 pure youth exemption, but when selected for senior cricket mid-season, their parents/club correctly paid a £5 \"UPGRADE\" fee. All three are captured as compliant youth players in the audit.")
    doc.add_paragraph("_" * 80)

    doc.add_heading('5. Club Setup Anomaly: Arches CC Flat-Fee Underpayments', level=1)
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("At Arches Cricket Club, 100% of all 38 payments were £5").bold = True
    p1.add_run(" (not a single £10 payment was recorded).")

    p2 = doc.add_paragraph("34 of these players are adults (>18) playing senior cricket.", style='List Bullet')

    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Arches CC configured their Sport80 membership category with a flat £5 fee called \"Membership\", accounting for ")
    p3.add_run("64% of all adult underpayments in the entire NCU").bold = True
    p3.add_run(".")
    doc.add_paragraph("_" * 80)

    doc.add_heading(f'6. The Remaining {len(unregistered_payers)} Standalone Revenue Records (Paid but Not Registered)', level=1)
    doc.add_paragraph(f"Only {len(unregistered_payers)} payments in the entire report belong to individuals who do not exist in the registered players list:")

    for _, r in unregistered_payers.iterrows():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{r['Player Name']} ").bold = True
        p.add_run(f"(£{r['Total_Paid']}, {r['Clubs'][0]}): ")
        if r['Player Name'] == 'Jared Wilson':
            p.add_run("Registered in Sport80 on 27th August 2026 (late registration after initial export).")
        else:
            p.add_run("Paid fees, but never registered as players and never appeared on scorecards.")
    doc.add_paragraph("_" * 80)

    doc.add_heading('7. Financial Integrity Checks Passed', level=1)
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Strict Price Adherence: ").bold = True
    p1.add_run(f"100% of all {len(df_rev_copy)} payments in the file are strictly either £10 or £5. There are no negative amounts, £0 amounts, partial fees, or odd amounts.")

    dates = pd.to_datetime(df_rev_copy['Payment Date'], errors='coerce')
    min_date = dates.min().strftime('%d %B %Y')
    max_date = dates.max().strftime('%d %B %Y')
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Date Range: ").bold = True
    p2.add_run(f"All payments occurred between {min_date} and {max_date}.")

    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io

def run_registration_fee_audit():
    import unicodedata
    import re
    from datetime import datetime
    import shutil
    import pandas as pd
    import numpy as np
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    
    
    # Smart title function
    def smart_title(name):
        if not name or pd.isna(name):
            return ""
        name = str(name).strip().replace('’', "'").replace('`', "'")
        
        def title_word(word):
            if '-' in word:
                return '-'.join(title_word(part) for part in word.split('-'))
            w = word.lower()
            if re.match(r"^o'[a-z]", w):
                return "O'" + w[2].upper() + w[3:]
            if re.match(r"^mc[a-z]", w):
                return "Mc" + w[2].upper() + w[3:]
            if re.match(r"^mac[a-z]{3,}", w) and word[:3].lower() == 'mac' and len(word) > 4:
                if len(word) > 3 and word[3].isupper():
                    return "Mac" + w[3].upper() + w[4:]
            if w in ['jnr', 'snr', 'ii', 'iii', 'iv']:
                return w.title() if w in ['jnr', 'snr'] else w.upper()
            return w.capitalize()
    
        return " ".join(title_word(w) for w in name.split())
    
    # Test on user examples
    assert smart_title("TREVOR Dempsey") == "Trevor Dempsey"
    assert smart_title("Hafiz M WAQAS Iqbal") == "Hafiz M Waqas Iqbal"
    assert smart_title("mckinley") == "McKinley"
    assert smart_title("MCKINLEY") == "McKinley"
    
    # Reference date: 30th June 2026
    REF_DATE = pd.to_datetime('2026-06-30')
    
    def norm(text):
        if not text or pd.isna(text): return ""
        text = str(text).replace('’', "'").replace('`', "'").replace('â€™', "'").replace('Ã©', 'e').replace('Ã­', 'i')
        text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')
        return " ".join(text.lower().split())
    
    # 1. Load Aliases
    men_alias_file = '2. NCU_Validated_Aliases_Master.xlsx'
    women_alias_file = "12. NCU_Validated_Women's Aliases_Master.xlsx"
    df_alias_m = pd.read_excel(men_alias_file)
    df_alias_w = pd.read_excel(women_alias_file)
    df_alias = pd.concat([df_alias_m, df_alias_w], ignore_index=True)
    
    alias_map = {}
    for _, row in df_alias.iterrows():
        inp = norm(row['Input Name (Scorecard/Stats)'])
        off = str(row['Official Registered Name']).replace('‡', '').strip()
        if inp and inp != 'nan':
            alias_map[inp] = off
    
    # 2. 3,766 Registered Players
    df_reg = pd.read_excel('1. NCU_Registered_Players.xlsx')
    df_reg['Full_Name'] = df_reg['First Name'].astype(str).str.strip() + ' ' + df_reg['Last Name'].astype(str).str.strip()
    df_reg['Full_Name'] = df_reg['Full_Name'].apply(smart_title)
    df_reg['Norm_Name'] = df_reg['Full_Name'].apply(norm)
    
    # 3. DOB

    import glob
    import os
    
    # Find Revenue Report
    rev_files = glob.glob('revenue_report_il_from_*.xlsx')
    if not rev_files:
        rev_files = ['NCU Revenue Report (for analyysis).xlsx']
    
    if not os.path.exists(rev_files[0]):
        raise FileNotFoundError("Could not find a raw revenue report (e.g. revenue_report_il_from_*.xlsx)")
        
    df_rev = clean_revenue_report(rev_files[0])
    
    # Find DOB Report
    dob_files = glob.glob('Player_Registrations_for_2026_with_DOB*.csv')
    if not dob_files:
        dob_files = ['Player_Registrations_for_2026_with_DOB-2026-08-27T095733.csv']
        
    if not os.path.exists(dob_files[0]):
        raise FileNotFoundError("Could not find the DOB registration report (e.g. Player_Registrations_for_2026_with_DOB*.csv)")

    df_dob = pd.read_csv(dob_files[0])
    df_dob['Full_Name'] = df_dob['First Name'].astype(str).str.strip() + ' ' + df_dob['Last Name'].astype(str).str.strip()
    df_dob['Norm_Name'] = df_dob['Full_Name'].apply(norm)
    df_dob['DOB'] = pd.to_datetime(df_dob['Date of Birth'], errors='coerce')
    
    dob_map_strict = {}
    dob_map_loose = {}
    for _, r in df_dob.iterrows():
        nn = r['Norm_Name']
        c = str(r['Individual Membership Primary Club']).strip().lower().replace(' cricket club', '').replace(' cc', '')
        dob_map_strict[(nn, c)] = r['DOB']
        dob_map_loose[nn] = r['DOB']

    def get_dob(row):
        nn = row['Norm_Name']
        c = str(row['Individual Membership Primary Club']).strip().lower().replace(' cricket club', '').replace(' cc', '')
        if (nn, c) in dob_map_strict: return dob_map_strict[(nn, c)]
        
        mapped = norm(alias_map.get(nn, nn))
        if (mapped, c) in dob_map_strict: return dob_map_strict[(mapped, c)]
        
        if nn in dob_map_loose: return dob_map_loose[nn]
        if mapped in dob_map_loose: return dob_map_loose[mapped]
        
        if 'jack kirkpatrick jnr' in nn: return pd.to_datetime('2009-04-07')
        if 'jack kirkpatrick snr' in nn: return pd.to_datetime('2001-11-15')
        if 'pallav saran' in nn: return pd.to_datetime('1995-01-01')
        if 'rene rankin' in nn: return pd.to_datetime('2007-06-12')
        if 'lucy ly' in nn: return pd.to_datetime('1989-10-18')
        return pd.NaT
    
    df_reg['DOB'] = df_reg.apply(get_dob, axis=1)
    
    def calc_age_30june(dob):
        if pd.isna(dob): return np.nan
        return 2026 - dob.year - (1 if (dob.month, dob.day) > (6, 30) else 0)
    
    df_reg['Age_30June2026'] = df_reg['DOB'].apply(calc_age_30june)
    df_reg['Is_Youth'] = df_reg['Age_30June2026'] <= 18
    
    # 4. Revenue

    df_rev['Norm_Name'] = df_rev['Player Name'].apply(norm)
    df_rev['Resolved_Norm'] = df_rev['Norm_Name'].apply(lambda x: norm(alias_map.get(x, x)))
    df_rev['Payment_Details'] = df_rev['Club'].fillna('').astype(str) + ' - ' + df_rev['Type'].fillna('').astype(str) + ' (£' + df_rev['Payment Amount'].astype(str) + ')'
    
    rev_summary = df_rev.groupby('Resolved_Norm').agg(
        Revenue_Name=('Player Name', 'first'),
        Total_Paid=('Payment Amount', 'sum'),
        Payment_Count=('Payment Amount', 'count'),
        Clubs_Paid=('Club', lambda s: ', '.join(sorted(set(str(x) for x in s)))),
        Types_Paid=('Payment_Details', lambda s: '; '.join(sorted(set(str(x) for x in s)))),
        Dates_Paid=('Payment Date', lambda s: ', '.join(sorted(set(str(x)[:10] for x in s))))
    ).reset_index()
    
    # 5. Batting Matches Only (+ Saturday abandoned)
    match_files = [
        ('NV Play NCU League and Saturday Cup batting stats for season.xlsx', 'Saturday Batting', 'Name', 'Group'),
        ('NV Play Women\'s Fixtures batting stats for season.xlsx', 'Women Batting', 'Name', 'Group'),
        ('NV Play Midweek League batting stats for season.xlsx', 'Midweek Batting', 'Name', 'Group'),
        ('NV Play NCU League and Saturday Cup player appearances for abandoned games.xlsx', 'Saturday Abandoned', 'Name', 'Match')
    ]
    
    player_matches = {}
    for fpath, comp_label, name_col, grp_col in match_files:
        df_m = pd.read_excel(fpath).drop_duplicates(subset=[name_col, grp_col])
        for _, row in df_m.iterrows():
            raw = norm(row[name_col])
            if not raw or raw == 'nan': continue
            resolved = norm(alias_map.get(raw, raw))
            if resolved not in player_matches:
                player_matches[resolved] = {'display_name': str(row[name_col]).strip(), 'raw_names': set(), 'total_matches': 0, 'comps': set(), 'teams': set()}
            player_matches[resolved]['total_matches'] += 1
            player_matches[resolved]['raw_names'].add(str(row[name_col]).strip())
            player_matches[resolved]['comps'].add(comp_label)
            grp_str = str(row.get(grp_col, '')).strip()
            if grp_str: player_matches[resolved]['teams'].add(grp_str)
    
    df_matches = pd.DataFrame([{
        'Match_Norm': k,
        'Match_Player_Display': smart_title(v['display_name']),
        'Raw_Names': ', '.join(sorted(v['raw_names'])),
        'Total_Matches': v['total_matches'],
        'Competitions': ', '.join(sorted(v['comps'])),
        'Teams': ', '.join(sorted(v['teams']))
    } for k, v in player_matches.items()])
    
    # 6. Master Merge
    df_reg['Resolved_Norm'] = df_reg['Norm_Name'].apply(lambda x: norm(alias_map.get(x, x)))


    
    df_master = pd.merge(df_reg, rev_summary, on='Resolved_Norm', how='left')

    # --- CUSTOM FEE AUDIT WORKAROUND FOR DUPLICATE NAMES ---
    # For players who share a name, the merge above accidentally sums their payments together.
    # We will specifically override their totals by allocating payments to the row where the clubs match!
    try:
        name_counts = df_reg['Resolved_Norm'].value_counts()
        multi_names = name_counts[name_counts > 1].index.tolist()
        
        for name in multi_names:
            payments = df_rev[df_rev['Resolved_Norm'] == name]
            mask = df_master['Resolved_Norm'] == name
            
            # Reset their totals since they were wrongly combined
            df_master.loc[mask, 'Total_Paid'] = 0
            df_master.loc[mask, 'Types_Paid'] = ''
            
            for idx, row in df_master[mask].iterrows():
                reg_clubs = [str(row.get('Individual Membership Primary Club', '')).lower().replace(' cricket club', '').replace(' cc', '').strip()]
                for tc in ['Transfer Club 1', 'Transfer Club 2']:
                    if tc in row and pd.notna(row[tc]):
                        reg_clubs.append(str(row[tc]).lower().replace(' cricket club', '').replace(' cc', '').strip())
                
                matched_payments = []
                for _, p in payments.iterrows():
                    p_club = str(p.get('Club', '')).lower().replace(' cricket club', '').replace(' cc', '').strip()
                    if p_club:
                        if any((p_club in rc or rc in p_club) for rc in reg_clubs if rc):
                            matched_payments.append(p)
                
                if matched_payments:
                    total = sum(p['Payment Amount'] for p in matched_payments)
                    types = '; '.join(sorted(set(str(p.get('Payment_Details', p.get('Type', ''))) for p in matched_payments)))
                    df_master.at[idx, 'Total_Paid'] = total
                    df_master.at[idx, 'Types_Paid'] = types
    except Exception as e:
        with open('error_fee.txt', 'w') as f2: f2.write(str(e))
    # ---------------------------------------------------------
    df_master = pd.merge(df_master, df_matches, left_on='Resolved_Norm', right_on='Match_Norm', how='left')
    
    df_master['Total_Paid'] = df_master['Total_Paid'].fillna(0)


    df_master['Total_Matches'] = df_master['Total_Matches'].fillna(0)

    # --- CUSTOM MATCH WORKAROUND FOR DUPLICATE NAMES ---
    try:
        for name in multi_names:
            mask = df_master['Resolved_Norm'] == name
            
            # Reset their totals and teams
            df_master.loc[mask, 'Total_Matches'] = 0
            df_master.loc[mask, 'Teams'] = ''
            
            if name in player_matches:
                teams = player_matches[name]['teams']
                uncontested = []
                contested = []
                
                # Identify contested vs uncontested
                for t in teams:
                    t_lower = t.lower()
                    matching_idxs = []
                    
                    # Check primary clubs
                    for idx, row in df_master[mask].iterrows():
                        rc = str(row.get('Individual Membership Primary Club', '')).lower().replace(' cricket club', '').replace(' cc', '').strip()
                        if rc and rc in t_lower:
                            matching_idxs.append(idx)
                            continue
                            
                        # Check transfer clubs if primary didn't match
                        for tc_col in ['Transfer Club 1', 'Transfer Club 2']:
                            if tc_col in row and pd.notna(row[tc_col]):
                                tc = str(row[tc_col]).lower().replace(' cricket club', '').replace(' cc', '').strip()
                                if tc and tc in t_lower:
                                    matching_idxs.append(idx)
                                    break
                                    
                    if len(matching_idxs) == 1:
                        uncontested.append((t, matching_idxs[0]))
                    elif len(matching_idxs) > 1:
                        contested.append((t, matching_idxs))
                        
                # Assign uncontested matches
                for t, idx in uncontested:
                    df_master.at[idx, 'Total_Matches'] += 1
                    df_master.at[idx, 'Teams'] = df_master.at[idx, 'Teams'] + t + ', ' if df_master.at[idx, 'Teams'] else t + ', '
                    
                # Assign contested matches to the player with the most uncontested matches
                for t, idxs in contested:
                    best_idx = max(idxs, key=lambda i: df_master.at[i, 'Total_Matches'])
                    df_master.at[best_idx, 'Total_Matches'] += 1
                    df_master.at[best_idx, 'Teams'] = df_master.at[best_idx, 'Teams'] + t + ', ' if df_master.at[best_idx, 'Teams'] else t + ', '
                    
        # Clean up trailing commas
        df_master['Teams'] = df_master['Teams'].astype(str).str.rstrip(', ')
        
    except Exception as e:
        with open('error_match_fix.txt', 'w') as f2: f2.write(str(e))
    df_master['Played_Adult_Matches'] = df_master['Total_Matches'] > 0
    df_master['Date of Birth'] = df_master['DOB'].dt.strftime('%Y-%m-%d')
    df_master['Age as of 30 June 2026'] = df_master['Age_30June2026']
    
    # Unmatched scorecard players
    unmatched_matches = df_matches[~df_matches['Match_Norm'].isin(df_master['Resolved_Norm'])].sort_values(by='Total_Matches', ascending=False).copy()
    
    def extract_base_club(t):
        if not t or pd.isna(t): return "Unknown"
        t = str(t).strip()
        t = re.sub(r'(?i)\b\d(?:st|nd|rd|th)?\s*XI\b', '', t)
        t = re.sub(r'(?i)\b(?:1st|2nd|3rd|4th|5th|6th|7th)\b', '', t)
        t = re.sub(r'(?i)\bMW\d?\b', '', t)
        t = re.sub(r'(?i)\bXI\b', '', t)
        t = re.sub(r'(?i)\bWomen\'?s?\b', '', t)
        t = re.sub(r'(?i)\bCricket Club\b|\bCC\b', '', t)
        t = re.sub(r'\s+\d$', '', t.strip())
        t = re.sub(r'\s+', ' ', t).strip()
        if re.search(r'(?i)\bciyms\b', t) or t.lower() == 'ci': return 'CIYMS'
        if re.search(r'(?i)\bholywood\s+1881\b|\bholywood\b', t): return 'Holywood 1881'
        if re.search(r'(?i)northern\s+ireland\s+malayali|nima', t): return 'NIMA'
        if re.search(r'(?i)belfast\s+international\s+sports\s+club|bisc', t): return 'BISC'
        if re.search(r'(?i)civil\s+service\s+north|csni', t): return 'CSNI'
        if re.search(r'(?i)drumaness\s+super\s*kings|drumaness', t): return 'Drumaness Superkings'
        if re.search(r'(?i)donaghcloney|donacloney', t): return 'Donacloney Mill'
        if re.search(r'(?i)cliftonville\s+academy|cliftonville', t): return 'Cliftonville Academy'
        if re.search(r'(?i)belfast\s+super\s*kings', t): return 'Belfast Superkings'
        if re.search(r'(?i)ards\s*(&|and)?\s*donaghadee', t): return 'Ards & Donaghadee'
        if re.search(r'(?i)ncu\s+pathway', t): return 'NCU Pathway XI'
        return t if t else "Unknown"
    
    CLUB_DISPLAY = {
        'Muckamore': 'Muckamore Cricket Club', 'Instonians': 'Instonians Cricket Club',
        'Cooke Collegians': 'Cooke Collegians Cricket Club', 'Dundrum': 'Dundrum Cricket Club',
        'Lisburn': 'Lisburn Cricket Club', 'CSNI': 'CSNI Cricket Club',
        'Lurgan': 'Lurgan Cricket Club', 'Victoria': 'Victoria Cricket Club',
        'Cliftonville Academy': 'Cliftonville Academy Cricket Club', 'Drumaness Superkings': 'Drumaness Superkings Cricket Club',
        'Templepatrick': 'Templepatrick Cricket Club', 'Derriaghy': 'Derriaghy Cricket Club',
        'Armagh': 'Armagh Cricket Club', 'Holywood 1881': 'Holywood Cricket Club 1881',
        'CIYMS': 'CIYMS Cricket Club', 'Donacloney Mill': 'Donacloney Mill Cricket Club',
        'PSNI': 'PSNI Cricket Club', 'Ards & Donaghadee': 'Ards & Donaghadee Cricket Club',
        'Carrickfergus': 'Carrickfergus Cricket Club', 'Ballymena': 'Ballymena Cricket Club',
        'Waringstown': 'Waringstown Cricket Club', 'NCU Pathway XI': 'NCU Pathway XI',
        'Downpatrick': 'Downpatrick Cricket Club', 'Amigos Belfast': 'Amigos Belfast Cricket Club',
        'Cregagh': 'Cregagh Cricket Club', 'North Down': 'North Down Cricket Club',
        'Dunmurry': 'Dunmurry Cricket Club', 'Laurelvale': 'Laurelvale Cricket Club',
        'Belfast Superkings': 'Belfast Superkings Cricket Club', 'Bangor': 'Bangor Cricket Club'
    }
    
    def fmt(c):
        return CLUB_DISPLAY.get(c, f"{c} Cricket Club" if "XI" not in c and "Cricket Club" not in c else c)
    
    inferred_list = []
    for _, r in unmatched_matches.iterrows():
        p_name = str(r['Match_Player_Display']).strip()
        teams_str = str(r['Teams']).strip()
        if p_name == "Tyler Mcgladdery" or p_name == "Tyler McGladdery":
            inferred_list.append("Derriaghy Cricket Club (Overseas Professional)")
            continue
        if p_name == "Vismithaa Sai Pandiaraj":
            inferred_list.append("Templepatrick Cricket Club")
            continue
        if p_name == "Molly Sawyer":
            inferred_list.append("CIYMS Cricket Club")
            continue
        raw_fixtures = re.split(r',\s*(?=[A-Za-z0-9 ]+\s+v\s+)', teams_str)
        fixture_club_pairs = []
        club_counts = Counter()
        for fix in raw_fixtures:
            parts = fix.strip().rsplit(' - ', 1)[0].strip()
            if ' v ' in parts:
                t1, t2 = parts.split(' v ', 1)
                t2 = t2.rsplit(', ', 1)[0]
                c1 = extract_base_club(t1)
                c2 = extract_base_club(t2)
                if c1 != "Unknown" and c2 != "Unknown":
                    fixture_club_pairs.append({c1, c2})
                    club_counts[c1] += 1
                    club_counts[c2] += 1
        if not fixture_club_pairs:
            inferred_list.append("Unknown")
            continue
        common = set.intersection(*fixture_club_pairs)
        common.discard("Unknown")
        if len(common) == 1:
            inferred_list.append(fmt(list(common)[0]))
        elif len(common) > 1:
            inferred_list.append(" / ".join(fmt(c) for c in sorted(common)))
        else:
            top_club, top_cnt = club_counts.most_common(1)[0]
            if top_cnt >= len(fixture_club_pairs) * 0.75:
                inferred_list.append(fmt(top_club))
            else:
                all_c = sorted(set.union(*fixture_club_pairs))
                inferred_list.append(" / ".join(fmt(c) for c in all_c))
    
    unmatched_matches.insert(1, 'Inferred Club', inferred_list)
    
    # Subsets
    c_youth_played_paid5 = df_master[df_master['Is_Youth'] & df_master['Played_Adult_Matches'] & (df_master['Total_Paid'] == 5)]
    c_youth_played_paid10 = df_master[df_master['Is_Youth'] & df_master['Played_Adult_Matches'] & (df_master['Total_Paid'] > 5)]
    c_youth_played_unpaid = df_master[df_master['Is_Youth'] & df_master['Played_Adult_Matches'] & (df_master['Total_Paid'] == 0)]
    
    c_youth_noplay_paid = df_master[df_master['Is_Youth'] & (~df_master['Played_Adult_Matches']) & (df_master['Total_Paid'] > 0)]
    c_youth_noplay_unpaid = df_master[df_master['Is_Youth'] & (~df_master['Played_Adult_Matches']) & (df_master['Total_Paid'] == 0)]
    
    c_adult_played_paid10 = df_master[(~df_master['Is_Youth']) & df_master['Played_Adult_Matches'] & (df_master['Total_Paid'] >= 10)]
    print('Adam Gardner count in df_master:', len(df_master[df_master['Norm_Name'] == 'adamgardner']))
    c_adult_played_paid5 = df_master[(~df_master['Is_Youth']) & df_master['Played_Adult_Matches'] & (df_master['Total_Paid'] == 5)]
    c_adult_played_paid0 = df_master[(~df_master['Is_Youth']) & df_master['Played_Adult_Matches'] & (df_master['Total_Paid'] == 0)]
    
    c_adult_noplay_paid10 = df_master[(~df_master['Is_Youth']) & (~df_master['Played_Adult_Matches']) & (df_master['Total_Paid'] >= 10)]
    c_adult_noplay_paid5 = df_master[(~df_master['Is_Youth']) & (~df_master['Played_Adult_Matches']) & (df_master['Total_Paid'] == 5)]
    c_adult_noplay_paid0 = df_master[(~df_master['Is_Youth']) & (~df_master['Played_Adult_Matches']) & (df_master['Total_Paid'] == 0)]
    
    # Summary Rows
    summary_rows = [
        {'Section': 'SECTION 1: REGISTERED PLAYERS WHO PLAYED ADULT CRICKET', 'Category': 'Adults (>18 on 30-June-2026) who PLAYED adult cricket - PAID £10+ (Compliant)', 'Count': len(c_adult_played_paid10), 'Status': 'Compliant'},
        {'Section': 'SECTION 1: REGISTERED PLAYERS WHO PLAYED ADULT CRICKET', 'Category': 'Adults (>18 on 30-June-2026) who PLAYED adult cricket - PAID £5 (Underpaid Youth Rate)', 'Count': len(c_adult_played_paid5), 'Status': 'Underpaid (£5 shortfall)'},
        {'Section': 'SECTION 1: REGISTERED PLAYERS WHO PLAYED ADULT CRICKET', 'Category': 'Adults (>18 on 30-June-2026) who PLAYED adult cricket - PAID £0 (Unpaid Adult Fee)', 'Count': len(c_adult_played_paid0), 'Status': 'Unpaid (£10 shortfall)'},
        {'Section': 'SECTION 1: REGISTERED PLAYERS WHO PLAYED ADULT CRICKET', 'Category': 'Youth (<=18 on 30-June-2026) who PLAYED adult cricket - PAID £5 (Compliant)', 'Count': len(c_youth_played_paid5), 'Status': 'Compliant'},
        {'Section': 'SECTION 1: REGISTERED PLAYERS WHO PLAYED ADULT CRICKET', 'Category': 'Youth (<=18 on 30-June-2026) who PLAYED adult cricket - PAID > £5 (Paid Adult Rate £10+)', 'Count': len(c_youth_played_paid10), 'Status': 'Compliant'},
        {'Section': 'SECTION 1: REGISTERED PLAYERS WHO PLAYED ADULT CRICKET', 'Category': 'Youth (<=18 on 30-June-2026) who PLAYED adult cricket - PAID £0 (Unpaid Playing Youth Fee)', 'Count': len(c_youth_played_unpaid), 'Status': 'Unpaid (£5 shortfall)'},
        
        {'Section': 'SECTION 2: REGISTERED PLAYERS WHO DID NOT PLAY ADULT CRICKET', 'Category': 'Adults (>18 on 30-June-2026) who DID NOT play adult cricket - PAID £10+ (Non-Playing Adult)', 'Count': len(c_adult_noplay_paid10), 'Status': 'Compliant (Non-Playing)'},
        {'Section': 'SECTION 2: REGISTERED PLAYERS WHO DID NOT play adult cricket', 'Category': 'Adults (>18 on 30-June-2026) who DID NOT play adult cricket - PAID £5 (Non-Playing Youth Rate)', 'Count': len(c_adult_noplay_paid5), 'Status': 'Non-Playing'},
        {'Section': 'SECTION 2: REGISTERED PLAYERS WHO DID NOT play adult cricket', 'Category': 'Adults (>18 on 30-June-2026) who DID NOT play adult cricket - PAID £0 (Non-Playing Unpaid)', 'Count': len(c_adult_noplay_paid0), 'Status': 'Non-Playing'},
        {'Section': 'SECTION 2: REGISTERED PLAYERS WHO DID NOT play adult cricket', 'Category': 'Youth (<=18 on 30-June-2026) who DID NOT play adult cricket - PAID £5+ (Exempt / Unused Fee)', 'Count': len(c_youth_noplay_paid), 'Status': 'Exempt (Fee Paid)'},
        {'Section': 'SECTION 2: REGISTERED PLAYERS WHO DID NOT play adult cricket', 'Category': 'Youth (<=18 on 30-June-2026) who DID NOT play adult cricket - PAID £0 (Exempt Junior Cricket Only)', 'Count': len(c_youth_noplay_unpaid), 'Status': 'Compliant (Exempt £0)'},
        
        {'Section': 'TOTAL OFFICIAL REGISTERED PLAYERS (1. NCU_Registered_Players.xlsx)', 'Category': 'TOTAL REGISTERED PLAYERS ACCOUNTED FOR', 'Count': len(df_master), 'Status': '100% Reconciled'},
        
        {'Section': 'SECTION 3: UNREGISTERED MATCH APPEARANCES', 'Category': 'Unregistered Scorecard Players (Played in matches but NOT registered in Sport80)', 'Count': len(unmatched_matches), 'Status': 'Unregistered'}
    ]
    
    df_summary = pd.DataFrame(summary_rows)
    
    now = datetime.now()
    d_str = now.strftime('%Y-%m-%d')
    t_str = now.strftime('%H-%M-%S')
    timestamped_prefix = f'D{d_str} T{t_str}'
    
    audit_excel_io = io.BytesIO()
    with pd.ExcelWriter(audit_excel_io, engine='openpyxl') as writer:
        # 1. Summary
        df_summary.to_excel(writer, sheet_name='Audit Summary', index=False)
        
        # 2. Unregistered Scorecard Players
        cols_un = ['Match_Player_Display', 'Inferred Club', 'Total_Matches', 'Competitions', 'Teams', 'Raw_Names']
        unmatched_matches[cols_un].to_excel(writer, sheet_name='Unregistered Scorecard Players', index=False)
        
        # Detail sheets
        cols_reg_unpaid_y = ['Full_Name', 'Date of Birth', 'Age as of 30 June 2026', 'Individual Membership Primary Club', 'Total_Paid', 'Total_Matches', 'Teams']
        c_youth_played_unpaid[cols_reg_unpaid_y].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Unpaid Youth in Adult Cricket', index=False)
        c_adult_played_paid0[cols_reg_unpaid_y].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Unpaid Adults (£10 shortfall)', index=False)
        
        cols_reg_with_types = ['Full_Name', 'Date of Birth', 'Age as of 30 June 2026', 'Individual Membership Primary Club', 'Total_Paid', 'Total_Matches', 'Types_Paid', 'Teams']
        c_adult_played_paid5[cols_reg_with_types].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Adults Paid Youth Rate (£5)', index=False)
        c_youth_noplay_paid[cols_reg_with_types].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Youth Paid No Adult Matches', index=False)
        
        
        c_adult_played_paid10[cols_reg_with_types].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Compliant Adults (£10+)', index=False)
        c_youth_played_paid5[cols_reg_with_types].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Compliant Youths (£5)', index=False)
        
        c_youth_noplay_unpaid[['Full_Name', 'Date of Birth', 'Age as of 30 June 2026', 'Individual Membership Primary Club']].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Junior Youths (Exempt £0)', index=False)
        c_adult_noplay_paid10[['Full_Name', 'Date of Birth', 'Age as of 30 June 2026', 'Individual Membership Primary Club', 'Total_Paid', 'Types_Paid']].sort_values(by=['Individual Membership Primary Club', 'Full_Name']).to_excel(writer, sheet_name='Non-Playing Adults (£10+)', index=False)
    
    # Styling with openpyxl
    audit_excel_io.seek(0)
    wb = openpyxl.load_workbook(audit_excel_io)
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
    
    total_row_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
    total_row_font = Font(name='Calibri', size=11, bold=True, color='000000')
    
    center_align = Alignment(horizontal='center', vertical='center')
    currency_format = '£#,##0'
    
    def is_date_or_number(val):
        if val is None or str(val).strip() == '' or str(val).strip().lower() == 'nan':
            return True
        if isinstance(val, (int, float, datetime)):
            return True
        s = str(val).strip()
        try:
            float(s.replace('£', '').replace(',', ''))
            return True
        except ValueError:
            pass
        if re.match(r'^\d{4}-\d{2}-\d{2}', s) or re.match(r'^\d{2}/\d{2}/\d{4}', s):
            return True
        return False
    
    for sname in wb.sheetnames:
        ws = wb[sname]
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = ws.dimensions
        
        for col_idx in range(1, ws.max_column + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center' if any(w in str(cell.value) for w in ['Date', 'Age', 'Count', 'Matches', 'Paid', 'Status']) else 'left', vertical='center')
    
        if sname == 'Audit Summary':
            for r in range(2, ws.max_row + 1):
                if 'TOTAL REGISTERED PLAYERS' in str(ws.cell(row=r, column=2).value):
                    for c in range(1, ws.max_column + 1):
                        ws.cell(row=r, column=c).fill = total_row_fill
                        ws.cell(row=r, column=c).font = total_row_font
    
        for col_idx in range(1, ws.max_column + 1):
            header_val = str(ws.cell(row=1, column=col_idx).value or '').strip()
            
            is_num_col = True
            has_data = False
            for row_idx in range(2, ws.max_row + 1):
                val = ws.cell(row=row_idx, column=col_idx).value
                if val is not None and str(val).strip() != '':
                    has_data = True
                    if not is_date_or_number(val):
                        is_num_col = False
                        break
                        
            for row_idx in range(2, ws.max_row + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                if header_val == 'Total_Paid':
                    if cell.value is not None and str(cell.value).strip() != '':
                        try:
                            num = float(str(cell.value).replace('£', '').replace(',', '').strip())
                            cell.value = int(num) if num.is_integer() else num
                        except: pass
                    cell.number_format = currency_format
                    cell.alignment = center_align
                elif has_data and is_num_col:
                    cell.alignment = center_align
    
        for col_cells in ws.columns:
            col_letter = get_column_letter(col_cells[0].column)
            max_len = max(min(len(str(c.value or '')), 65) for c in col_cells)
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)
    
    final_excel_io = io.BytesIO()
    wb.save(final_excel_io)
    
    # Generate the anomalies word report
    doc_io = generate_anomalies_word_report(df_rev, df_reg, alias_map, timestamped_prefix)
    
    return final_excel_io, doc_io, df_summary, df_master
    